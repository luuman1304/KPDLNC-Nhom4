from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


DOCX = Path("Document/Bai_Bao_Nghien_Cuu_M5_Cau_Truc_Chuan_Hoc_Thuat.docx")
METHOD_FIGURE = Path("outputs_method_figures/method_workflow.png")
EDA_MANIFEST = Path("outputs_eda_figures/eda_figure_manifest.csv")
RESULT_MANIFEST = Path("outputs_research_figures/figure_manifest.csv")
APPENDIX_TITLE = "Phụ lục A. Toàn bộ biểu đồ nghiên cứu"


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.size = Pt(10)


def add_figure(doc: Document, title: str, path: Path, caption: str, width: float = 6.3) -> None:
    doc.add_heading(title, level=2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, f"Chú thích: {caption}")


def remove_existing_appendix(doc: Document) -> None:
    body = doc._body._element
    start = None
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == APPENDIX_TITLE:
            start = idx
            break
    if start is None:
        return
    for paragraph in doc.paragraphs[start:]:
        body.remove(paragraph._element)


def main() -> None:
    doc = Document(DOCX)
    remove_existing_appendix(doc)

    doc.add_page_break()
    heading = doc.add_heading(APPENDIX_TITLE, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro = doc.add_paragraph(
        "Phụ lục này tập hợp toàn bộ biểu đồ được sử dụng trong nghiên cứu, bao gồm sơ đồ phương pháp, "
        "biểu đồ phân tích khám phá dữ liệu và biểu đồ kết quả thực nghiệm. Một số biểu đồ có thể đã xuất hiện ở phần nội dung chính; "
        "việc gom lại trong phụ lục giúp người đọc kiểm tra trực quan toàn bộ pipeline và kết quả."
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    add_figure(
        doc,
        "Hình A1. Sơ đồ quy trình phương pháp nghiên cứu",
        METHOD_FIGURE,
        "Sơ đồ mô tả luồng từ dữ liệu M5, EDA, feature engineering, clustering, huấn luyện A0/B1/C, recursive forecasting đến đánh giá accuracy, stability và robustness.",
        width=6.6,
    )

    doc.add_heading("A.1. Biểu đồ phân tích khám phá dữ liệu", level=1)
    eda = pd.read_csv(EDA_MANIFEST)
    for _, row in eda.iterrows():
        fig_no = int(row["figure"])
        add_figure(
            doc,
            f"Hình EDA {fig_no}. {row['title']}",
            Path(row["path"]),
            str(row["caption"]),
        )

    doc.add_heading("A.2. Biểu đồ kết quả thực nghiệm", level=1)
    results = pd.read_csv(RESULT_MANIFEST)
    for _, row in results.iterrows():
        fig_no = int(row["figure"])
        add_figure(
            doc,
            f"Hình R{fig_no}. {row['title']}",
            Path(row["path"]),
            str(row["caption"]),
        )

    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
