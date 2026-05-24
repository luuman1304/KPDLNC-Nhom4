from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


INPUT_DOCX = Path("Document/Bai_Viet_Nghien_Cuu_Khoa_Hoc_Chi_Tiet_Hoc_Thuat_M5.docx")
OUTPUT_DOCX = Path("Document/Bai_Viet_Nghien_Cuu_Khoa_Hoc_Chi_Tiet_Hoc_Thuat_M5_Co_Bieu_Do.docx")
EDA_MANIFEST = Path("outputs_eda_figures/eda_figure_manifest.csv")
RESULT_MANIFEST = Path("outputs_research_figures/figure_manifest.csv")


def main() -> None:
    doc = Document(INPUT_DOCX)
    doc.add_page_break()
    heading = doc.add_heading("Phụ lục biểu đồ phân tích khám phá dữ liệu và kết quả nghiên cứu", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro = doc.add_paragraph(
        "Các biểu đồ trong phụ lục được sinh trực tiếp từ dữ liệu EDA và kết quả full-scale đã chạy. "
        "Mục đích của phụ lục là trực quan hóa đặc điểm dữ liệu trước mô hình hóa, sau đó minh họa các kết luận chính về accuracy, stability, ablation, phân cụm, overfitting và feature importance."
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if EDA_MANIFEST.exists():
        eda_heading = doc.add_heading("Phụ lục A. Biểu đồ phân tích khám phá dữ liệu", level=1)
        eda_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        eda_intro = doc.add_paragraph(
            "Các hình EDA dưới đây mô tả mức độ intermittent demand, tỷ lệ ngày không bán, cấu trúc ADI-CV2, khác biệt theo hierarchy và đặc điểm giá. "
            "Những biểu đồ này được dùng để kiểm tra giả định dữ liệu và giải thích vì sao nghiên cứu cần cluster-aware global forecasting, objective phù hợp với dữ liệu không âm nhiều giá trị 0, và metric stability có scale floor."
        )
        eda_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        append_manifest(doc, pd.read_csv(EDA_MANIFEST), "Hình EDA")

    result_heading = doc.add_heading("Phụ lục B. Biểu đồ kết quả mô hình", level=1)
    result_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    result_intro = doc.add_paragraph(
        "Các hình kết quả trực quan hóa so sánh A0, B1 và C, tác động ablation, profile cụm, overfitting gap và feature importance."
    )
    result_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    append_manifest(doc, pd.read_csv(RESULT_MANIFEST), "Hình")

    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


def append_manifest(doc: Document, manifest: pd.DataFrame, label: str) -> None:
    for _, row in manifest.iterrows():
        title = str(row["title"])
        caption = str(row["caption"])
        path = Path(row["path"])
        fig_no = int(row["figure"])

        doc.add_heading(f"{label} {fig_no}. {title}", level=2)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(6.3))

        cap = doc.add_paragraph(f"Chú thích: {caption}")
        cap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in cap.runs:
            run.font.size = Pt(10)


if __name__ == "__main__":
    main()
