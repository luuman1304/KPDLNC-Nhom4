from __future__ import annotations

import zipfile
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "Document" / "Nhom3.dotx"
TARGET = ROOT / "Document" / "Bai_Bao_Nghien_Cuu_M5_Cau_Truc_Chuan_Hoc_Thuat.docx"
OUTPUT = ROOT / "Document" / "Bai_Bao_Nghien_Cuu_M5_Nhom3_Format.docx"


DOCX_MAIN_CT = (
    b"application/vnd.openxmlformats-officedocument."
    b"wordprocessingml.document.main+xml"
)
DOTX_MAIN_CT = (
    b"application/vnd.openxmlformats-officedocument."
    b"wordprocessingml.template.main+xml"
)


def convert_dotx_to_docx(source: Path, destination: Path) -> None:
    """Convert a Word template package to a document package without changing content."""
    with zipfile.ZipFile(source) as zin, zipfile.ZipFile(
        destination, "w", zipfile.ZIP_DEFLATED
    ) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "[Content_Types].xml":
                data = data.replace(DOTX_MAIN_CT, DOCX_MAIN_CT)
            zout.writestr(item, data)


def validate_document(path: Path) -> dict[str, object]:
    doc = Document(path)
    headings = [
        p.text.strip()
        for p in doc.paragraphs
        if p.style.name.startswith("Heading") and p.text.strip()
    ]
    captions = [
        p.text.strip()
        for p in doc.paragraphs
        if p.style.name == "Caption" and p.text.strip()
    ]
    return {
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "figures": sum(1 for p in captions if p.startswith("Hình ")),
        "captions": len(captions),
        "headings": headings,
    }


def main() -> None:
    if not TEMPLATE.exists():
        raise FileNotFoundError(TEMPLATE)

    convert_dotx_to_docx(TEMPLATE, OUTPUT)

    # Save over the working article so the requested file follows Nhom3.dotx.
    TARGET.write_bytes(OUTPUT.read_bytes())

    summary = validate_document(TARGET)
    print(f"Updated: {TARGET}")
    print(f"Also saved: {OUTPUT}")
    print(
        "Counts:",
        f"paragraphs={summary['paragraphs']}",
        f"tables={summary['tables']}",
        f"figures={summary['figures']}",
        f"captions={summary['captions']}",
    )
    print("Headings:")
    for heading in summary["headings"]:
        print(f"- {heading}")


if __name__ == "__main__":
    main()
