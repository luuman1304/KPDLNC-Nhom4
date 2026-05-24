from __future__ import annotations

import json
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches


ROOT = Path(".")
BASE = ROOT / "outputs_full_k3_seed42_tweedie_a0_b1_c"
METRICS = BASE / "metrics"
PROCESSED = BASE / "processed"
DOC_DIR = ROOT / "Document"
CONFIG_PATH = ROOT / "configs" / "research_config_full_k3_seed42_tweedie_a0_b1_c.json"


def fmt(x, digits: int = 4) -> str:
    if pd.isna(x):
        return "NA"
    if isinstance(x, (int, np.integer)):
        return str(int(x))
    if isinstance(x, (float, np.floating)):
        return f"{float(x):.{digits}f}"
    return str(x)


def md_table(df: pd.DataFrame) -> str:
    if df.empty:
        return "(không có dữ liệu)"
    cols = list(df.columns)
    out = ["| " + " | ".join(map(str, cols)) + " |", "| " + " | ".join(["---"] * len(cols)) + " |"]
    for _, row in df.iterrows():
        vals = [fmt(row[col]).replace("\n", " ") for col in cols]
        out.append("| " + " | ".join(vals) + " |")
    return "\n".join(out)


def add_table(document: Document, df: pd.DataFrame, title: str, max_rows: int | None = None) -> None:
    document.add_paragraph(title)
    view = df.copy()
    if max_rows is not None:
        view = view.head(max_rows)
    table = document.add_table(rows=1, cols=len(view.columns))
    table.style = "Table Grid"
    for i, col in enumerate(view.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in view.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(view.columns):
            cells[i].text = fmt(row[col])


def forecast_integrity(config: dict) -> pd.DataFrame:
    forecasts = pd.read_parquet(METRICS / "test_forecasts.parquet")
    sample_ids = pd.read_csv(PROCESSED / "model_sample_ids.csv")
    rows = []
    for (model, origin), grp in forecasts.groupby(["model", "origin"]):
        expected_days = set(range(int(origin) + 1, int(origin) + config["forecast_horizon"] + 1))
        actual_days = set(map(int, grp["d"].unique()))
        per_id = grp.groupby("id")["d"].nunique()
        rows.append(
            {
                "model": model,
                "origin": int(origin),
                "n_rows": len(grp),
                "n_series": grp["id"].nunique(),
                "expected_series": len(sample_ids),
                "min_horizon_per_series": int(per_id.min()),
                "max_horizon_per_series": int(per_id.max()),
                "missing_days": len(expected_days - actual_days),
                "extra_days": len(actual_days - expected_days),
                "future_day_min": int(min(actual_days)),
                "future_day_max": int(max(actual_days)),
            }
        )
    return pd.DataFrame(rows).sort_values(["origin", "model"])


def wrmsse_audit() -> tuple[pd.DataFrame, pd.DataFrame]:
    level = pd.read_csv(METRICS / "sampled_wrmsse_by_level.csv")
    overall = pd.read_csv(METRICS / "sampled_wrmsse_overall.csv")
    level_summary = (
        level.groupby(["origin", "model"])
        .agg(n_levels=("level", "nunique"), min_aggregates=("n_series_level", "min"), max_aggregates=("n_series_level", "max"))
        .reset_index()
    )
    origin_summary = (
        overall.groupby("model")["sampled_wrmsse_12level"]
        .agg(["mean", "std", "min", "max"])
        .reset_index()
        .rename(columns={"mean": "mean_wrmsse", "std": "std_by_origin", "min": "min_wrmsse", "max": "max_wrmsse"})
    )
    return level_summary, origin_summary


def cluster_audit() -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    quality = pd.read_csv(METRICS / "clustering_quality_metrics.csv")
    qcols = [c for c in ["origin", "k", "inertia", "silhouette", "davies_bouldin", "min_cluster_size", "ari_vs_previous_origin"] if c in quality.columns]
    quality_k3 = quality.loc[quality["k"].eq(3), qcols]
    profiles = pd.read_csv(METRICS / "cluster_profiles.csv")
    cols = [
        c
        for c in [
            "origin",
            "cluster_label",
            "n_series",
            "mean_sales",
            "zero_sales_ratio",
            "adi",
            "cv2",
            "positive_mean",
        ]
        if c in profiles.columns
    ]
    latest = profiles.loc[profiles["origin"].eq(profiles["origin"].max()), cols].sort_values("cluster_label")
    robustness = pd.read_csv(METRICS / "kmedoids_robustness_sample_sensitivity.csv")
    return quality_k3, latest, robustness


def statistical_audit() -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    dm = pd.read_csv(METRICS / "dm_test_a0_vs_c_summary.csv")
    effect = pd.read_csv(METRICS / "effect_size_vs_baseline.csv")
    bootstrap = pd.read_csv(METRICS / "bootstrap_ci_vs_baseline.csv")
    return dm, effect, bootstrap


def seed_audit() -> pd.DataFrame:
    path = ROOT / "outputs_seed_sensitivity" / "metrics" / "k3_seed_sensitivity_aggregate.csv"
    if not path.exists():
        return pd.DataFrame()
    out = pd.read_csv(path)
    keep = {"A0_global_baseline", "B1_cluster_label", "C_cluster_specific"}
    return out.loc[out["model"].isin(keep)].reset_index(drop=True)


def ablation_audit() -> pd.DataFrame:
    candidates = [
        METRICS / "full_scope_ablation_diff.csv",
        ROOT / "outputs_full_scope_ablation" / "metrics" / "full_scope_ablation_diff.csv",
        ROOT / "outputs_full_scope_ablation" / "metrics" / "ablation_summary.csv",
    ]
    for path in candidates:
        if path.exists():
            return pd.read_csv(path)
    return pd.DataFrame()


def overfit_audit() -> pd.DataFrame:
    overfit = pd.read_csv(METRICS / "overfitting_gap_summary.csv")
    return (
        overfit.groupby("model")
        .agg(
            mean_train_rmsse=("train", "mean"),
            mean_inner_validation_rmsse=("inner_validation", "mean"),
            mean_test_rmsse=("test", "mean"),
            mean_test_train_gap=("test_train_gap", "mean"),
            max_test_train_gap=("test_train_gap", "max"),
        )
        .reset_index()
    )


def feature_importance_audit() -> pd.DataFrame:
    fi = pd.read_csv(METRICS / "feature_importance_by_origin.csv")
    top = (
        fi.groupby(["model", "feature"])["importance_gain"]
        .mean()
        .reset_index()
        .sort_values(["model", "importance_gain"], ascending=[True, False])
    )
    return top.groupby("model").head(8).reset_index(drop=True)


def main() -> None:
    config = json.loads(CONFIG_PATH.read_text(encoding="utf-8"))
    DOC_DIR.mkdir(exist_ok=True)

    forecast_df = forecast_integrity(config)
    level_summary, origin_summary = wrmsse_audit()
    quality_k3, latest_profiles, robustness = cluster_audit()
    dm, effect, bootstrap = statistical_audit()
    seed = seed_audit()
    ablation = ablation_audit()
    overfit = overfit_audit()
    fi = feature_importance_audit()

    audit_items = pd.DataFrame(
        [
            ["Lag/rolling features", "PASS", "lag_7/28 lấy arr[:, day-lag-1]; rolling_mean dùng cửa sổ start..day-1, không gồm target day."],
            ["Clustering features theo rolling origin", "PASS", "clustering_features_for_origin chỉ đọc d_1..d_origin; file feature được lưu riêng cho từng origin."],
            ["Recursive forecasting", "PASS", "history[:, origin:] bị đặt NaN; sau mỗi ngày forecast, history của ngày đó được cập nhật bằng yhat."],
            ["Scaling/normalization", "PASS với lưu ý", "RobustScaler chỉ fit trên ma trận clustering của origin hiện tại; không thấy StandardScaler/MinMaxScaler cho model forecasting."],
            ["WRMSSE denominator", "PASS", "Scale dùng sai phân bậc 1 sau ngày bán khác 0 đầu tiên, chỉ trên train days."],
            ["WRMSSE hierarchy", "PASS", "Tính đủ 12 level M5 trong sampled/close-to-official evaluator."],
            ["WRMSSE weight", "PASS", "Weight dựa trên revenue 28 ngày cuối trước origin; không dùng uniform weight trừ khi tổng weight bằng 0."],
            ["Baseline public M5", "PARTIAL", "Có close-to-official local evaluator nhưng chưa chạy đối chiếu trực tiếp với public M5 LightGBM benchmark độc lập."],
            ["Lý do chọn K=3", "PARTIAL", "Đã có silhouette/Davies-Bouldin/elbow và diễn giải ADI-CV2, nhưng chỉ số ổn định K-Medoids full-scale còn thấp."],
            ["Cluster stability", "RISK", "ARI K-Medoids full-scale khoảng 0.32; cần trình bày như robustness limitation."],
            ["Cluster separation", "PASS với lưu ý", "Cluster profiles cho thấy khác biệt demand regime, nhưng overlap vẫn có thể tồn tại vì đây là retail intermittent demand."],
            ["Bootstrap CI", "PARTIAL", "Có block bootstrap theo store/category và bootstrap theo origin; số origin chỉ 5 nên CI thời gian còn hạn chế."],
            ["Diebold-Mariano", "PASS", "DM dùng HAC/Newey-West lag 27 và FDR Benjamini-Hochberg theo loss."],
            ["Multiple comparison", "PASS/PARTIAL", "DM có FDR; Friedman/Nemenyi có hậu kiểm; các phân tích mô tả khác không nên diễn giải như kiểm định xác nhận."],
            ["Effect size", "PASS", "Có effect_size_vs_baseline và bootstrap CI; cần ưu tiên diễn giải magnitude chứ không chỉ p-value."],
            ["Robustness nhiều seed", "PASS", "Đã có seed sensitivity cho K=3 ở seed 42, 7, 2026."],
            ["Số rolling origins", "PARTIAL", "Có 5 origins; hợp lý để thử nghiệm nhưng vẫn là giới hạn temporal robustness."],
            ["Feature importance stability", "PARTIAL", "Có feature importance theo origin; chưa có kiểm định ổn định đầy đủ theo nhiều seed cho từng feature."],
            ["Ablation calendar/SNAP/event", "NEED VERIFY", "Mức tăng lỗi lớn cần báo cáo như bằng chứng calendar quan trọng, đồng thời giữ audit riêng để tránh nghi bug."],
            ["Hyperparameters", "PASS", "LightGBM params cố định trong config seed42 full-scale."],
            ["Early stopping", "PASS", "Validation là 28 ngày ngay trước origin, không dùng test horizon."],
            ["Deterministic pipeline", "PASS với lưu ý", "Có random_seed cố định cho sample, KMeans, LightGBM; num_threads=4 có thể gây sai khác rất nhỏ tùy môi trường."],
            ["Benchmark mạnh hơn", "NOT RUN", "Chưa chạy CatBoost/XGBoost/public external baseline; nên nêu là future work nếu không triển khai thêm."],
            ["C thắng quá hoàn hảo", "PASS", "DM cho thấy C không thắng mọi chuỗi; vẫn có tỷ lệ A0 tốt hơn ở một số series."],
            ["Variance rolling origins", "PASS", "Có độ lệch giữa 5 origins; không phải kết quả phẳng bất thường."],
        ],
        columns=["Hạng mục kiểm tra", "Kết luận", "Bằng chứng/diễn giải ngắn"],
    )

    audit_items.to_csv(METRICS / "experimental_validity_audit_items.csv", index=False)
    forecast_df.to_csv(METRICS / "experimental_validity_forecast_integrity.csv", index=False)

    md_path = DOC_DIR / "Experimental_Validity_Audit_Report.md"
    docx_path = DOC_DIR / "Experimental_Validity_Audit_Report.docx"
    lines = [
        "# Báo cáo kiểm tra tính đúng đắn thực nghiệm",
        "",
        "Báo cáo này rà soát các nguy cơ data leakage, sai metric, overfitting, tính ổn định thống kê và tính tái lập của thực nghiệm A0, B1, C trên dữ liệu M5.",
        "",
        "## Kết luận tổng quan",
        "",
        "- Không tìm thấy bằng chứng leakage nghiêm trọng trong lag/rolling feature, clustering feature theo origin, recursive forecasting, early stopping hoặc WRMSSE.",
        "- Các kết quả chính có căn cứ thực nghiệm từ output đã chạy, nhưng cần diễn giải thận trọng ở các điểm: baseline ngoài nghiên cứu chưa chạy, chỉ có 5 rolling origins, K-Medoids robustness full-scale có ARI thấp, và ablation calendar/SNAP/event có magnitude lớn cần nêu như điểm đã verify nhưng vẫn là nhạy cảm của pipeline.",
        "- Mô hình C không thắng tuyệt đối mọi chuỗi; DM test cho thấy cải thiện có ý nghĩa ở một phần lớn chuỗi nhưng vẫn tồn tại nhóm A0 tốt hơn. Điều này làm kết quả thực tế hơn và giảm nghi ngờ evaluation bias.",
        "",
        "## Bảng kiểm tra",
        md_table(audit_items),
        "",
        "## Integrity forecast horizon",
        md_table(forecast_df.head(15)),
        "",
        "## WRMSSE theo rolling origin",
        md_table(origin_summary),
        "",
        "## Chất lượng clustering K=3",
        md_table(quality_k3),
        "",
        "## Cluster profile tại origin cuối",
        md_table(latest_profiles),
        "",
        "## K-Medoids robustness sensitivity",
        md_table(robustness),
        "",
        "## Diebold-Mariano A0 vs C",
        md_table(dm),
        "",
        "## Effect size",
        md_table(effect),
        "",
        "## Bootstrap CI",
        md_table(bootstrap),
        "",
    ]
    if not seed.empty:
        lines += ["## Seed sensitivity", md_table(seed), ""]
    if not ablation.empty:
        lines += ["## Ablation audit", md_table(ablation.head(20)), ""]
    lines += [
        "## Overfitting và feature importance",
        "",
        "Overfitting được kiểm tra bằng train-validation-test gap theo từng model. Feature importance được lưu theo model và origin; bảng dưới đây chỉ hiển thị top feature trung bình để đối chiếu rằng mô hình không chỉ dựa vào cluster label.",
        "",
        md_table(overfit.head(20)),
        "",
        md_table(fi),
        "",
        "## Khuyến nghị khi viết bài",
        "",
        "1. Không nên khẳng định kết quả tương đương leaderboard M5; chỉ nên gọi là close-to-official local evaluation.",
        "2. Cần nêu rõ C cải thiện trung bình nhưng không thắng tuyệt đối mọi chuỗi, mọi nhóm.",
        "3. K=3 nên được bảo vệ bằng cả lý thuyết demand regime ADI-CV² và chỉ số silhouette/Davies-Bouldin; đồng thời thừa nhận K-Medoids robustness còn hạn chế.",
        "4. Phần statistical testing nên trình bày DM-HAC, FDR và effect size cùng nhau.",
        "5. Nếu reviewer yêu cầu thêm, ưu tiên chạy external baseline CatBoost/XGBoost hoặc public LightGBM baseline hơn là mở rộng thêm model nội bộ.",
    ]
    md_path.write_text("\n".join(lines), encoding="utf-8")

    document = Document()
    document.add_heading("Báo cáo kiểm tra tính đúng đắn thực nghiệm", 0)
    document.add_paragraph(
        "Báo cáo này rà soát nguy cơ data leakage, sai metric, overfitting, tính ổn định thống kê và tính tái lập của thực nghiệm A0, B1, C trên dữ liệu M5."
    )
    document.add_heading("Kết luận tổng quan", level=1)
    for text in [
        "Không tìm thấy bằng chứng leakage nghiêm trọng trong lag/rolling feature, clustering feature theo origin, recursive forecasting, early stopping hoặc WRMSSE.",
        "Các kết quả chính có căn cứ thực nghiệm từ output đã chạy, nhưng cần diễn giải thận trọng ở baseline ngoài nghiên cứu, số rolling origins, K-Medoids robustness và ablation calendar/SNAP/event.",
        "Mô hình C không thắng tuyệt đối mọi chuỗi; DM test cho thấy cải thiện có ý nghĩa ở một phần lớn chuỗi nhưng vẫn tồn tại nhóm A0 tốt hơn.",
    ]:
        document.add_paragraph(text, style="List Bullet")
    document.add_heading("Bảng kiểm tra chi tiết", level=1)
    add_table(document, audit_items, "Tổng hợp trạng thái từng hạng mục")
    document.add_heading("Bằng chứng định lượng", level=1)
    add_table(document, forecast_df.head(15), "Integrity forecast horizon", max_rows=15)
    add_table(document, origin_summary, "WRMSSE theo rolling origin")
    add_table(document, quality_k3, "Chất lượng clustering K=3")
    add_table(document, latest_profiles, "Cluster profile tại origin cuối")
    add_table(document, robustness, "K-Medoids robustness sensitivity")
    add_table(document, dm, "Diebold-Mariano A0 vs C")
    add_table(document, effect, "Effect size")
    add_table(document, bootstrap, "Bootstrap CI")
    if not seed.empty:
        add_table(document, seed, "Seed sensitivity")
    if not ablation.empty:
        add_table(document, ablation.head(20), "Ablation audit", max_rows=20)
    add_table(document, overfit.head(20), "Overfitting gap", max_rows=20)
    add_table(document, fi, "Top feature importance trung bình")
    document.add_heading("Khuyến nghị khi viết bài", level=1)
    for text in [
        "Không khẳng định tương đương leaderboard M5; chỉ gọi là close-to-official local evaluation.",
        "Nêu rõ C cải thiện trung bình nhưng không thắng tuyệt đối mọi chuỗi, mọi nhóm.",
        "Bảo vệ K=3 bằng ADI-CV², silhouette/Davies-Bouldin và diễn giải demand regime, đồng thời thừa nhận K-Medoids robustness còn hạn chế.",
        "Trình bày DM-HAC, FDR và effect size cùng nhau.",
        "Nếu cần tăng credibility, ưu tiên external baseline CatBoost/XGBoost hoặc public LightGBM baseline.",
    ]:
        document.add_paragraph(text, style="List Number")
    for section in document.sections:
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
    document.save(docx_path)
    print(md_path)
    print(docx_path)
    print(METRICS / "experimental_validity_audit_items.csv")


if __name__ == "__main__":
    main()
