from __future__ import annotations

from pathlib import Path
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


DEFAULT_INPUT = Path("Document/Detailed_Method_Results_Assessment.md")
DEFAULT_OUTPUT = Path("Document/Detailed_Method_Results_Assessment.docx")


def split_table_line(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_separator(line: str) -> bool:
    cells = split_table_line(line)
    return bool(cells) and all(set(cell.replace(":", "").strip()) <= {"-"} for cell in cells)


def add_markdown_table(doc: Document, rows: list[str]) -> None:
    parsed = [split_table_line(row) for row in rows if row.strip()]
    parsed = [row for row in parsed if not is_separator("| " + " | ".join(row) + " |")]
    if not parsed:
        return

    n_cols = max(len(row) for row in parsed)
    table = doc.add_table(rows=len(parsed), cols=n_cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(parsed):
        for c_idx in range(n_cols):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            cell.text = text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    if r_idx == 0:
                        run.bold = True


def add_paragraph_with_inline_code(doc: Document, text: str, style: str | None = None) -> None:
    paragraph = doc.add_paragraph(style=style)
    parts = text.split("`")
    for idx, part in enumerate(parts):
        run = paragraph.add_run(part)
        if idx % 2 == 1:
            run.font.name = "Courier New"
            run.font.size = Pt(9)


def main() -> None:
    input_path = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_INPUT
    output_path = Path(sys.argv[2]) if len(sys.argv) > 2 else DEFAULT_OUTPUT
    text = input_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    title_text = "Báo cáo chi tiết phương pháp, kết quả và đánh giá nghiên cứu"
    if lines and lines[0].startswith("# "):
        title_text = lines[0][2:].strip()
        text = "\n".join(lines[1:])

    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(title_text)
    run.bold = True
    run.font.size = Pt(16)

    in_code = False
    code_lines: list[str] = []
    table_lines: list[str] = []

    def flush_table() -> None:
        nonlocal table_lines
        if table_lines:
            add_markdown_table(doc, table_lines)
            table_lines = []

    def flush_code() -> None:
        nonlocal code_lines
        if code_lines:
            paragraph = doc.add_paragraph()
            for line in code_lines:
                run = paragraph.add_run(line + "\n")
                run.font.name = "Courier New"
                run.font.size = Pt(9)
            code_lines = []

    for raw_line in text.splitlines():
        line = raw_line.rstrip()

        if line.startswith("```"):
            flush_table()
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if line.startswith("|") and line.endswith("|"):
            table_lines.append(line)
            continue

        flush_table()

        if not line.strip():
            doc.add_paragraph()
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- "):
            add_paragraph_with_inline_code(doc, line[2:].strip(), style="List Bullet")
        elif line.startswith("+"):
            add_paragraph_with_inline_code(doc, line.strip(), style="List Bullet")
        else:
            add_paragraph_with_inline_code(doc, line)

    flush_table()
    flush_code()

    for section in doc.sections:
        section.top_margin = Pt(54)
        section.bottom_margin = Pt(54)
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)

    doc.save(output_path)
    print(output_path)


if __name__ == "__main__":
    main()
