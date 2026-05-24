from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


DOCX_OUT = Path("Document/Bai_Bao_Nghien_Cuu_M5_Theo_Cau_Truc_Reference.docx")
MD_OUT = Path("Document/Bai_Bao_Nghien_Cuu_M5_Theo_Cau_Truc_Reference.md")

BASE = Path("outputs_full_k3_seed42_tweedie_a0_b1_c")
FULL_SCOPE = Path("outputs_extended_fullscale_checks/metrics")
FIGURES = Path("outputs_research_figures")
EDA_FIGURES = Path("outputs_eda_figures")
METHOD_FIGURES = Path("outputs_method_figures")


def fmt(x: float, ndigits: int = 6) -> str:
    return f"{float(x):.{ndigits}f}"


def add_para(doc: Document, text: str, style: str | None = None, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    doc.add_paragraph()


def add_figure(doc: Document, path: Path, caption: str) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(6.3))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in cap.runs:
        run.font.size = Pt(10)


def md_table(headers: list[str], rows: list[list[str]]) -> str:
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    lines.extend("| " + " | ".join(row) + " |" for row in rows)
    return "\n".join(lines)


def build_content() -> dict[str, object]:
    all_metrics = pd.read_csv(FULL_SCOPE / "extended_fullscale_all_metrics.csv")
    base = all_metrics[all_metrics["run"].eq("base_tweedie")].copy()
    base = base[base["model"].isin(["A0_global_baseline", "B1_cluster_label", "C_cluster_specific"])]
    base_order = ["A0_global_baseline", "B1_cluster_label", "C_cluster_specific"]
    base["model"] = pd.Categorical(base["model"], categories=base_order, ordered=True)
    base = base.sort_values("model")

    ablation = all_metrics[
        all_metrics["run"].isin(["no_price", "no_calendar", "no_hierarchy"])
        & all_metrics["model"].eq("C_cluster_specific")
    ].copy()
    ablation["run"] = pd.Categorical(ablation["run"], ["no_price", "no_calendar", "no_hierarchy"], ordered=True)
    ablation = ablation.sort_values("run")

    eda_overview = pd.read_csv(BASE / "eda" / "dataset_overview.csv").set_index("metric")["value"]
    demand_counts = pd.read_csv(BASE / "eda" / "demand_class_counts.csv")
    describe = pd.read_csv(BASE / "eda" / "series_summary_describe.csv", index_col=0)
    clusters = pd.read_csv(BASE / "metrics" / "cluster_profiles.csv")
    clusters_1913 = clusters[clusters["origin"].eq(1913)].sort_values("mean_sales")
    kmedoids = pd.read_csv(BASE / "metrics" / "kmedoids_robustness.csv").iloc[0]
    kmed_sens = pd.read_csv(BASE / "metrics" / "kmedoids_robustness_sample_sensitivity.csv")

    return {
        "base": base,
        "ablation": ablation,
        "eda_overview": eda_overview,
        "demand_counts": demand_counts,
        "describe": describe,
        "clusters_1913": clusters_1913,
        "kmedoids": kmedoids,
        "kmed_sens": kmed_sens,
    }


def main() -> None:
    data = build_content()
    base: pd.DataFrame = data["base"]  # type: ignore[assignment]
    ablation: pd.DataFrame = data["ablation"]  # type: ignore[assignment]
    eda = data["eda_overview"]  # type: ignore[assignment]
    demand_counts: pd.DataFrame = data["demand_counts"]  # type: ignore[assignment]
    describe: pd.DataFrame = data["describe"]  # type: ignore[assignment]
    clusters_1913: pd.DataFrame = data["clusters_1913"]  # type: ignore[assignment]
    kmedoids = data["kmedoids"]  # type: ignore[assignment]
    kmed_sens: pd.DataFrame = data["kmed_sens"]  # type: ignore[assignment]

    model_labels = {
        "A0_global_baseline": "A0 - Global LightGBM",
        "B1_cluster_label": "B1 - Global LightGBM + cluster label",
        "C_cluster_specific": "C - Cluster-specific LightGBM",
    }

    base_rows = [
        [
            model_labels[row.model],
            fmt(row.sampled_wrmsse_12level),
            fmt(row.validation_origin_wrmsse_mean),
            fmt(row.scale_aware_stability_loss),
            fmt(row.wape),
            fmt(row.bias),
            fmt(row.mean_test_train_gap),
        ]
        for row in base.itertuples()
    ]

    ablation_rows = [
        [
            {"no_price": "No price", "no_calendar": "No calendar/SNAP/event", "no_hierarchy": "No hierarchy/ID"}[row.run],
            fmt(row.sampled_wrmsse_12level),
            fmt(row.validation_origin_wrmsse_mean),
            fmt(row.scale_aware_stability_loss),
            fmt(row.sampled_wrmsse_12level - float(base[base["model"].eq("C_cluster_specific")]["sampled_wrmsse_12level"].iloc[0])),
        ]
        for row in ablation.itertuples()
    ]

    cluster_rows = [
        [
            str(int(row.cluster_label)),
            str(int(row.n_series)),
            fmt(row.mean_sales, 4),
            fmt(row.zero_sales_ratio, 4),
            fmt(row.adi, 4),
            fmt(row.cv2, 4),
        ]
        for row in clusters_1913.itertuples()
    ]

    kmed_rows = [
        [
            str(int(row.sample_n)),
            fmt(row.ari_kmeans_vs_kmedoids),
            str(int(row.kmedoids_min_cluster_size)),
            str(int(row.kmedoids_max_cluster_size)),
            str(row.note),
        ]
        for row in kmed_sens.itertuples()
    ]

    c = base[base["model"].eq("C_cluster_specific")].iloc[0]
    a0 = base[base["model"].eq("A0_global_baseline")].iloc[0]
    b1 = base[base["model"].eq("B1_cluster_label")].iloc[0]

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)

    title = doc.add_heading(
        "Cluster-aware Global Forecasting cho dự báo nhu cầu bán lẻ quy mô lớn: "
        "thực nghiệm LightGBM trên dữ liệu M5",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "Bài báo nghiên cứu thực nghiệm dựa trên Proposal, dữ liệu M5 và các kết quả full-scale đã chạy", align=WD_ALIGN_PARAGRAPH.CENTER)

    add_heading(doc, "Tóm tắt", 1)
    add_para(
        doc,
        "Nghiên cứu này đánh giá một framework cluster-aware global forecasting cho bài toán dự báo nhu cầu bán lẻ quy mô lớn trên dữ liệu M5. "
        "Khác với các nghiên cứu tối ưu leaderboard thuần túy, trọng tâm của nghiên cứu là đánh giá đồng thời độ chính xác, tính ổn định của dự báo, khả năng kiểm soát data leakage và vai trò của phân cụm nhu cầu. "
        "Ba cấu hình được so sánh gồm A0 - global LightGBM baseline, B1 - global LightGBM có bổ sung cluster label, và C - LightGBM huấn luyện riêng theo cụm. "
        f"Kết quả full-scale cho thấy C đạt rolling WRMSSE {fmt(c.sampled_wrmsse_12level)}, thấp hơn A0 ({fmt(a0.sampled_wrmsse_12level)}) và B1 ({fmt(b1.sampled_wrmsse_12level)}). "
        f"C cũng đạt scale-aware stability loss thấp nhất ({fmt(c.scale_aware_stability_loss)}), cho thấy cải thiện không chỉ nằm ở accuracy mà còn ở độ ổn định dự báo qua các rolling origins. "
        "Ablation study xác nhận calendar/SNAP/event và hierarchy/ID là hai nhóm feature quan trọng, trong khi price có đóng góp nhỏ hơn. "
        f"K-Medoids robustness full-scale trên 30,490 chuỗi cho ARI {fmt(kmedoids.ari_kmeans_vs_kmedoids)}, hàm ý rằng cấu trúc cụm có tương đồng thấp-vừa giữa hai thuật toán và cần được diễn giải thận trọng. "
        "Kết quả ủng hộ việc dùng cluster-specific learning như một cơ chế mô hình hóa demand regimes, nhưng không nên xem cluster như các phân khúc kinh doanh cố định tuyệt đối.",
    )
    add_para(doc, "Từ khóa: M5, LightGBM, global forecasting, cluster-aware forecasting, WRMSSE, forecast stability, intermittent demand.", style=None)

    add_heading(doc, "1. Giới thiệu", 1)
    add_para(
        doc,
        "Dự báo nhu cầu bán lẻ quy mô lớn là một bài toán có ý nghĩa trực tiếp đối với quản trị tồn kho, lập kế hoạch cung ứng và giảm chi phí thiếu hàng hoặc dư hàng. "
        "Dữ liệu M5 được sử dụng rộng rãi vì phản ánh bối cảnh bán lẻ thực tế với nhiều chuỗi item-store, cấu trúc phân cấp, thông tin lịch, sự kiện, SNAP và giá bán. "
        "Bài báo M5 Accuracy Competition của Makridakis, Spiliotis và Assimakopoulos cho thấy LightGBM là một trong những phương pháp chủ đạo của các đội có kết quả cao, đồng thời nhấn mạnh tính phức tạp của dữ liệu hierarchy và intermittent demand. "
        "Tuy nhiên, mục tiêu của nghiên cứu hiện tại không phải tái tạo leaderboard M5, mà là kiểm tra một câu hỏi phương pháp: liệu cluster-aware global forecasting có thể cải thiện đồng thời độ chính xác và tính ổn định của dự báo hay không.",
    )
    add_para(
        doc,
        "Khoảng trống nghiên cứu nằm ở chỗ nhiều pipeline forecasting tập trung chủ yếu vào accuracy, trong khi tính ổn định qua các lần cập nhật mô hình cũng rất quan trọng trong vận hành bán lẻ. "
        "Nếu forecast thay đổi mạnh giữa các rolling origins, kế hoạch tồn kho và phân bổ hàng hóa có thể bị dao động dù accuracy trung bình có vẻ tốt. "
        "Do đó, nghiên cứu này kết hợp WRMSSE close-to-official với scale-aware stability loss, đồng thời kiểm tra overfitting, ablation, K-Medoids robustness và data leakage.",
    )

    add_heading(doc, "2. Triển khai và thiết kế thực nghiệm", 1)
    add_heading(doc, "2.1. Dữ liệu và phân tích khám phá", 2)
    add_para(
        doc,
        f"Nghiên cứu sử dụng dữ liệu M5 gồm {int(eda['n_series']):,} chuỗi item-store, {int(eda['n_items']):,} sản phẩm, {int(eda['n_stores'])} cửa hàng, "
        f"{int(eda['n_states'])} bang, {int(eda['n_categories'])} ngành hàng và {int(eda['n_calendar_days']):,} ngày lịch. "
        f"Tổng số dòng giá bán là {int(eda['n_price_rows']):,}. Median zero-sales ratio bằng {fmt(eda['median_zero_sales_ratio'], 4)}, cho thấy chuỗi điển hình có hơn 63% số ngày không phát sinh doanh số.",
    )
    demand_rows = [[str(r.demand_class), str(int(r.n_series))] for r in demand_counts.itertuples()]
    add_table(doc, ["Nhóm nhu cầu", "Số chuỗi"], demand_rows)
    add_para(
        doc,
        f"Phân loại demand regime cho thấy {int(demand_counts[demand_counts['demand_class'].eq('intermittent')]['n_series'].iloc[0]):,} chuỗi intermittent và "
        f"{int(demand_counts[demand_counts['demand_class'].eq('lumpy')]['n_series'].iloc[0]):,} chuỗi lumpy. "
        f"Median ADI là {fmt(describe.loc['50%', 'adi'], 4)} và median CV2 là {fmt(describe.loc['50%', 'cv2'], 4)}. "
        "Đặc điểm này giải thích vì sao nghiên cứu cần objective phù hợp với dữ liệu không âm nhiều giá trị 0 và metric stability không bị phóng đại bởi mẫu số gần 0.",
    )
    add_figure(doc, EDA_FIGURES / "eda_fig01_demand_class_counts.png", "Hình 1. Phân bố chuỗi theo nhóm nhu cầu trong dữ liệu M5.")
    add_figure(doc, EDA_FIGURES / "eda_fig03_adi_cv2_by_demand_class.png", "Hình 2. Bản đồ ADI-CV2 cho thấy sự tồn tại của nhiều demand regimes.")

    add_heading(doc, "2.2. Thiết kế rolling-origin và kiểm soát leakage", 2)
    add_para(
        doc,
        "Thực nghiệm sử dụng rolling-origin evaluation với các origin 1885, 1892, 1899, 1906 và 1913. "
        "Tại mỗi origin, pipeline chỉ sử dụng dữ liệu lịch sử đến thời điểm origin để tạo feature, phân cụm và huấn luyện mô hình; sau đó dự báo 28 ngày tiếp theo. "
        "Thiết kế này mô phỏng tình huống vận hành khi mô hình được cập nhật định kỳ.",
    )
    add_para(
        doc,
        "Nguyên tắc anti-leakage được áp dụng xuyên suốt: lag và rolling features được shift theo thời gian; clustering features được tính riêng tại từng origin; forecast horizon dùng recursive prediction để cập nhật lag bằng dự báo đã sinh ra, không sử dụng actual demand tương lai. "
        "Release day và availability cũng được xử lý để tránh xem ngày trước khi sản phẩm được bán là zero demand thật.",
    )
    add_figure(doc, METHOD_FIGURES / "method_workflow.png", "Hình 3. Quy trình phương pháp thực hiện nghiên cứu.")

    add_heading(doc, "2.3. Thước đo đánh giá", 2)
    add_para(
        doc,
        "Độ chính xác được đánh giá bằng WRMSSE close-to-official, phù hợp với cấu trúc phân cấp của M5 và nguyên tắc lower is better. "
        "Bên cạnh accuracy, nghiên cứu dùng scale-aware stability loss để đo mức thay đổi dự báo cho cùng target date khi rolling origin thay đổi. "
        "Metric stability này sử dụng scale floor để giảm sai lệch trên chuỗi gần 0, một vấn đề quan trọng trong intermittent demand.",
    )

    add_heading(doc, "3. Các mô hình thực nghiệm", 1)
    add_para(
        doc,
        "Ba cấu hình mô hình được giữ làm trọng tâm nghiên cứu. A0 là global LightGBM baseline, huấn luyện chung trên toàn bộ panel và không dùng thông tin cụm. "
        "B1 là global LightGBM có bổ sung cluster label như một feature, nhằm kiểm tra liệu nhãn cụm dạng biến giải thích có đủ cải thiện mô hình global hay không. "
        "C là cluster-specific LightGBM, trong đó mỗi cụm nhu cầu có một mô hình riêng. Cấu hình C thể hiện giả thuyết chính của nghiên cứu: các demand regimes khác nhau có thể cần cơ chế học khác nhau.",
    )
    add_para(
        doc,
        "Phân cụm chính sử dụng Mini-batch K-Means với K=3 vì thuật toán này phù hợp quy mô 30,490 chuỗi và có khả năng chạy qua nhiều rolling origins. "
        "K-Medoids không được dùng làm clustering chính, mà được dùng như robustness check hậu nghiệm để kiểm tra độ nhạy của cấu trúc cụm khi đổi thuật toán.",
    )

    add_heading(doc, "4. Kết quả và phát hiện chính", 1)
    add_heading(doc, "4.1. Kết quả chính của A0, B1 và C", 2)
    add_table(
        doc,
        ["Mô hình", "Rolling WRMSSE", "Validation WRMSSE", "Stability loss", "WAPE", "Bias", "Train-test gap"],
        base_rows,
    )
    add_para(
        doc,
        f"C đạt rolling WRMSSE {fmt(c.sampled_wrmsse_12level)}, thấp nhất trong ba mô hình. "
        f"So với A0, C giảm WRMSSE khoảng {fmt(a0.sampled_wrmsse_12level - c.sampled_wrmsse_12level)}; so với B1, C giảm khoảng {fmt(b1.sampled_wrmsse_12level - c.sampled_wrmsse_12level)}. "
        f"C cũng có stability loss thấp nhất ({fmt(c.scale_aware_stability_loss)}), trong khi A0 và B1 lần lượt là {fmt(a0.scale_aware_stability_loss)} và {fmt(b1.scale_aware_stability_loss)}. "
        "Kết quả này cho thấy lợi ích của cluster-aware forecasting không xuất hiện khi chỉ thêm cluster label vào global model, mà xuất hiện rõ hơn khi huấn luyện mô hình riêng theo cụm.",
    )
    add_figure(doc, FIGURES / "fig01_base_wrmsse_a0_b1_c.png", "Hình 4. So sánh rolling WRMSSE của A0, B1 và C.")
    add_figure(doc, FIGURES / "fig02_base_stability_a0_b1_c.png", "Hình 5. So sánh scale-aware stability loss của A0, B1 và C.")
    add_figure(doc, FIGURES / "fig03_accuracy_stability_tradeoff.png", "Hình 6. Trade-off giữa accuracy và stability.")

    add_heading(doc, "4.2. Diễn giải cụm nhu cầu", 2)
    add_table(doc, ["Cluster", "Số chuỗi", "Mean sales", "Zero-sales ratio", "ADI", "CV2"], cluster_rows)
    add_para(
        doc,
        "Tại origin 1913, các cụm thể hiện sự khác biệt rõ về mức bán, độ thưa và tần suất phát sinh nhu cầu. "
        "Một cụm có mean sales thấp, zero-sales ratio cao và ADI cao, đại diện cho nhóm long-tail/intermittent. "
        "Một cụm có mean sales cao hơn và zero-sales ratio thấp hơn, đại diện cho nhóm core hoặc high-demand. "
        "Sự phân hóa này giải thích vì sao mô hình C có thể cải thiện kết quả: mỗi cụm được học bằng một mô hình phù hợp hơn với demand regime của nó.",
    )
    add_figure(doc, FIGURES / "fig06_cluster_profile_k3.png", "Hình 7. Đặc trưng ba cụm nhu cầu K=3.")

    add_heading(doc, "4.3. Ablation study", 2)
    add_table(doc, ["Biến thể C", "Rolling WRMSSE", "Validation WRMSSE", "Stability loss", "Chênh lệch WRMSSE so với C base"], ablation_rows)
    add_para(
        doc,
        "Ablation cho thấy calendar/SNAP/event là nhóm feature quan trọng nhất: khi loại bỏ nhóm này, WRMSSE của C tăng mạnh lên 0.972886. "
        "Khi loại bỏ hierarchy/ID, WRMSSE tăng lên 0.633918, cho thấy thông tin item, category, department, store và state là nền tảng quan trọng của mô hình. "
        "Khi loại bỏ price, WRMSSE tăng nhẹ lên 0.583373 và stability loss xấu hơn, hàm ý price có đóng góp nhưng nhỏ hơn calendar và hierarchy.",
    )
    add_figure(doc, FIGURES / "fig04_ablation_wrmsse_c.png", "Hình 8. Tác động của ablation lên WRMSSE của model C.")
    add_figure(doc, FIGURES / "fig05_ablation_delta_wrmsse_c.png", "Hình 9. Mức tăng WRMSSE khi loại từng nhóm feature.")

    add_heading(doc, "4.4. Overfitting, robustness và kiểm định bổ trợ", 2)
    add_para(
        doc,
        f"Mean train-test gap của C là {fmt(c.mean_test_train_gap)}, cao hơn A0 ({fmt(a0.mean_test_train_gap)}) và B1 ({fmt(b1.mean_test_train_gap)}) nhưng không bất thường khi đặt cạnh WRMSSE và stability tốt hơn. "
        "Do C huấn luyện nhiều mô hình riêng theo cụm, rủi ro overfitting cần được theo dõi, nhưng kết quả hiện tại không cho thấy overfitting nghiêm trọng.",
    )
    add_figure(doc, FIGURES / "fig07_overfitting_gap_a0_b1_c.png", "Hình 10. Mean test-train gap của A0, B1 và C.")
    add_table(doc, ["Sample", "ARI", "Min cluster size", "Max cluster size", "Ghi chú"], kmed_rows)
    add_para(
        doc,
        f"K-Medoids robustness full-scale trên {int(kmedoids.sample_n):,} chuỗi cho ARI {fmt(kmedoids.ari_kmeans_vs_kmedoids)}. "
        "Mức này được diễn giải là tương đồng thấp-vừa giữa Mini-batch K-Means và K-Medoids. "
        "Do đó, cluster có giá trị như một cơ chế kỹ thuật cho cluster-specific learning, nhưng không nên diễn giải như các phân khúc kinh doanh cố định tuyệt đối.",
    )

    add_heading(doc, "4.5. So sánh với bài M5 Accuracy Competition", 2)
    add_para(
        doc,
        "Bài M5 Accuracy Competition báo cáo các đội đứng đầu đạt WRMSSE khoảng 0.520 và cho thấy LightGBM được sử dụng rộng rãi trong nhóm phương pháp hiệu quả nhất. "
        "Kết quả tốt nhất của nghiên cứu hiện tại không vượt winner M5 vì mục tiêu và thiết kế khác nhau: nghiên cứu này không tối ưu leaderboard, không dùng ensemble phức tạp và không submit official leaderboard. "
        "Đóng góp chính nằm ở việc đánh giá framework cluster-aware theo rolling-origin, bổ sung stability metric, kiểm soát leakage và phân tích ablation.",
    )

    add_heading(doc, "5. Thảo luận, hạn chế và hướng nghiên cứu", 1)
    add_heading(doc, "5.1. Thảo luận", 2)
    add_para(
        doc,
        "Kết quả cho thấy khác biệt giữa các demand regimes có ý nghĩa đối với forecasting. "
        "Tuy nhiên, cách tích hợp cụm quyết định hiệu quả: B1 không cải thiện so với A0, trong khi C tốt nhất trên cả accuracy và stability. "
        "Điều này hàm ý cluster label như một feature đơn lẻ chưa đủ để thay đổi cấu trúc học của global model; lợi ích rõ hơn khi mỗi cụm được mô hình hóa bằng một learner riêng.",
    )
    add_para(
        doc,
        "Ablation củng cố vai trò của calendar/SNAP/event và hierarchy features. "
        "Đây là kết quả phù hợp với bán lẻ thực tế vì nhu cầu chịu ảnh hưởng bởi chu kỳ thời gian, sự kiện, chương trình SNAP, đặc điểm sản phẩm và khác biệt địa lý. "
        "Price features có đóng góp nhỏ hơn nhưng vẫn nên giữ vì hỗ trợ stability.",
    )

    add_heading(doc, "5.2. Hạn chế", 2)
    for text in [
        "WRMSSE trong nghiên cứu là close-to-official local evaluation, không phải official Kaggle leaderboard score.",
        "Rolling-origin chỉ gồm 5 origins, do đó kiểm định thống kê cần diễn giải thận trọng.",
        "K-Medoids robustness full-scale cho ARI thấp-vừa, nên không nên diễn giải cluster như phân khúc kinh doanh cố định.",
        "Model C có chi phí vận hành cao hơn A0/B1 vì phải huấn luyện nhiều mô hình theo cụm.",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    add_heading(doc, "5.3. Hướng nghiên cứu tiếp theo", 2)
    for text in [
        "Mở rộng số rolling origins hoặc dùng nhiều seed hơn để đánh giá độ ổn định thống kê.",
        "Kiểm tra thêm các phương pháp clustering có khả năng mở rộng, ví dụ CLARA hoặc approximate medoids.",
        "Kết hợp cluster-specific learning với ensemble hoặc hierarchical reconciliation nếu mục tiêu là tiến gần hơn leaderboard M5.",
        "Phân tích sâu hơn bias và stability theo category, store, state và nhóm intermittent demand.",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    add_heading(doc, "6. Kết luận", 1)
    add_para(
        doc,
        "Nghiên cứu đã thực hiện một đánh giá full-scale framework cluster-aware global forecasting trên dữ liệu M5 với ba cấu hình A0, B1 và C. "
        "Kết quả cho thấy C - cluster-specific LightGBM - là phương án tốt nhất trong phạm vi nghiên cứu, với WRMSSE và stability loss thấp nhất. "
        "B1 không vượt A0, cho thấy việc thêm cluster label vào global model chưa đủ để khai thác đầy đủ thông tin cụm. "
        "Các kiểm tra ablation xác nhận calendar/SNAP/event và hierarchy/ID là các nhóm feature thiết yếu. "
        "K-Medoids full-scale cho thấy cấu trúc cụm cần được diễn giải thận trọng, nhưng không làm thay đổi kết luận chính về hiệu quả của C. "
        "Tổng thể, kết quả đáp ứng mục tiêu của proposal: đánh giá một framework cluster-aware nhằm cải thiện đồng thời độ chính xác và tính ổn định trong dự báo nhu cầu bán lẻ quy mô lớn.",
    )

    add_heading(doc, "Tài liệu tham khảo", 1)
    add_para(
        doc,
        "Makridakis, S., Spiliotis, E., & Assimakopoulos, V. (2022). M5 accuracy competition: Results, findings, and conclusions. International Journal of Forecasting, 38(4), 1346-1364.",
    )

    doc.save(DOCX_OUT)

    md_sections: list[str] = [
        "# Cluster-aware Global Forecasting cho dự báo nhu cầu bán lẻ quy mô lớn",
        "",
        "## Tóm tắt",
        "Nghiên cứu đánh giá A0, B1 và C trên dữ liệu M5, tập trung vào WRMSSE, stability, leakage control, ablation và robustness.",
        "",
        "## Bảng kết quả chính",
        md_table(["Mô hình", "Rolling WRMSSE", "Validation WRMSSE", "Stability loss", "WAPE", "Bias", "Train-test gap"], base_rows),
        "",
        "## Ablation của model C",
        md_table(["Biến thể C", "Rolling WRMSSE", "Validation WRMSSE", "Stability loss", "Chênh lệch WRMSSE"], ablation_rows),
        "",
        "## K-Medoids robustness",
        md_table(["Sample", "ARI", "Min cluster size", "Max cluster size", "Ghi chú"], kmed_rows),
        "",
        "File Word chứa bản trình bày đầy đủ theo cấu trúc bài báo tham chiếu M5.",
    ]
    MD_OUT.write_text("\n".join(md_sections), encoding="utf-8")

    print(DOCX_OUT)
    print(MD_OUT)


if __name__ == "__main__":
    main()
