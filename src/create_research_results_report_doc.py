from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "Document" / "Ket_Qua_Nghien_Cuu_M5.docx"
OUT_MD = ROOT / "Document" / "Ket_Qua_Nghien_Cuu_M5.md"
FIG_DIR = ROOT / "Document" / "figures_results_report"

BASE = ROOT / "outputs_full_k3_seed42_tweedie_a0_b1_c"
MULTI = ROOT / "outputs_full_multiseed_summary"
EXT = ROOT / "outputs_extended_fullscale_checks"

MODEL_LABELS = {
    "A0_global_baseline": "A0: Global LightGBM",
    "B1_cluster_label": "B1: Global LightGBM + cluster label",
    "C_cluster_specific": "C: Cluster-specific LightGBM",
}


def read_csv(path: Path) -> pd.DataFrame:
    if not path.exists():
        raise FileNotFoundError(path)
    return pd.read_csv(path)


def fmt(value, digits: int = 6) -> str:
    if pd.isna(value):
        return ""
    if isinstance(value, (int,)) or (isinstance(value, float) and value.is_integer() and abs(value) > 10):
        return f"{int(value):,}"
    if isinstance(value, float):
        return f"{value:.{digits}f}"
    return str(value)


def add_table(doc: Document, df: pd.DataFrame, title: str, digits: int = 6) -> None:
    doc.add_paragraph(title, style="Caption")
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr[i].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            cells[i].text = fmt(row[col], digits)


def add_figure(doc: Document, path: Path, caption: str, width: float = 6.3) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption, style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def save_multiseed_plots() -> list[tuple[Path, str]]:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    agg = read_csv(MULTI / "metrics" / "full_multiseed_aggregate_metrics.csv").copy()
    agg["label"] = agg["model"].map(MODEL_LABELS)

    plots = []
    metric_specs = [
        ("sampled_wrmsse_12level", "WRMSSE close-to-official", "multiseed_wrmsse.png"),
        ("scale_aware_stability_loss", "Scale-aware stability loss", "multiseed_stability.png"),
        ("wape", "WAPE", "multiseed_wape.png"),
        ("train_test_gap", "Train-test gap", "multiseed_train_test_gap.png"),
    ]
    for metric, title, filename in metric_specs:
        mean_col = f"{metric}_mean"
        low_col = f"{metric}_ci95_low"
        high_col = f"{metric}_ci95_high"
        if mean_col not in agg:
            continue
        out = FIG_DIR / filename
        draw_bar_chart(
            out,
            title,
            agg["label"].tolist(),
            agg[mean_col].tolist(),
            agg[low_col].tolist(),
            agg[high_col].tolist(),
        )
        plots.append((out, f"Hình. Multi-seed {title} của A0, B1 và C."))

    per_origin = read_csv(MULTI / "metrics" / "full_multiseed_per_origin_metrics.csv").copy()
    if "sampled_wrmsse_12level" in per_origin:
        per_origin["label"] = per_origin["model"].map(MODEL_LABELS)
        out = FIG_DIR / "rolling_origin_wrmsse.png"
        series = {}
        for model, group in per_origin.groupby("label"):
            g = group.groupby("origin")["sampled_wrmsse_12level"].mean().reset_index()
            series[model] = list(zip(g["origin"].tolist(), g["sampled_wrmsse_12level"].tolist()))
        draw_line_chart(out, "WRMSSE theo rolling origin", series)
        plots.append((out, "Hình. WRMSSE trung bình theo từng rolling origin trong thiết kế multi-seed."))

    return plots


def load_font(size: int, bold: bool = False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Times New Roman Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Times New Roman.ttf",
    ]
    for path in candidates:
        try:
            return ImageFont.truetype(path, size)
        except OSError:
            continue
    return ImageFont.load_default()


def draw_bar_chart(path: Path, title: str, labels: list[str], means: list[float], lows: list[float], highs: list[float]) -> None:
    width, height = 1400, 850
    margin_l, margin_r, margin_t, margin_b = 120, 70, 105, 170
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    f_title = load_font(34, True)
    f_label = load_font(22)
    f_small = load_font(19)
    colors = ["#7FB3D5", "#82E0AA", "#F8C471"]

    title_box = draw.textbbox((0, 0), title, font=f_title)
    draw.text(((width - (title_box[2] - title_box[0])) / 2, 32), title, fill="#17202A", font=f_title)

    y_min = min(lows)
    y_max = max(highs)
    pad = (y_max - y_min) * 0.25 if y_max > y_min else 0.01
    y_min = max(0, y_min - pad)
    y_max = y_max + pad
    plot_x1, plot_y1 = margin_l, margin_t
    plot_x2, plot_y2 = width - margin_r, height - margin_b

    draw.line((plot_x1, plot_y2, plot_x2, plot_y2), fill="#2C3E50", width=3)
    draw.line((plot_x1, plot_y1, plot_x1, plot_y2), fill="#2C3E50", width=3)

    def y_pos(v: float) -> float:
        return plot_y2 - (v - y_min) / (y_max - y_min) * (plot_y2 - plot_y1)

    for i in range(5):
        val = y_min + i * (y_max - y_min) / 4
        y = y_pos(val)
        draw.line((plot_x1, y, plot_x2, y), fill="#E5E7E9", width=1)
        draw.text((25, y - 12), f"{val:.4f}", fill="#566573", font=f_small)

    n = len(labels)
    slot = (plot_x2 - plot_x1) / n
    bar_w = slot * 0.46
    for i, (label, mean, low, high) in enumerate(zip(labels, means, lows, highs)):
        cx = plot_x1 + slot * (i + 0.5)
        y_mean = y_pos(mean)
        y0 = y_pos(0)
        draw.rectangle((cx - bar_w / 2, y_mean, cx + bar_w / 2, y0), fill=colors[i % len(colors)], outline="#2C3E50", width=2)
        draw.line((cx, y_pos(low), cx, y_pos(high)), fill="#1B2631", width=3)
        draw.line((cx - 13, y_pos(low), cx + 13, y_pos(low)), fill="#1B2631", width=3)
        draw.line((cx - 13, y_pos(high), cx + 13, y_pos(high)), fill="#1B2631", width=3)
        draw.text((cx - 42, y_mean - 30), f"{mean:.4f}", fill="#17202A", font=f_small)
        # Split long labels over two lines.
        parts = label.replace(" + ", "\n+ ").replace(": ", ":\n").split("\n")
        y_text = plot_y2 + 18
        for part in parts:
            tb = draw.textbbox((0, 0), part, font=f_label)
            draw.text((cx - (tb[2] - tb[0]) / 2, y_text), part, fill="#17202A", font=f_label)
            y_text += 25

    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path)


def draw_line_chart(path: Path, title: str, series: dict[str, list[tuple[int, float]]]) -> None:
    width, height = 1400, 850
    margin_l, margin_r, margin_t, margin_b = 120, 270, 105, 100
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    f_title = load_font(34, True)
    f_label = load_font(21)
    f_small = load_font(18)
    colors = ["#2874A6", "#229954", "#D68910"]

    title_box = draw.textbbox((0, 0), title, font=f_title)
    draw.text(((width - (title_box[2] - title_box[0])) / 2, 32), title, fill="#17202A", font=f_title)
    xs = sorted({x for points in series.values() for x, _ in points})
    vals = [v for points in series.values() for _, v in points]
    y_min, y_max = min(vals), max(vals)
    pad = (y_max - y_min) * 0.2 if y_max > y_min else 0.01
    y_min -= pad
    y_max += pad
    plot_x1, plot_y1 = margin_l, margin_t
    plot_x2, plot_y2 = width - margin_r, height - margin_b
    draw.line((plot_x1, plot_y2, plot_x2, plot_y2), fill="#2C3E50", width=3)
    draw.line((plot_x1, plot_y1, plot_x1, plot_y2), fill="#2C3E50", width=3)

    def x_pos(x):
        return plot_x1 + (x - min(xs)) / (max(xs) - min(xs)) * (plot_x2 - plot_x1)

    def y_pos(v):
        return plot_y2 - (v - y_min) / (y_max - y_min) * (plot_y2 - plot_y1)

    for i in range(5):
        val = y_min + i * (y_max - y_min) / 4
        y = y_pos(val)
        draw.line((plot_x1, y, plot_x2, y), fill="#E5E7E9", width=1)
        draw.text((25, y - 10), f"{val:.4f}", fill="#566573", font=f_small)

    for x in xs:
        xp = x_pos(x)
        draw.line((xp, plot_y2, xp, plot_y2 + 8), fill="#2C3E50", width=2)
        draw.text((xp - 22, plot_y2 + 18), str(x), fill="#17202A", font=f_small)

    for idx, (name, points) in enumerate(series.items()):
        pts = [(x_pos(x), y_pos(v)) for x, v in points]
        color = colors[idx % len(colors)]
        if len(pts) >= 2:
            draw.line(pts, fill=color, width=4)
        for x, y in pts:
            draw.ellipse((x - 6, y - 6, x + 6, y + 6), fill=color, outline="#1B2631")
        ly = plot_y1 + idx * 34
        draw.rectangle((plot_x2 + 35, ly, plot_x2 + 60, ly + 18), fill=color)
        draw.text((plot_x2 + 70, ly - 4), name, fill="#17202A", font=f_label)

    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path)


def build_seed42_summary() -> pd.DataFrame:
    train_val_test = read_csv(BASE / "metrics" / "model_train_val_test_metrics.csv")
    sampled = read_csv(BASE / "metrics" / "sampled_wrmsse_overall.csv")
    stability = read_csv(BASE / "metrics" / "forecast_stability_metrics.csv")

    test = train_val_test[train_val_test["split"] == "test"].groupby("model").agg(
        RMSSE=("rmsse_item_store", "mean"),
        WAPE=("wape", "mean"),
        Bias=("bias", "mean"),
    )
    wrmsse = sampled.groupby("model")["sampled_wrmsse_12level"].mean().rename("Rolling WRMSSE")
    stab = stability.groupby("model").agg(
        **{
            "Stability loss": ("scale_aware_stability_loss", "mean"),
            "JumpRate@0.3": ("jump_rate_tau_0.3", "mean"),
            "JumpRate@0.5": ("jump_rate_tau_0.5", "mean"),
        }
    )
    train = train_val_test[train_val_test["split"] == "train"].groupby("model")["rmsse_item_store"].mean()
    test_r = train_val_test[train_val_test["split"] == "test"].groupby("model")["rmsse_item_store"].mean()
    gap = (test_r - train).rename("Train-test gap")

    out = pd.concat([wrmsse, test, stab, gap], axis=1).reset_index()
    out["Mô hình"] = out["model"].map(MODEL_LABELS)
    cols = ["Mô hình", "Rolling WRMSSE", "RMSSE", "WAPE", "Bias", "Stability loss", "JumpRate@0.3", "JumpRate@0.5", "Train-test gap"]
    return out[cols].sort_values("Rolling WRMSSE")


def build_multiseed_summary() -> pd.DataFrame:
    agg = read_csv(MULTI / "metrics" / "full_multiseed_aggregate_metrics.csv")
    rows = []
    for _, r in agg.iterrows():
        rows.append(
            {
                "Mô hình": MODEL_LABELS.get(r["model"], r["model"]),
                "WRMSSE mean": r["sampled_wrmsse_12level_mean"],
                "WRMSSE std": r["sampled_wrmsse_12level_std"],
                "WRMSSE CI95": f"[{r['sampled_wrmsse_12level_ci95_low']:.6f}; {r['sampled_wrmsse_12level_ci95_high']:.6f}]",
                "Stability mean": r["scale_aware_stability_loss_mean"],
                "WAPE mean": r["wape_mean"],
                "Train-test gap": r["train_test_gap_mean"],
            }
        )
    return pd.DataFrame(rows).sort_values("WRMSSE mean")


def build_rolling_table() -> pd.DataFrame:
    sampled = read_csv(BASE / "metrics" / "sampled_wrmsse_overall.csv")
    sampled["Mô hình"] = sampled["model"].map(MODEL_LABELS)
    pivot = sampled.pivot(index="origin", columns="Mô hình", values="sampled_wrmsse_12level").reset_index()
    pivot = pivot.rename(columns={"origin": "Origin"})
    return pivot


def build_ablation_table() -> pd.DataFrame:
    df = read_csv(EXT / "metrics" / "extended_fullscale_ablation_diff_vs_base.csv")
    c = df[df["model"] == "C_cluster_specific"].copy()
    c["Biến thể"] = c["run"].replace(
        {
            "no_price": "Bỏ price",
            "no_calendar": "Bỏ calendar/SNAP/event",
            "no_hierarchy": "Bỏ hierarchy/ID",
        }
    )
    return c[
        [
            "Biến thể",
            "rolling_wrmsse_diff_vs_base",
            "validation_wrmsse_diff_vs_base",
            "stability_diff_vs_base",
            "gap_diff_vs_base",
        ]
    ].rename(
        columns={
            "rolling_wrmsse_diff_vs_base": "Delta Rolling WRMSSE",
            "validation_wrmsse_diff_vs_base": "Delta Validation WRMSSE",
            "stability_diff_vs_base": "Delta Stability",
            "gap_diff_vs_base": "Delta Gap",
        }
    )


def build_dm_table() -> pd.DataFrame:
    df = read_csv(BASE / "metrics" / "dm_test_a0_vs_c_summary.csv")
    return df[
        [
            "loss",
            "n_series",
            "mean_loss_diff_a0_minus_c",
            "share_c_better_mean_loss",
            "share_c_better_p05",
            "share_c_better_fdr05",
        ]
    ].rename(
        columns={
            "loss": "Loss",
            "n_series": "Số chuỗi",
            "mean_loss_diff_a0_minus_c": "Mean diff A0-C",
            "share_c_better_mean_loss": "Tỷ lệ C tốt hơn",
            "share_c_better_p05": "Tỷ lệ p<0.05",
            "share_c_better_fdr05": "Tỷ lệ FDR<0.05",
        }
    )


def build_cluster_profile_table() -> pd.DataFrame:
    df = read_csv(BASE / "metrics" / "cluster_profiles.csv")
    keep = [c for c in ["origin", "cluster", "n_series", "mean_sales_mean", "zero_sales_ratio_mean", "adi_mean", "cv2_mean"] if c in df.columns]
    out = df[keep].copy()
    if "origin" in out:
        out = out[out["origin"] == out["origin"].max()]
    return out.rename(
        columns={
            "origin": "Origin",
            "cluster": "Cluster",
            "n_series": "Số chuỗi",
            "mean_sales_mean": "Mean sales",
            "zero_sales_ratio_mean": "Zero-sales ratio",
            "adi_mean": "ADI",
            "cv2_mean": "CV²",
        }
    )


def build_kmedoids_table() -> pd.DataFrame:
    df = read_csv(BASE / "metrics" / "kmedoids_robustness.csv")
    return df.rename(
        columns={
            "origin": "Origin",
            "k": "K",
            "sample_n": "Số chuỗi",
            "ari_kmeans_vs_kmedoids": "ARI K-Means vs K-Medoids",
            "kmedoids_min_cluster_size": "Min cluster size",
            "kmedoids_max_cluster_size": "Max cluster size",
        }
    )


def build_paired_tests_table() -> pd.DataFrame:
    df = read_csv(MULTI / "metrics" / "full_multiseed_paired_tests_a0_vs_c.csv")
    return df[
        [
            "metric",
            "n_seeds",
            "mean_diff_A0_minus_C",
            "paired_t_p_value",
            "wilcoxon_p_value",
        ]
    ].rename(
        columns={
            "metric": "Metric",
            "n_seeds": "Số seed",
            "mean_diff_A0_minus_C": "Mean diff A0-C",
            "paired_t_p_value": "Paired t p-value",
            "wilcoxon_p_value": "Wilcoxon p-value",
        }
    )


def add_metric_explanations(doc: Document) -> None:
    data = [
        ("WRMSSE", "Sai số chính theo cấu trúc phân cấp M5; càng thấp càng tốt."),
        ("RMSSE", "Sai số đã chuẩn hóa theo scale lịch sử ở cấp item-store; càng thấp càng tốt."),
        ("WAPE", "Tổng sai số tuyệt đối chia cho tổng actual; dễ diễn giải theo tỷ lệ; càng thấp càng tốt."),
        ("Bias", "Độ lệch tổng thể của dự báo; dương là dự báo cao hơn actual, âm là thấp hơn actual."),
        ("Scale-aware stability loss", "Mức dao động forecast giữa hai rolling origins, có chặn mẫu số để tránh phóng đại near-zero demand; càng thấp càng ổn định."),
        ("JumpRate@0.3/@0.5", "Tỷ lệ forecast có mức thay đổi chuẩn hóa vượt ngưỡng 0.3 hoặc 0.5; càng thấp càng ổn định."),
        ("Train-test gap", "Chênh lệch sai số test và train; dùng để phát hiện overfitting."),
        ("DM test", "Kiểm định Diebold-Mariano cho cặp A0-C, dùng để đánh giá khác biệt loss có ý nghĩa thống kê hay không."),
        ("ARI", "Adjusted Rand Index giữa K-Means và K-Medoids; đo độ tương đồng phân cụm."),
    ]
    add_table(doc, pd.DataFrame(data, columns=["Chỉ số", "Ý nghĩa"]), "Bảng. Ý nghĩa các chỉ số đánh giá.", digits=4)


def create_report() -> None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    multiseed_plots = save_multiseed_plots()

    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(13)
    styles["Heading 1"].font.name = "Times New Roman"
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 2"].font.name = "Times New Roman"
    styles["Heading 2"].font.size = Pt(14)
    styles["Caption"].font.name = "Times New Roman"
    styles["Caption"].font.size = Pt(10)

    title = doc.add_heading("KẾT QUẢ NGHIÊN CỨU M5 CLUSTER-AWARE FORECASTING", 1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Phần này tổng hợp các kết quả thực nghiệm đã chạy cho ba mô hình trọng tâm: "
        "A0 Global LightGBM, B1 Global LightGBM có cluster label và C Cluster-specific "
        "LightGBM. Các kết quả được trình bày theo cùng thiết kế leakage-aware rolling-origin, "
        "không sử dụng dữ liệu tương lai trong feature engineering, clustering, scaler hoặc "
        "recursive forecasting."
    )

    doc.add_heading("1. Phân tích khám phá dữ liệu", 2)
    doc.add_paragraph(
        "EDA cho thấy dữ liệu M5 có độ thưa cao và hành vi nhu cầu rất dị thể. Median "
        "zero-sales ratio đạt 0.6337, median ADI là 2.7300 và median CV² là 0.3485. "
        "Intermittent và lumpy demand chiếm tỷ trọng lớn, tạo cơ sở cho việc dùng "
        "scale-aware stability và cluster-aware learning."
    )
    add_figure(doc, ROOT / "outputs_eda_figures" / "eda_fig01_demand_class_counts.png", "Hình. Phân bố demand regime theo ADI-CV².")
    add_figure(doc, ROOT / "outputs_eda_figures" / "eda_fig02_zero_sales_ratio_distribution.png", "Hình. Phân phối zero-sales ratio của các chuỗi item-store.")
    add_figure(doc, ROOT / "outputs_eda_figures" / "eda_fig03_adi_cv2_by_demand_class.png", "Hình. Bản đồ ADI-CV² theo demand regime.")

    doc.add_heading("2. Kết quả tổng hợp A0, B1 và C", 2)
    summary = build_seed42_summary()
    add_table(doc, summary, "Bảng. Kết quả full-scale seed 42 của A0, B1 và C.")
    doc.add_paragraph(
        "Ở cấu hình full-scale seed 42, C đạt Rolling WRMSSE thấp nhất trong ba mô hình. "
        "A0 đứng sau C, trong khi B1 không cải thiện ổn định so với A0. Kết quả này cho "
        "thấy việc chỉ thêm cluster label vào global model chưa đủ; lợi ích rõ hơn khi "
        "cluster được dùng để tách cơ chế học như trong C."
    )
    add_figure(doc, ROOT / "outputs_research_figures" / "fig01_base_wrmsse_a0_b1_c.png", "Hình. So sánh Rolling WRMSSE của A0, B1 và C.")
    add_figure(doc, ROOT / "outputs_research_figures" / "fig02_base_stability_a0_b1_c.png", "Hình. So sánh scale-aware stability loss của A0, B1 và C.")
    add_figure(doc, ROOT / "outputs_research_figures" / "fig03_accuracy_stability_tradeoff.png", "Hình. Trade-off giữa accuracy và stability.")

    doc.add_heading("3. Kết quả theo rolling origin", 2)
    add_table(doc, build_rolling_table(), "Bảng. WRMSSE theo từng rolling origin.")
    doc.add_paragraph(
        "Kết quả theo rolling origin cho thấy C không chỉ tốt ở một split đơn lẻ. Tuy nhiên, "
        "mức cải thiện thay đổi theo origin, phản ánh tính mùa vụ, sự kiện và biến động nhu "
        "cầu ở cuối chuỗi M5."
    )

    doc.add_heading("4. Kết quả phân cụm và demand regime", 2)
    add_table(doc, build_cluster_profile_table(), "Bảng. Profile cụm K=3 tại rolling origin cuối.")
    add_figure(doc, ROOT / "outputs_research_figures" / "fig06_cluster_profile_k3.png", "Hình. Profile các cụm K=3 theo mean sales, zero-sales ratio, ADI và CV².")
    doc.add_paragraph(
        "Các cụm K=3 phản ánh khác biệt về scale, độ thưa và mức biến động của nhu cầu. "
        "Điều này hỗ trợ giả thuyết rằng cluster-specific learning có thể giảm dị thể so "
        "với một global learner duy nhất."
    )

    doc.add_heading("5. Overfitting và feature importance", 2)
    add_figure(doc, ROOT / "outputs_research_figures" / "fig07_overfitting_gap_a0_b1_c.png", "Hình. Train-test gap của A0, B1 và C.")
    add_figure(doc, ROOT / "outputs_research_figures" / "fig08_feature_importance_c.png", "Hình. Top feature importance của mô hình C.")
    doc.add_paragraph(
        "C có train-test gap cao hơn A0 trong một số thiết lập, cho thấy lợi ích accuracy "
        "đi kèm trade-off về độ phức tạp. Kết quả không cho thấy overfitting nghiêm trọng, "
        "nhưng cần đọc kết quả C cùng train-test gap và multi-seed robustness."
    )

    doc.add_heading("6. Ablation study", 2)
    add_table(doc, build_ablation_table(), "Bảng. Tác động của ablation đối với mô hình C.")
    add_figure(doc, ROOT / "outputs_research_figures" / "fig04_ablation_wrmsse_c.png", "Hình. WRMSSE của C khi loại từng nhóm feature.")
    add_figure(doc, ROOT / "outputs_research_figures" / "fig05_ablation_delta_wrmsse_c.png", "Hình. Delta WRMSSE của C trong ablation study.")
    doc.add_paragraph(
        "Ablation cho thấy calendar/SNAP/event là nhóm feature có tác động lớn nhất. "
        "Khi loại nhóm này, WRMSSE tăng mạnh, phù hợp với đặc điểm bán lẻ chịu ảnh hưởng "
        "lớn của ngày trong tuần, sự kiện và SNAP. Hierarchy/ID cũng quan trọng, còn price "
        "features tạo tác động nhỏ hơn trong cấu hình hiện tại."
    )

    doc.add_heading("7. Kiểm định thống kê và robustness", 2)
    add_table(doc, build_dm_table(), "Bảng. Diebold-Mariano test giữa A0 và C.")
    add_table(doc, build_kmedoids_table(), "Bảng. K-Medoids robustness full-scale.")
    doc.add_paragraph(
        "DM test cho thấy C tốt hơn A0 trên tỷ lệ lớn chuỗi theo mean loss, nhưng không phải "
        "mọi chuỗi đều được cải thiện. K-Medoids robustness cho ARI ở mức thấp-vừa, phản ánh "
        "cấu trúc cụm có độ nhạy khi đổi thuật toán. Vì vậy, cluster được diễn giải như một "
        "phân hoạch vận hành hỗ trợ mô hình, không phải cấu trúc bất biến tuyệt đối."
    )

    doc.add_heading("8. Multi-seed robustness", 2)
    add_table(doc, build_multiseed_summary(), "Bảng. Tổng hợp kết quả multi-seed.")
    add_table(doc, build_paired_tests_table(), "Bảng. Paired tests A0-C qua 5 seed.")
    for path, caption in multiseed_plots:
        add_figure(doc, path, caption)
    doc.add_paragraph(
        "Multi-seed full-scale được chạy với các seed 42, 52, 62, 72 và 82. C có WRMSSE "
        "mean thấp nhất, 0.585790, so với A0 là 0.595689 và B1 là 0.594660. Paired t-test "
        "A0-C có ý nghĩa cho WRMSSE (p=0.011982) và WAPE (p=0.000560). Stability improvement "
        "nhỏ và không có ý nghĩa thống kê rõ, trong khi train-test gap của C cao hơn A0."
    )

    doc.add_heading("9. Ý nghĩa các chỉ số đánh giá", 2)
    add_metric_explanations(doc)

    doc.add_heading("10. Kết luận kết quả", 2)
    doc.add_paragraph(
        "Kết quả thực nghiệm ủng hộ mục tiêu chính của nghiên cứu: cluster-specific LightGBM "
        "C cải thiện độ chính xác dự báo so với global baseline A0 trong thiết kế rolling-origin "
        "kiểm soát leakage. B1 không tạo cải thiện ổn định, cho thấy cluster label chỉ như một "
        "feature chưa đủ để khai thác đầy đủ khác biệt demand regime. Về stability, C có xu hướng "
        "tốt hơn trong seed 42 và một số tổng hợp, nhưng lợi thế này yếu hơn accuracy và không "
        "ổn định như WRMSSE qua multi-seed. Do đó, đóng góp chính của C nằm ở cải thiện accuracy, "
        "trong khi stability cần tiếp tục được tối ưu trong các nghiên cứu sau."
    )

    doc.save(OUT_DOCX)

    md_lines = [
        "# Kết quả nghiên cứu M5 Cluster-aware Forecasting",
        "",
        "File Word đã được tạo kèm bảng kết quả, biểu đồ EDA, biểu đồ mô hình, ablation, robustness và multi-seed.",
        "",
        f"- Word: `{OUT_DOCX}`",
        f"- Figures: `{FIG_DIR}`",
    ]
    OUT_MD.write_text("\n".join(md_lines), encoding="utf-8")


if __name__ == "__main__":
    create_report()
    print(f"Created: {OUT_DOCX}")
    print(f"Created: {OUT_MD}")
