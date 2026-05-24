from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


OUT_DOCX = Path("Document/Bai_Bao_Nghien_Cuu_M5_Hoc_Thuat_Hoan_Thien.docx")
OUT_MD = Path("Document/Bai_Bao_Nghien_Cuu_M5_Hoc_Thuat_Hoan_Thien.md")

BASE = Path("outputs_full_k3_seed42_tweedie_a0_b1_c")
EXT = Path("outputs_extended_fullscale_checks/metrics")
OFFICIAL = Path("outputs_full_official_like_wrmsse/metrics")
FIG = Path("outputs_research_figures")
EDA_FIG = Path("outputs_eda_figures")
METHOD_FIG = Path("outputs_method_figures")


MODEL_LABEL = {
    "A0_global_baseline": "A0: Global LightGBM",
    "B1_cluster_label": "B1: Global LightGBM + cluster label",
    "C_cluster_specific": "C: Cluster-specific LightGBM",
}
MODEL_ORDER = list(MODEL_LABEL)


def fmt(value: float, digits: int = 6) -> str:
    return f"{float(value):.{digits}f}"


def p(doc: Document, text: str, style: str | None = None) -> None:
    para = doc.add_paragraph(style=style)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = para.add_run(text)
    run.font.size = Pt(11)


def h(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, header in enumerate(headers):
        t.rows[0].cells[i].text = header
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


def fig(doc: Document, path: Path, caption: str) -> None:
    if not path.exists():
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(str(path), width=Inches(6.35))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in cap.runs:
        run.font.size = Pt(10)


def eq(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(10)


def load() -> dict[str, pd.DataFrame | pd.Series]:
    all_metrics = pd.read_csv(EXT / "extended_fullscale_all_metrics.csv")
    base = all_metrics[all_metrics["run"].eq("base_tweedie") & all_metrics["model"].isin(MODEL_ORDER)].copy()
    base["model"] = pd.Categorical(base["model"], MODEL_ORDER, ordered=True)
    base = base.sort_values("model")
    ablation = all_metrics[all_metrics["model"].eq("C_cluster_specific") & all_metrics["run"].isin(["no_price", "no_calendar", "no_hierarchy"])].copy()
    ablation["run"] = pd.Categorical(ablation["run"], ["no_price", "no_calendar", "no_hierarchy"], ordered=True)
    ablation = ablation.sort_values("run")
    per_origin = pd.read_csv(OFFICIAL / "official_like_wrmsse_by_origin.csv")
    train_val_test = pd.read_csv(BASE / "metrics" / "model_train_val_test_metrics.csv")
    clusters = pd.read_csv(BASE / "metrics" / "cluster_profiles.csv")
    residual = pd.read_csv(BASE / "metrics" / "residual_bias_diagnostics.csv")
    hierarchy = pd.read_csv(BASE / "metrics" / "sample_hierarchy_metrics.csv")
    bootstrap = pd.read_csv(BASE / "metrics" / "bootstrap_ci_vs_baseline.csv")
    nemenyi = pd.read_csv("outputs_full_nemenyi_tweedie_a0_b1_c/nemenyi_pairwise.csv")
    dm_summary = pd.read_csv(BASE / "metrics" / "dm_test_a0_vs_c_summary.csv")
    dm_group = pd.read_csv(BASE / "metrics" / "dm_test_a0_vs_c_by_group.csv")
    kmed = pd.read_csv(BASE / "metrics" / "kmedoids_robustness.csv")
    kmed_sens = pd.read_csv(BASE / "metrics" / "kmedoids_robustness_sample_sensitivity.csv")
    eda = pd.read_csv(BASE / "eda" / "dataset_overview.csv").set_index("metric")["value"]
    demand = pd.read_csv(BASE / "eda" / "demand_class_counts.csv")
    desc = pd.read_csv(BASE / "eda" / "series_summary_describe.csv", index_col=0)
    config = pd.read_json("configs/research_config_full_k3_seed42_tweedie_a0_b1_c.json", typ="series")
    return {
        "base": base,
        "ablation": ablation,
        "per_origin": per_origin,
        "train_val_test": train_val_test,
        "clusters": clusters,
        "residual": residual,
        "hierarchy": hierarchy,
        "bootstrap": bootstrap,
        "nemenyi": nemenyi,
        "dm_summary": dm_summary,
        "dm_group": dm_group,
        "kmed": kmed,
        "kmed_sens": kmed_sens,
        "eda": eda,
        "demand": demand,
        "desc": desc,
        "config": config,
    }


def pivot_rows(df: pd.DataFrame, index: str, value: str, digits: int = 6) -> list[list[str]]:
    piv = df.pivot(index=index, columns="model", values=value)
    rows: list[list[str]] = []
    for idx, row in piv.iterrows():
        rows.append([str(idx)] + [fmt(row[m], digits) for m in MODEL_ORDER])
    return rows


def main() -> None:
    d = load()
    base = d["base"]  # type: ignore[assignment]
    ablation = d["ablation"]  # type: ignore[assignment]
    per_origin = d["per_origin"]  # type: ignore[assignment]
    train_val_test = d["train_val_test"]  # type: ignore[assignment]
    clusters = d["clusters"]  # type: ignore[assignment]
    residual = d["residual"]  # type: ignore[assignment]
    hierarchy = d["hierarchy"]  # type: ignore[assignment]
    bootstrap = d["bootstrap"]  # type: ignore[assignment]
    nemenyi = d["nemenyi"]  # type: ignore[assignment]
    dm_summary = d["dm_summary"]  # type: ignore[assignment]
    dm_group = d["dm_group"]  # type: ignore[assignment]
    kmed = d["kmed"]  # type: ignore[assignment]
    kmed_sens = d["kmed_sens"]  # type: ignore[assignment]
    eda = d["eda"]  # type: ignore[assignment]
    demand = d["demand"]  # type: ignore[assignment]
    desc = d["desc"]  # type: ignore[assignment]
    config = d["config"]  # type: ignore[assignment]

    c = base[base["model"].eq("C_cluster_specific")].iloc[0]
    a0 = base[base["model"].eq("A0_global_baseline")].iloc[0]
    b1 = base[base["model"].eq("B1_cluster_label")].iloc[0]

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    title = doc.add_heading("Cluster-aware Global Forecasting nhằm cải thiện độ chính xác và tính ổn định trong dự báo nhu cầu bán lẻ quy mô lớn", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p(doc, "Thực nghiệm với LightGBM và dữ liệu M5 Accuracy", None)

    h(doc, "Tóm tắt", 1)
    p(
        doc,
        "Bối cảnh: Dự báo nhu cầu bán lẻ quy mô lớn thường phải xử lý hàng chục nghìn chuỗi item-store có cấu trúc phân cấp, nhiều giá trị 0 và hành vi nhu cầu không đồng nhất. "
        "Khoảng trống: Các nghiên cứu về global forecasting và clustering đã chứng minh lợi ích của việc học trên nhiều chuỗi hoặc nhóm chuỗi tương tự, nhưng vẫn còn thiếu bằng chứng thực nghiệm về việc liệu cluster-specific learning có cải thiện đồng thời WRMSSE và forecast stability trong bối cảnh M5, dưới kiểm soát data leakage và đánh giá rolling-origin hay không. "
        "Phương pháp: Nghiên cứu so sánh ba cấu hình tập trung: A0 - global LightGBM baseline, B1 - global LightGBM với cluster label, và C - LightGBM huấn luyện riêng theo cụm. Phân cụm chính sử dụng Mini-batch K-Means K=3; đánh giá dùng rolling-origin, WRMSSE close-to-official, scale-aware stability loss, ablation, bootstrap confidence interval, Nemenyi post-hoc và K-Medoids robustness full-scale. "
        f"Kết quả: C đạt rolling WRMSSE {fmt(c.sampled_wrmsse_12level)}, thấp hơn A0 ({fmt(a0.sampled_wrmsse_12level)}) và B1 ({fmt(b1.sampled_wrmsse_12level)}), đồng thời đạt stability loss thấp nhất ({fmt(c.scale_aware_stability_loss)}). "
        f"Bootstrap CI cho chênh lệch RMSSE của C so với A0 là [{fmt(bootstrap[(bootstrap.metric.eq('rmsse_item_store')) & (bootstrap.model.eq('C_cluster_specific'))]['ci_2.5'].iloc[0])}, {fmt(bootstrap[(bootstrap.metric.eq('rmsse_item_store')) & (bootstrap.model.eq('C_cluster_specific'))]['ci_97.5'].iloc[0])}], và Nemenyi test cho thấy C tốt hơn A0 có ý nghĩa ở RMSSE. "
        f"Diebold-Mariano test theo từng chuỗi cho thấy C có mean loss thấp hơn A0 trên {fmt(dm_summary[dm_summary.loss.eq('abs_error')]['share_c_better_mean_loss'].iloc[0] * 100, 2)}% chuỗi theo absolute error và {fmt(dm_summary[dm_summary.loss.eq('sq_error')]['share_c_better_mean_loss'].iloc[0] * 100, 2)}% chuỗi theo squared error. "
        "Đóng góp: Bài báo đóng khung cluster-specific forecasting như một cơ chế học theo demand regimes, chứng minh B1 không đủ để khai thác thông tin cụm, và đề xuất stability loss có scale floor để tránh sai lệch trên intermittent demand.",
    )
    p(doc, "Từ khóa: M5, LightGBM, cluster-aware forecasting, global forecasting, WRMSSE, forecast stability, intermittent demand.")

    h(doc, "1. Giới thiệu", 1)
    p(
        doc,
        "Dữ liệu bán lẻ quy mô lớn đặt ra ba thách thức chính cho dự báo: số lượng chuỗi lớn, cấu trúc phân cấp phức tạp và nhu cầu thưa. "
        "Nghiên cứu M5 Accuracy Competition cho thấy machine learning, đặc biệt LightGBM, có vai trò nổi bật trong dự báo bán lẻ và cung cấp bối cảnh chuẩn về dữ liệu phân cấp cùng thước đo WRMSSE. Tuy nhiên, bài này không sử dụng M5 như một bài toán xếp hạng competition, mà dùng M5 làm testbed để đánh giá một framework cluster-aware forecasting dưới rolling-origin evaluation, nơi forecast stability cũng quan trọng như accuracy.",
    )
    p(
        doc,
        "Các nghiên cứu về global forecasting cho thấy mô hình học trên nhiều chuỗi có thể tận dụng cross-learning. Các nghiên cứu dùng clustering hoặc localized models cũng cho thấy việc nhóm các chuỗi tương tự có thể cải thiện forecast. "
        "Tuy nhiên, nhiều nghiên cứu trước tập trung vào accuracy hoặc kiến trúc mô hình, chưa đóng khung rõ bài toán stability trong bối cảnh dữ liệu M5 có intermittent demand và hierarchy. "
        "Khoảng trống nghiên cứu của bài này là đánh giá một framework cluster-aware bằng rolling-origin evaluation, đồng thời kiểm tra accuracy, stability, leakage, overfitting, ablation và robustness của clustering.",
    )
    p(doc, "Câu hỏi nghiên cứu được đặt ra như sau:")
    for rq in [
        "RQ1: Cluster-specific LightGBM có cải thiện WRMSSE so với global baseline A0 hay không?",
        "RQ2: Cluster-specific learning có cải thiện forecast stability qua rolling origins hay không?",
        "RQ3: Việc thêm cluster label như một feature trong global model B1 có đủ để cải thiện kết quả hay không?",
        "RQ4: Kết quả có robust với feature ablation, kiểm định thống kê và clustering method hay không?",
    ]:
        doc.add_paragraph(rq, style="List Bullet")
    p(doc, "Từ các câu hỏi trên, nghiên cứu kiểm tra các giả thuyết: H1 - C có WRMSSE thấp hơn A0; H2 - C có stability loss thấp hơn A0; H3 - B1 không nhất thiết vượt A0; H4 - kết luận chính vẫn giữ khi kiểm tra ablation và robustness.")

    h(doc, "2. Dữ liệu và phân tích khám phá", 1)
    p(
        doc,
        f"Dữ liệu gồm {int(eda['n_series']):,} chuỗi item-store, {int(eda['n_items']):,} sản phẩm, {int(eda['n_stores'])} cửa hàng, {int(eda['n_states'])} bang, {int(eda['n_categories'])} ngành hàng và {int(eda['n_calendar_days']):,} ngày lịch. "
        f"Median zero-sales ratio là {fmt(eda['median_zero_sales_ratio'], 4)}, median ADI là {fmt(desc.loc['50%', 'adi'], 4)} và median CV2 là {fmt(desc.loc['50%', 'cv2'], 4)}.",
    )
    demand_rows = [[str(r.demand_class), str(int(r.n_series))] for r in demand.itertuples()]
    table(doc, ["Demand regime", "Số chuỗi"], demand_rows)
    p(doc, "Tỷ trọng intermittent và lumpy cao cho thấy các thước đo dạng phần trăm thay đổi có nguy cơ bị phóng đại khi demand gần 0. Đây là lý do nghiên cứu dùng scale-aware stability loss thay vì percentage forecast change đơn giản.")
    fig(doc, EDA_FIG / "eda_fig01_demand_class_counts.png", "Hình 1. Phân bố demand regimes trong dữ liệu M5.")
    fig(doc, EDA_FIG / "eda_fig03_adi_cv2_by_demand_class.png", "Hình 2. Bản đồ ADI-CV2 theo demand regimes.")

    h(doc, "3. Phương pháp nghiên cứu", 1)
    h(doc, "3.1. Thiết kế tái lập", 2)
    config_rows = [
        ["Data directory", str(config["data_dir"])],
        ["Output directory", str(config["outputs_dir"])],
        ["Random seed", str(config["random_seed"])],
        ["Selected K", str(config["selected_k"])],
        ["Rolling origins", str(config["rolling_origins"])],
        ["Forecast horizon", str(config["forecast_horizon"])],
        ["Train lookback days", str(config["train_lookback_days"])],
        ["Objective", str(config["lightgbm_params"]["objective"])],
        ["Tweedie variance power", str(config["lightgbm_params"].get("tweedie_variance_power", "default"))],
        ["Boosting rounds", str(config["num_boost_round"])],
    ]
    table(doc, ["Thành phần", "Giá trị"], config_rows)
    p(
        doc,
        "Tại mỗi rolling origin T, pipeline chỉ sử dụng dữ liệu lịch sử đến T để tạo feature, phân cụm và huấn luyện mô hình. Forecast horizon 28 ngày được dự báo theo cơ chế recursive forecasting: lag của ngày tương lai được cập nhật bằng prediction đã sinh ở các bước trước, không dùng actual demand tương lai.",
    )
    fig(doc, METHOD_FIG / "method_workflow.png", "Hình 3. Quy trình phương pháp và kiểm soát anti-leakage.")

    h(doc, "3.2. Mô hình so sánh", 2)
    for item in [
        "A0: Global LightGBM baseline, huấn luyện chung trên toàn bộ panel và không sử dụng cluster label.",
        "B1: Global LightGBM có bổ sung cluster label như một feature.",
        "C: Cluster-specific LightGBM, huấn luyện một mô hình riêng cho từng cụm K=3.",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    p(
        doc,
        "K=3 được chọn để cân bằng giữa khả năng diễn giải và độ ổn định thực nghiệm. Về mặt trực quan, bản đồ ADI-CV2 cho thấy dữ liệu có ba vùng hành vi nổi bật: nhóm nhu cầu thưa/long-tail với ADI và zero-sales ratio cao, nhóm nhu cầu trung bình, và nhóm high-demand/core có ADI thấp hơn. Về mặt lý thuyết, phân loại nhu cầu kinh điển dựa trên ADI và CV2 phân biệt các regime như smooth, intermittent, erratic và lumpy; trong dữ liệu M5, hai nhóm sparse demand chiếm tỷ trọng lớn nên việc gom thành ba cụm vận hành giúp tránh tạo cụm quá nhỏ nhưng vẫn giữ được khác biệt demand regimes chính.",
    )
    p(doc, "Nghiên cứu không trình bày các mô hình ngoài phạm vi A0, B1, C để giữ trọng tâm vào câu hỏi cluster label versus cluster-specific learning.")

    h(doc, "3.3. Công thức đánh giá", 2)
    p(doc, "Với y_i,t là actual demand, ŷ_i,t là forecast, h là horizon, S_i là scale RMSSE của chuỗi i, w_i là trọng số hierarchy, và eps là hằng số nhỏ:")
    eq(doc, "RMSSE_i = sqrt( (1/h) * sum_t (y_i,t - yhat_i,t)^2 / S_i )")
    eq(doc, "WRMSSE = sum_i w_i * RMSSE_i")
    p(doc, "Scale-aware stability loss đo thay đổi forecast cho cùng chuỗi i và target date t giữa hai rolling origins liên tiếp o-1 và o, với m_i là mean sales lịch sử và gamma là scale floor:")
    eq(doc, "Stability = mean_{i,t,o} | yhat_i,t,o - yhat_i,t,o-1 | / max(m_i, gamma)")
    eq(doc, "ADI_i = active_days_i / nonzero_days_i")
    eq(doc, "CV2_i = Var(y_i,t | y_i,t > 0) / Mean(y_i,t | y_i,t > 0)^2")
    eq(doc, "WAPE = sum |y - yhat| / sum |y|")
    eq(doc, "Bias = sum(yhat - y) / sum(y)")
    eq(doc, "Train-test gap = RMSSE_test - RMSSE_train")
    eq(doc, "ARI = adjusted index measuring agreement between two partitions, corrected for chance")
    p(doc, "Diebold-Mariano test được dùng để so sánh A0 và C theo từng chuỗi. Với d_t = L(e_A0,t) - L(e_C,t), giả thuyết H0 là E[d_t] = 0; H1 là E[d_t] > 0, tức C có loss kỳ vọng thấp hơn A0. Phương sai của d_t được ước lượng bằng Newey-West/HAC với lag 27 để phản ánh forecast horizon 28 ngày.")
    eq(doc, "DM = mean(d_t) / sqrt( HACVar(d_t) / n )")

    h(doc, "4. Kết quả", 1)
    h(doc, "4.1. Kết quả tổng hợp của A0, B1 và C", 2)
    base_rows = [
        [MODEL_LABEL[r.model], fmt(r.sampled_wrmsse_12level), fmt(r.validation_origin_wrmsse_mean), fmt(r.scale_aware_stability_loss), fmt(r.wape), fmt(r.bias), fmt(r.mean_test_train_gap)]
        for r in base.itertuples()
    ]
    table(doc, ["Mô hình", "Rolling WRMSSE", "Validation WRMSSE", "Stability loss", "WAPE", "Bias", "Train-test gap"], base_rows)
    p(doc, f"C giảm WRMSSE so với A0 là {fmt(a0.sampled_wrmsse_12level - c.sampled_wrmsse_12level)} và giảm stability loss so với A0 là {fmt(a0.scale_aware_stability_loss - c.scale_aware_stability_loss)}. B1 không cải thiện so với A0, trả lời RQ3 rằng cluster label đơn thuần chưa đủ.")
    fig(doc, FIG / "fig01_base_wrmsse_a0_b1_c.png", "Hình 4. Rolling WRMSSE của A0, B1 và C.")
    fig(doc, FIG / "fig03_accuracy_stability_tradeoff.png", "Hình 5. Trade-off accuracy-stability.")

    h(doc, "4.2. Kết quả theo từng rolling origin", 2)
    origin_rows = pivot_rows(per_origin, "origin", "official_like_wrmsse_12level")
    table(doc, ["Origin", "A0", "B1", "C"], origin_rows)
    p(doc, "Kết quả theo origin cho thấy C đặc biệt tốt ở origin 1913, trong khi ở một số origin sớm hơn chênh lệch giữa các mô hình nhỏ hơn. Điều này cho thấy lợi ích của C nên được đánh giá trên toàn bộ rolling design thay vì một split duy nhất.")

    h(doc, "4.3. Kết quả theo cluster", 2)
    cluster_metric = pd.read_csv(BASE / "metrics" / "model_metrics_by_cluster.csv")
    cm = cluster_metric.groupby(["model", "cluster_label"])[["mae", "wape", "rmsse_item_store", "bias"]].mean().reset_index()
    rows = [[MODEL_LABEL[r.model], str(int(r.cluster_label)), fmt(r.rmsse_item_store), fmt(r.mae), fmt(r.wape), fmt(r.bias)] for r in cm.itertuples()]
    table(doc, ["Mô hình", "Cluster", "RMSSE", "MAE", "WAPE", "Bias"], rows)
    fig(doc, FIG / "fig06_cluster_profile_k3.png", "Hình 6. Profile các cluster K=3.")

    h(doc, "4.4. Kết quả theo demand regime", 2)
    dr = residual[residual["group_type"].eq("demand_class")].copy()
    rows = [[str(r.group), MODEL_LABEL[r.model], fmt(r.mae), fmt(r.bias), fmt(r.p90_abs_error)] for r in dr.itertuples()]
    table(doc, ["Demand regime", "Mô hình", "MAE", "Bias", "P90 absolute error"], rows)
    p(doc, "C có MAE thấp nhất trong cả bốn nhóm erratic, intermittent, lumpy và smooth. Điều này củng cố RQ1 ở cấp demand regime, không chỉ ở metric tổng hợp.")

    h(doc, "4.5. Kết quả theo category, store và state", 2)
    cat = hierarchy[hierarchy["level"].eq("cat_id")]
    rows = [[str(r.group), MODEL_LABEL[r.model], fmt(r.mae), fmt(r.wape), fmt(r.bias)] for r in cat.itertuples()]
    table(doc, ["Category", "Mô hình", "MAE", "WAPE", "Bias"], rows)
    store = hierarchy[hierarchy["level"].eq("store_id")]
    rows = [[str(r.group), MODEL_LABEL[r.model], fmt(r.mae), fmt(r.wape), fmt(r.bias)] for r in store.itertuples()]
    table(doc, ["Store", "Mô hình", "MAE", "WAPE", "Bias"], rows)
    forecasts = pd.read_parquet(BASE / "metrics" / "test_forecasts.parquet")
    forecasts["state_id"] = forecasts["store_id"].str.split("_").str[0]
    forecasts["abs_error"] = (forecasts["y"] - forecasts["yhat"]).abs()
    state_rows = []
    for (state_id, model), g in forecasts.groupby(["state_id", "model"], observed=True):
        state_rows.append([state_id, MODEL_LABEL[model], fmt((g["y"] - g["yhat"]).abs().mean()), fmt((g["y"] - g["yhat"]).abs().sum() / (g["y"].abs().sum() + 1e-9)), fmt((g["yhat"] - g["y"]).sum() / (g["y"].sum() + 1e-9))])
    table(doc, ["State", "Mô hình", "MAE", "WAPE", "Bias"], state_rows)

    h(doc, "4.6. Ablation study", 2)
    base_c = float(c.sampled_wrmsse_12level)
    ab_rows = []
    for r in ablation.itertuples():
        label = {"no_price": "No price", "no_calendar": "No calendar/SNAP/event", "no_hierarchy": "No hierarchy/ID"}[r.run]
        ab_rows.append([label, fmt(r.sampled_wrmsse_12level), fmt(r.validation_origin_wrmsse_mean), fmt(r.scale_aware_stability_loss), fmt(r.sampled_wrmsse_12level - base_c)])
    table(doc, ["Biến thể C", "Rolling WRMSSE", "Validation WRMSSE", "Stability loss", "Delta WRMSSE"], ab_rows)
    p(doc, "No-calendar/SNAP/event gây suy giảm lớn nhất, tiếp theo là no-hierarchy/ID; no-price gây suy giảm nhỏ hơn nhưng làm stability xấu đi. Kết quả trả lời RQ4 rằng kết luận về feature importance nhất quán với logic bán lẻ.")
    fig(doc, FIG / "fig04_ablation_wrmsse_c.png", "Hình 7. WRMSSE của C trong ablation study.")

    h(doc, "4.7. Confidence interval và kiểm định thống kê", 2)
    ci_rows = [[r.metric, MODEL_LABEL[r.model], fmt(r.mean_diff), fmt(r["ci_2.5"]), fmt(r["ci_97.5"])] for _, r in bootstrap.iterrows() if r["model"] in MODEL_ORDER]
    table(doc, ["Metric", "Mô hình so với A0", "Mean diff", "CI 2.5%", "CI 97.5%"], ci_rows)
    nem_rows = [
        [
            str(r["metric"]),
            MODEL_LABEL[str(r["model_a"])],
            MODEL_LABEL[str(r["model_b"])],
            fmt(r["rank_diff"]),
            fmt(r["critical_difference"]),
            str(bool(r["significant_alpha_0.05"])),
        ]
        for _, r in nemenyi.iterrows()
    ]
    table(doc, ["Metric", "Model A", "Model B", "Rank diff", "Critical difference", "Significant 0.05"], nem_rows)
    p(doc, "Bootstrap CI theo origin cho thấy C cải thiện RMSSE, MAE và WAPE so với A0 với khoảng tin cậy không cắt 0. Nemenyi post-hoc chỉ xác nhận khác biệt có ý nghĩa giữa C và A0 trên RMSSE; các metric khác cần diễn giải thận trọng do chỉ có 5 rolling origins.")

    h(doc, "4.8. Diebold-Mariano test giữa A0 và C", 2)
    dm_rows = [
        [
            "Absolute error" if r.loss == "abs_error" else "Squared error",
            str(int(r.n_series)),
            fmt(r.mean_loss_diff_a0_minus_c),
            fmt(r.median_loss_diff_a0_minus_c),
            fmt(r.share_c_better_mean_loss),
            fmt(r.share_c_better_p05),
            fmt(r.share_c_better_fdr05),
        ]
        for r in dm_summary.itertuples()
    ]
    table(doc, ["Loss", "Số chuỗi", "Mean loss diff A0-C", "Median loss diff", "Share C better", "Share p<0.05", "Share FDR<0.05"], dm_rows)
    p(
        doc,
        "DM test được thực hiện ở cấp chuỗi item-store, sử dụng toàn bộ cặp forecast A0 và C qua 5 rolling origins và horizon 28 ngày. "
        "Mean loss diff dương nghĩa là loss của A0 lớn hơn loss của C. Kết quả cho thấy C có loss trung bình thấp hơn A0 trên đa số chuỗi: 59.28% chuỗi theo absolute error và 56.49% chuỗi theo squared error. "
        "Tỷ lệ chuỗi có C tốt hơn với p<0.05 là 31.25% theo absolute error và 21.70% theo squared error; sau hiệu chỉnh FDR, tỷ lệ còn 21.19% và 8.45%. "
        "Do đó, DM test cung cấp bằng chứng bổ sung rằng lợi thế của C không chỉ đến từ may mắn ở một vài origin, nhưng kết quả cũng cho thấy hiệu quả của C không đồng đều trên toàn bộ chuỗi.",
    )
    dmg = dm_group[dm_group["loss"].eq("abs_error") & dm_group["group_type"].isin(["cat_id", "state_id", "cluster_label"])].copy()
    dm_group_rows = [
        [
            str(r.group_type),
            str(r.group),
            str(int(r.n_series)),
            fmt(r.mean_loss_diff_a0_minus_c),
            fmt(r.share_c_better_mean_loss),
            fmt(r.share_c_better_p05),
            fmt(r.share_c_better_fdr05),
        ]
        for r in dmg.itertuples()
    ]
    table(doc, ["Group type", "Group", "Số chuỗi", "Mean loss diff A0-C", "Share C better", "Share p<0.05", "Share FDR<0.05"], dm_group_rows)
    p(doc, "Theo nhóm, C có mean absolute loss thấp hơn A0 ở hầu hết category, state và cluster; ngoại lệ đáng chú ý là cluster 1 có mean loss diff âm nhẹ theo absolute error. Điều này phù hợp với kết luận tổng quát: C tốt hơn về trung bình nhưng không thống trị tuyệt đối mọi phân đoạn.")

    h(doc, "4.9. Robustness clustering", 2)
    rows = [[str(int(r.sample_n)), fmt(r.ari_kmeans_vs_kmedoids), str(int(r.kmedoids_min_cluster_size)), str(int(r.kmedoids_max_cluster_size)), str(r.note)] for r in kmed_sens.itertuples()]
    table(doc, ["Sample", "ARI", "Min cluster size", "Max cluster size", "Ghi chú"], rows)
    p(doc, f"K-Medoids full-scale trên {int(kmed.iloc[0].sample_n):,} chuỗi cho ARI {fmt(kmed.iloc[0].ari_kmeans_vs_kmedoids)}. Kết quả này cho thấy tương đồng thấp-vừa, nên cụm được diễn giải như demand regimes phục vụ mô hình hóa, không phải phân khúc kinh doanh cố định.")

    h(doc, "5. Thảo luận", 1)
    h(doc, "5.1. Vì sao C hoạt động tốt hơn", 2)
    p(doc, "C hoạt động tốt hơn vì cluster-specific learning cho phép các chuỗi long-tail/intermittent và core/high-demand được học bởi các mô hình khác nhau. Điều này giảm áp lực cho một global learner phải mô tả đồng thời các cơ chế nhu cầu khác biệt.")
    h(doc, "5.2. Vì sao B1 không đủ", 2)
    p(doc, "B1 chỉ đưa cluster label vào như một biến giải thích. Với mô hình cây, label này có thể hỗ trợ split nhưng không thay đổi toàn bộ cấu trúc học, loss landscape và feature interaction theo từng demand regime. Do đó, B1 không vượt A0 trong kết quả chính.")
    h(doc, "5.3. Hàm ý vận hành", 2)
    p(doc, "Trong vận hành, C phù hợp khi doanh nghiệp ưu tiên cả accuracy và stability, chấp nhận chi phí huấn luyện nhiều mô hình theo cụm. A0 có thể là baseline đơn giản hơn; B1 không nên được xem là giải pháp cluster-aware đủ mạnh nếu không có cải thiện thực nghiệm.")
    h(doc, "5.4. Hạn chế và khả năng khái quát", 2)
    for item in [
        "Đánh giá là close-to-official local evaluation phục vụ mục tiêu nghiên cứu framework, không phải một quy trình competition submission.",
        "Số rolling origins là 5, nên kiểm định thống kê có power hạn chế.",
        "K-Medoids full-scale cho ARI thấp-vừa, nên cần tránh overclaim về ý nghĩa quản trị của cụm.",
        "Khả năng khái quát sang ngành bán lẻ khác cần kiểm tra thêm trên dữ liệu ngoài M5.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    h(doc, "6. Kết luận", 1)
    p(doc, "Nghiên cứu chứng minh rằng cluster-specific LightGBM cải thiện đồng thời WRMSSE và forecast stability so với global baseline trong thiết kế rolling-origin trên M5. Đóng góp chính là kết hợp cluster-aware forecasting với scale-aware stability loss, kiểm soát leakage và đánh giá robustness. Kết quả cũng làm rõ rằng cluster label như một feature chưa đủ thay thế cluster-specific learning.")

    h(doc, "Tài liệu tham khảo", 1)
    refs = [
        "Bandara, K., Bergmeir, C., & Smyl, S. (2020). Forecasting across time series databases using recurrent neural networks on groups of similar series: A clustering approach. Expert Systems with Applications, 140, 112896.",
        "Breiman, L. (1996). Bagging predictors. Machine Learning, 24, 123-140.",
        "Croston, J. D. (1972). Forecasting and stock control for intermittent demands. Operational Research Quarterly, 23(3), 289-303.",
        "Fildes, R., Ma, S., & Kolassa, S. (2019). Retail forecasting: Research and practice. International Journal of Forecasting, 35(1), 1-7.",
        "Godahewa, R., Bandara, K., Webb, G. I., Smyl, S., & Bergmeir, C. (2021). Ensembles of localised models for time series forecasting.",
        "Hewamalage, H., Bergmeir, C., & Bandara, K. (2021). Global models for time series forecasting: A simulation study.",
        "Hyndman, R. J., & Koehler, A. B. (2006). Another look at measures of forecast accuracy. International Journal of Forecasting, 22(4), 679-688.",
        "Ke, G., Meng, Q., Finley, T., Wang, T., Chen, W., Ma, W., Ye, Q., & Liu, T. Y. (2017). LightGBM: A highly efficient gradient boosting decision tree. Advances in Neural Information Processing Systems.",
        "Makridakis, S., Spiliotis, E., & Assimakopoulos, V. (2022). M5 accuracy competition: Results, findings, and conclusions. International Journal of Forecasting, 38(4), 1346-1364.",
        "Montero-Manso, P., Athanasopoulos, G., Hyndman, R. J., & Talagala, T. S. (2020). FFORMA: Feature-based forecast model averaging. International Journal of Forecasting, 36(1), 86-92.",
        "Semenoglou, A. A., Spiliotis, E., Makridakis, S., & Assimakopoulos, V. (2021). Investigating the accuracy of cross-learning time series forecasting methods. International Journal of Forecasting.",
        "Syntetos, A. A., Boylan, J. E., & Croston, J. D. (2005). On the categorization of demand patterns. Journal of the Operational Research Society, 56, 495-503.",
        "Talagala, T. S., Hyndman, R. J., & Athanasopoulos, G. (2021). Meta-learning how to forecast time series. Journal of Forecasting.",
        "Tibshirani, R., Walther, G., & Hastie, T. (2001). Estimating the number of clusters in a data set via the gap statistic. Journal of the Royal Statistical Society: Series B, 63(2), 411-423.",
        "Vermorel, J. (2013). Quantile forecasting for retail inventory optimization. International Journal of Forecasting, 29(4), 595-604.",
    ]
    for ref in refs:
        doc.add_paragraph(ref, style="List Number")

    doc.save(OUT_DOCX)

    md = [
        "# Bài báo nghiên cứu M5 hoàn thiện",
        "",
        "Bản Word chứa nội dung học thuật đầy đủ: research gap, RQ/hypotheses, công thức, methodology tái lập, kết quả theo rolling origin/cluster/demand regime/category/store/state, confidence interval, kiểm định thống kê, discussion và references mở rộng.",
        "",
        f"File Word: `{OUT_DOCX}`",
    ]
    OUT_MD.write_text("\n".join(md), encoding="utf-8")
    print(OUT_DOCX)
    print(OUT_MD)


if __name__ == "__main__":
    main()
