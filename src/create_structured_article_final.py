from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


OUT_DOCX = Path("Document/Bai_Bao_Nghien_Cuu_M5_Cau_Truc_Chuan_Hoc_Thuat.docx")
OUT_MD = Path("Document/Bai_Bao_Nghien_Cuu_M5_Cau_Truc_Chuan_Hoc_Thuat.md")

BASE = Path("outputs_full_k3_seed42_tweedie_a0_b1_c")
EXT = Path("outputs_extended_fullscale_checks/metrics")
OFFICIAL = Path("outputs_full_official_like_wrmsse/metrics")
FIG = Path("outputs_research_figures")
EDA_FIG = Path("outputs_eda_figures")
METHOD_FIG = Path("outputs_method_figures")

MODEL_ORDER = ["A0_global_baseline", "B1_cluster_label", "C_cluster_specific"]
MODEL_LABEL = {
    "A0_global_baseline": "A0: Global LightGBM",
    "B1_cluster_label": "B1: Global LightGBM + cluster label",
    "C_cluster_specific": "C: Cluster-specific LightGBM",
}


def fmt(x: float, digits: int = 6) -> str:
    return f"{float(x):.{digits}f}"


def add_p(doc: Document, text: str, style: str | None = None, center: bool = False) -> None:
    para = doc.add_paragraph(style=style)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    run = para.add_run(text)
    run.font.size = Pt(11)


def add_h(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    doc.add_paragraph()


def add_fig(doc: Document, path: Path, caption: str) -> None:
    if not path.exists():
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(str(path), width=Inches(6.3))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in cap.runs:
        run.font.size = Pt(10)


def add_eq(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(10)


def load_data() -> dict[str, pd.DataFrame | pd.Series]:
    all_metrics = pd.read_csv(EXT / "extended_fullscale_all_metrics.csv")
    base = all_metrics[all_metrics["run"].eq("base_tweedie") & all_metrics["model"].isin(MODEL_ORDER)].copy()
    base["model"] = pd.Categorical(base["model"], MODEL_ORDER, ordered=True)
    base = base.sort_values("model")
    ablation = all_metrics[all_metrics["model"].eq("C_cluster_specific") & all_metrics["run"].isin(["no_price", "no_calendar", "no_hierarchy"])].copy()
    ablation["run"] = pd.Categorical(ablation["run"], ["no_price", "no_calendar", "no_hierarchy"], ordered=True)
    ablation = ablation.sort_values("run")
    return {
        "base": base,
        "ablation": ablation,
        "eda": pd.read_csv(BASE / "eda" / "dataset_overview.csv").set_index("metric")["value"],
        "demand": pd.read_csv(BASE / "eda" / "demand_class_counts.csv"),
        "desc": pd.read_csv(BASE / "eda" / "series_summary_describe.csv", index_col=0),
        "per_origin": pd.read_csv(OFFICIAL / "official_like_wrmsse_by_origin.csv"),
        "cluster_profiles": pd.read_csv(BASE / "metrics" / "cluster_profiles.csv"),
        "cluster_metrics": pd.read_csv(BASE / "metrics" / "model_metrics_by_cluster.csv"),
        "hierarchy": pd.read_csv(BASE / "metrics" / "sample_hierarchy_metrics.csv"),
        "residual": pd.read_csv(BASE / "metrics" / "residual_bias_diagnostics.csv"),
        "bootstrap": pd.read_csv(BASE / "metrics" / "bootstrap_ci_vs_baseline.csv"),
        "nemenyi": pd.read_csv("outputs_full_nemenyi_tweedie_a0_b1_c/nemenyi_pairwise.csv"),
        "dm_summary": pd.read_csv(BASE / "metrics" / "dm_test_a0_vs_c_summary.csv"),
        "dm_group": pd.read_csv(BASE / "metrics" / "dm_test_a0_vs_c_by_group.csv"),
        "kmed": pd.read_csv(BASE / "metrics" / "kmedoids_robustness.csv"),
        "kmed_sens": pd.read_csv(BASE / "metrics" / "kmedoids_robustness_sample_sensitivity.csv"),
        "seed_per_seed": pd.read_csv("outputs_full_multiseed_summary/metrics/full_multiseed_per_seed_metrics.csv"),
        "seed_agg": pd.read_csv("outputs_full_multiseed_summary/metrics/full_multiseed_aggregate_metrics.csv"),
        "seed_paired": pd.read_csv("outputs_full_multiseed_summary/metrics/full_multiseed_paired_tests_a0_vs_c.csv"),
        "seed_nemenyi": pd.read_csv("outputs_full_multiseed_summary/metrics/full_multiseed_nemenyi_posthoc.csv"),
        "config": pd.read_json("configs/research_config_full_k3_seed42_tweedie_a0_b1_c.json", typ="series"),
    }


def main() -> None:
    data = load_data()
    base = data["base"]  # type: ignore[assignment]
    ablation = data["ablation"]  # type: ignore[assignment]
    eda = data["eda"]  # type: ignore[assignment]
    demand = data["demand"]  # type: ignore[assignment]
    desc = data["desc"]  # type: ignore[assignment]
    per_origin = data["per_origin"]  # type: ignore[assignment]
    cluster_profiles = data["cluster_profiles"]  # type: ignore[assignment]
    cluster_metrics = data["cluster_metrics"]  # type: ignore[assignment]
    hierarchy = data["hierarchy"]  # type: ignore[assignment]
    residual = data["residual"]  # type: ignore[assignment]
    bootstrap = data["bootstrap"]  # type: ignore[assignment]
    nemenyi = data["nemenyi"]  # type: ignore[assignment]
    dm_summary = data["dm_summary"]  # type: ignore[assignment]
    dm_group = data["dm_group"]  # type: ignore[assignment]
    kmed = data["kmed"]  # type: ignore[assignment]
    kmed_sens = data["kmed_sens"]  # type: ignore[assignment]
    seed_per_seed = data["seed_per_seed"]  # type: ignore[assignment]
    seed_agg = data["seed_agg"]  # type: ignore[assignment]
    seed_paired = data["seed_paired"]  # type: ignore[assignment]
    seed_nemenyi = data["seed_nemenyi"]  # type: ignore[assignment]
    config = data["config"]  # type: ignore[assignment]

    a0 = base[base["model"].eq("A0_global_baseline")].iloc[0]
    b1 = base[base["model"].eq("B1_cluster_label")].iloc[0]
    c = base[base["model"].eq("C_cluster_specific")].iloc[0]

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    title = doc.add_heading(
        "Cluster-aware Global Forecasting nhằm cải thiện độ chính xác và tính ổn định trong dự báo nhu cầu bán lẻ quy mô lớn",
        0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_p(doc, "Thực nghiệm với LightGBM và dữ liệu M5", center=True)

    add_h(doc, "Abstract", 1)
    abstract = (
        "Dự báo nhu cầu bán lẻ quy mô lớn đòi hỏi mô hình xử lý đồng thời tính phân cấp, nhu cầu thưa và dị thể giữa các chuỗi item-store. "
        "Nghiên cứu này đánh giá một framework cluster-aware global forecasting trên dữ liệu M5, tập trung vào câu hỏi liệu cluster-specific learning có cải thiện độ chính xác và độ ổn định so với global LightGBM hay không. "
        "Ba cấu hình được so sánh: A0 global LightGBM, B1 global LightGBM với cluster label, và C cluster-specific LightGBM. "
        "Thiết kế thực nghiệm dùng rolling-origin evaluation, WRMSSE 12 cấp, scale-aware stability loss, ablation, Diebold-Mariano test, Nemenyi post-hoc, K-Medoids robustness và full multi-seed evaluation. "
        f"Trên cấu hình full-scale chính, C đạt WRMSSE {fmt(c.sampled_wrmsse_12level)}, thấp hơn A0 ({fmt(a0.sampled_wrmsse_12level)}) và B1 ({fmt(b1.sampled_wrmsse_12level)}). "
        "Full multi-seed cho thấy cải thiện WRMSSE của C so với A0 có ý nghĩa thống kê, trong khi lợi thế stability nhỏ hơn và không ổn định ở mọi seed. "
        "Kết quả cho thấy tách mô hình theo demand regime hữu ích hơn việc chỉ thêm cluster label vào global learner."
    )
    add_p(doc, abstract)
    add_p(doc, f"Số từ Abstract: {len(abstract.split())} từ.")
    add_p(doc, "Keywords: M5, LightGBM, global forecasting, cluster-aware forecasting, WRMSSE, forecast stability, intermittent demand.")

    add_h(doc, "1. Introduction", 1)
    add_p(
        doc,
        "Dự báo nhu cầu bán lẻ quy mô lớn là nền tảng của quản trị tồn kho, phân bổ hàng hóa và lập kế hoạch cung ứng. "
        "Với chuỗi nhu cầu y_i,t ở cấp item-store, bài toán là dự báo y_i,t+h trong horizon 28 ngày sao cho sai số thấp ở nhiều cấp phân cấp và forecast không biến động quá mạnh khi origin dịch chuyển. "
        "Khó khăn đến từ zero sales, vòng đời sản phẩm, sự kiện lịch, SNAP, biến động giá và dị biệt giữa cửa hàng, bang và ngành hàng.",
    )
    add_p(
        doc,
        "Global forecasting tận dụng thông tin liên chuỗi nhưng buộc các demand regimes rất khác nhau đi qua cùng một hàm học. "
        "Bài nghiên cứu kiểm tra liệu phân cụm chuỗi nhu cầu và huấn luyện mô hình theo cụm có tạo lợi ích thực nghiệm so với một global model duy nhất hay không.",
    )
    add_p(doc, "Các câu hỏi nghiên cứu được thiết kế riêng rẽ, tránh trộn nhiều mệnh đề trong cùng một câu hỏi:")
    for rq in [
        "RQ1: C có làm giảm WRMSSE so với A0 không?",
        "RQ2: C có làm giảm scale-aware stability loss so với A0 không?",
        "RQ3: B1 có cải thiện WRMSSE so với A0 không?",
        "RQ4: Ablation feature có làm thay đổi WRMSSE của C không?",
        "RQ5: Phân cụm K=3 có ổn định khi thay đổi thuật toán không?",
    ]:
        doc.add_paragraph(rq, style="List Bullet")
    add_p(doc, "Các giả thuyết tương ứng là H1: C có WRMSSE thấp hơn A0; H2: C có stability loss thấp hơn A0; H3: B1 không nhất thiết vượt A0; H4: calendar/SNAP/event và hierarchy/ID là các nhóm feature quan trọng; H5: cấu trúc cụm có độ nhạy khi đổi thuật toán.")

    add_h(doc, "2. Related work", 1)
    related_rows = [
        [
            "Makridakis et al. (2022)",
            "Tổng kết M5 Accuracy Competition, dữ liệu Walmart, WRMSSE, hierarchy forecasting.",
            "LightGBM và ensemble là hướng hiệu quả trong M5; dữ liệu có tính phân cấp, intermittent và erratic.",
            "Trọng tâm là kết quả competition và accuracy; chưa tập trung vào forecast stability của cluster-aware framework.",
        ],
        [
            "Ma and Fildes (2022)",
            "Kiểm tra robustness của global bottom-up approach trên M5 bằng nhiều rolling test sets.",
            "Global LightGBM bottom-up mạnh ở bottom level nhưng thiếu ổn định khi tổng hợp lên các cấp phân cấp.",
            "Gợi ý cần đánh giá qua nhiều origins; chưa xử lý riêng cluster-aware stability theo demand regime.",
        ],
        [
            "E et al. (2022)",
            "Dynamic model selection theo demand pattern dựa trên ADI và CV2 trong retail sales forecasting.",
            "Không một mô hình nào phù hợp cho mọi demand pattern; intermittent/lumpy cần chiến lược khác smooth/erratic.",
            "Tập trung chọn/weight mô hình thống kê, chưa đánh giá cluster-specific LightGBM trên M5 full-scale.",
        ],
        [
            "Van Ruitenbeek et al. (2023)",
            "Hierarchical agglomerative clustering cho product sales forecasting.",
            "Clustering hữu ích với intermittent demand và variation lớn; hiệu quả phụ thuộc đặc trưng chuỗi.",
            "Chủ yếu dùng cluster-based aggregation; chưa kiểm tra recursive LightGBM và WRMSSE-stability.",
        ],
        [
            "Hoeltgebaum et al. (2023)",
            "Score-driven models cho lumpy và intermittent retail demand.",
            "Nhấn mạnh tầm quan trọng của zero-inflated/intermittent retail demand và mô hình hóa phần zero/non-zero.",
            "Không tập trung vào global tabular ML hoặc phân cụm demand regime ở quy mô M5.",
        ],
        [
            "Mitchell et al. (2024)",
            "Demetra: hierarchical retail demand forecasting kết hợp top-down và bottom-level machine learning.",
            "Khai thác bổ sung giữa bottom-level ML và top-down time-series methods trong dữ liệu retail rất dị thể.",
            "Tập trung hierarchical ensemble trong bối cảnh doanh nghiệp; không kiểm tra cluster-specific learning như contribution chính.",
        ],
        [
            "Klee and Xia (2025)",
            "Đo forecast stability trong demand planning bằng biến thiên forecast qua nhiều runs.",
            "Nhấn mạnh trade-off accuracy-stability và rủi ro metric bị méo khi forecast gần 0.",
            "Đo stability theo multi-run CV; bài này mở rộng theo rolling-origin scale-aware change cho intermittent demand.",
        ],
        [
            "Recent global-local LightGBM study (2025)",
            "So sánh global pooling, cluster-level pooling và local modeling trong demand forecasting.",
            "Global LightGBM với identifier có thể vượt cluster-level pooling khi cụm còn dị thể.",
            "Củng cố nhu cầu kiểm chứng B1/C thay vì giả định clustering luôn cải thiện.",
        ],
        [
            "Bandara et al. (2020)",
            "Dùng clustering để chia nhóm chuỗi tương tự rồi huấn luyện RNN/LSTM theo nhóm.",
            "Clustering giúp giảm bất lợi của dữ liệu dị thể trong global neural forecasting.",
            "Không đánh giá LightGBM trên M5 với WRMSSE-stability và kiểm soát leakage rolling-origin như bài này.",
        ],
        [
            "Montero-Manso et al. (2020)",
            "FFORMA dùng đặc trưng chuỗi để học trọng số kết hợp mô hình dự báo.",
            "Feature-based learning giúp chọn hoặc kết hợp forecast model theo đặc điểm chuỗi.",
            "Không trực tiếp kiểm tra cluster-specific LightGBM hoặc stability loss trên intermittent retail panel.",
        ],
        [
            "Syntetos et al. (2005)",
            "Phân loại demand pattern bằng ADI và CV2.",
            "Cung cấp nền tảng lý thuyết để phân biệt smooth, intermittent, erratic và lumpy demand.",
            "Không phải framework dự báo global-scale; cần tích hợp với machine learning và evaluation hiện đại.",
        ],
        [
            "Ke et al. (2017)",
            "Đề xuất LightGBM, gradient boosting tree hiệu quả trên dữ liệu lớn.",
            "LightGBM phù hợp cho dữ liệu tabular quy mô lớn và nhiều feature.",
            "Không giải quyết riêng bài toán cluster-aware forecasting hoặc intermittent stability.",
        ],
    ]
    add_table(doc, ["Nghiên cứu", "Phương pháp/nội dung", "Kết quả chính", "Hạn chế liên quan"], related_rows)
    add_p(
        doc,
        "Khoảng trống của bài nằm ở giao điểm giữa cluster-specific learning, retail hierarchy và forecast stability. "
        "Các nghiên cứu trước đã chứng minh sức mạnh của global LightGBM, feature-based learning và clustering, nhưng ít đánh giá cluster-specific LightGBM trên M5 full-scale với rolling-origin leakage control, WRMSSE 12 cấp và kiểm định robustness.",
    )
    add_p(
        doc,
        "Hai điểm từ literature định hình thiết kế thực nghiệm. Global LightGBM là baseline mạnh trong retail panel, còn clustering chỉ hữu ích khi phân hoạch làm giảm dị thể có ý nghĩa cho learner. "
        "Contribution của bài nằm ở kiểm định trực tiếp cách dùng cluster trong mô hình dự báo: cluster label trong global model hay mô hình riêng theo demand regime.",
    )

    add_h(doc, "3. Background", 1)
    add_p(
        doc,
        "Bài toán có động cơ vận hành rõ ràng: sai số dự báo dẫn đến tồn kho dư hoặc thiếu hàng, còn forecast dao động mạnh làm kế hoạch bổ sung hàng kém ổn định. "
        "M5 cung cấp đầy đủ các tín hiệu cần thiết cho ngữ cảnh này: sales, calendar, SNAP, event, sell price, store, state, category và department.",
    )
    add_p(
        doc,
        "Một chuỗi item-store phản ánh đồng thời đặc tính sản phẩm và bối cảnh bán tại cửa hàng. "
        "Các chuỗi khác category hoặc state vẫn có thể chia sẻ demand regime, trong khi các sản phẩm cùng category có thể khác nhau mạnh do giá, SNAP hoặc sự kiện địa phương. "
        "Cluster-aware learning khai thác chính cấu trúc dị thể này.",
    )
    add_p(
        doc,
        "Intermittent demand là đặc trưng trung tâm của M5. Nhiều chuỗi có tỷ lệ zero sales cao; các metric phần trăm dễ méo khi mẫu số nhỏ, còn learner có thể tối ưu bằng cách dự báo gần 0. "
        "ADI, CV2, zero-sales ratio và scale-aware stability được dùng để mô tả độ thưa, biến động và ổn định theo thang đo phù hợp hơn.",
    )
    add_p(
        doc,
        "Nền tảng phương pháp gồm LightGBM, Mini-batch K-Means, K-Medoids robustness, ADI-CV2 demand classification và rolling-origin evaluation. "
        "LightGBM phù hợp với dữ liệu tabular lớn; Mini-batch K-Means cung cấp phân cụm có khả năng mở rộng; K-Medoids đóng vai trò kiểm tra độ nhạy hậu nghiệm.",
    )
    add_p(
        doc,
        "ADI đo khoảng cách trung bình giữa các lần bán khác 0, còn CV2 đo mức biến động tương đối của lượng bán trong các ngày có phát sinh nhu cầu. "
        "Theo cách phân loại nhu cầu kinh điển, ADI cao thường liên quan đến nhu cầu gián đoạn, trong khi CV2 cao liên quan đến nhu cầu erratic hoặc lumpy. "
        "Trong nghiên cứu này, ADI-CV2 không được dùng như nhãn mục tiêu mà được dùng như nền tảng lý thuyết để diễn giải cluster và kiểm tra xem các cụm học được có phản ánh khác biệt demand behavior hay không.",
    )
    add_p(
        doc,
        "WRMSSE phản ánh bản chất phân cấp của M5. Metric này đánh giá đồng thời item-store và các cấp tổng hợp như store, state, category và department, nhờ đó hạn chế việc tối ưu cục bộ ở bottom level nhưng làm sai lệch các aggregate quan trọng.",
    )
    add_p(
        doc,
        "K=3 cân bằng giữa khả năng diễn giải và kích thước cụm. Bản đồ ADI-CV2 gợi ý ba vùng vận hành chính: long-tail/intermittent, medium demand và high-demand/core. "
        "Cách chọn này tránh tạo các cụm quá nhỏ khi huấn luyện C.",
    )
    add_fig(doc, EDA_FIG / "eda_fig03_adi_cv2_by_demand_class.png", "Hình 1. Bản đồ ADI-CV2 hỗ trợ lựa chọn K=3 theo các demand regimes chính.")

    add_h(doc, "4. Research methodology", 1)
    add_fig(doc, METHOD_FIG / "method_workflow.png", "Hình 2. Mô hình phương pháp nghiên cứu.")
    add_p(
        doc,
        "Framework gồm năm lớp: data validation và EDA; feature engineering; clustering theo từng rolling origin; huấn luyện A0, B1 và C; cuối cùng là evaluation bằng WRMSSE, stability, overfitting gap, ablation và kiểm định thống kê.",
    )
    add_p(
        doc,
        "Tại mỗi origin T, mọi feature dựa trên actual sales chỉ dùng dữ liệu đến T. "
        "Lag, rolling mean, ADI, CV2, mean sales và zero-sales ratio đều được tạo theo nguyên tắc này. Calendar, event, SNAP và weekly price được xử lý như biến ngoại sinh biết trước trong bối cảnh M5.",
    )
    add_p(
        doc,
        "Ba cấu hình mô hình được thiết kế để tách riêng vai trò của thông tin cụm. A0 là baseline global LightGBM không dùng cluster, đóng vai trò mốc so sánh chính. "
        "B1 giữ cùng cấu trúc global model như A0 nhưng bổ sung cluster label như một categorical feature, qua đó kiểm tra liệu chỉ đưa nhãn cụm vào feature space có đủ tạo cải thiện hay không. "
        "C huấn luyện mô hình riêng cho từng cụm, qua đó cho phép mỗi demand regime có hàm học và cấu trúc split khác nhau trong LightGBM.",
    )
    add_p(
        doc,
        "C giảm dị thể trong từng tập huấn luyện con nhưng làm tăng độ phức tạp và rủi ro phân mảnh dữ liệu. "
        "Accuracy được đọc cùng train-test gap, stability, feature importance và robustness của clustering.",
    )
    add_p(
        doc,
        "Forecast 28 ngày được sinh đệ quy: dự báo của h+1 trở thành history để tạo lag/rolling cho h+2. "
        "Thiết kế này phản ánh triển khai thực tế và tránh teacher forcing trong horizon.",
    )
    add_p(doc, "Mã giả quy trình nghiên cứu:")
    pseudo_rows = [
        ["1", "For each rolling origin T"],
        ["2", "Use only historical data up to T"],
        ["3", "Create lag, rolling, calendar, price, availability and hierarchy features"],
        ["4", "Compute clustering features from historical data only"],
        ["5", "Assign K=3 clusters by Mini-batch K-Means"],
        ["6", "Train A0, B1 and C"],
        ["7", "Generate 28-day recursive forecasts"],
        ["8", "Evaluate WRMSSE, stability, overfitting, ablation and statistical tests"],
    ]
    add_table(doc, ["Bước", "Mã giả"], pseudo_rows)

    add_h(doc, "5. Experiment", 1)
    add_h(doc, "5.1. Mô tả thực nghiệm", 2)
    add_p(
        doc,
        f"Dữ liệu gồm {int(eda['n_series']):,} chuỗi item-store, {int(eda['n_items']):,} sản phẩm, {int(eda['n_stores'])} cửa hàng, {int(eda['n_states'])} bang, {int(eda['n_categories'])} ngành hàng, {int(eda['n_departments'])} department và {int(eda['n_calendar_days']):,} ngày lịch. "
        f"Median zero-sales ratio là {fmt(eda['median_zero_sales_ratio'], 4)}; median ADI là {fmt(desc.loc['50%', 'adi'], 4)}; median CV2 là {fmt(desc.loc['50%', 'cv2'], 4)}.",
    )
    add_p(
        doc,
        "Dữ liệu sales được sử dụng ở dạng evaluation vì có actual sales đến d_1941, cho phép tạo các rolling-origin test windows sau d_1913. "
        "Mỗi rolling origin tạo một bài toán dự báo 28 ngày; các origin cách nhau 7 ngày để vừa kiểm tra được nhiều trạng thái thời gian, vừa giữ chi phí huấn luyện ở mức khả thi cho full-scale 30,490 chuỗi. "
        "Cách này không nhằm mô phỏng một lần nộp kết quả duy nhất, mà nhằm đánh giá hành vi mô hình khi origin thay đổi.",
    )
    demand_rows = [[str(r.demand_class), str(int(r.n_series))] for r in demand.itertuples()]
    add_table(doc, ["Demand regime", "Số chuỗi"], demand_rows)
    config_rows = [
        ["Config", "configs/research_config_full_k3_seed42_tweedie_a0_b1_c.json"],
        ["Rolling origins", str(config["rolling_origins"])],
        ["Forecast horizon", str(config["forecast_horizon"])],
        ["Selected K", str(config["selected_k"])],
        ["Objective", str(config["lightgbm_params"]["objective"])],
        ["Random seed", str(config["random_seed"])],
        ["Output chính", "outputs_full_k3_seed42_tweedie_a0_b1_c/"],
    ]
    add_table(doc, ["Thành phần", "Giá trị"], config_rows)
    add_p(doc, "Field chính của dữ liệu gồm id, item_id, dept_id, cat_id, store_id, state_id, d_1...d_1941 trong sales; date, wm_yr_wk, weekday, event, SNAP trong calendar; item-store-week price trong sell_prices.")
    add_p(
        doc,
        "Feature engineering được chia thành bốn nhóm. Nhóm thứ nhất là historical demand features gồm lag_7, lag_14, lag_28, lag_56 và rolling mean 7/28/56 ngày. "
        "Nhóm thứ hai là calendar features gồm weekday, month, year, event flag, event type và SNAP theo state. Nhóm thứ ba là price features gồm sell_price và relative_price so với giá trung bình category-store-week. "
        "Nhóm thứ tư là hierarchy/identity features gồm item, department, category, store và state. Sự phân nhóm này cũng là cơ sở cho ablation study.",
    )
    add_p(
        doc,
        "Clustering features khác forecasting features ở mục đích sử dụng. Forecasting features được tạo theo từng ngày để dự báo y_i,t, còn clustering features được tính ở cấp chuỗi tại mỗi origin để mô tả hành vi lịch sử của chuỗi. "
        "Các clustering features gồm total sales, mean sales, median sales, standard deviation, zero-sales ratio, nonzero days, ADI, CV2, max gap, positive mean, spike frequency, event lift, weekend ratio và price statistics. "
        "Sau đó các biến lệch phải được log-transform, clipping theo quantile và RobustScaler để giảm ảnh hưởng outlier trước khi phân cụm.",
    )
    add_p(
        doc,
        "LightGBM được huấn luyện với objective Tweedie vì nhu cầu bán lẻ là biến không âm, có nhiều giá trị 0 và phần dương lệch phải. "
        "Tweedie objective phù hợp hơn squared error thuần túy trong bối cảnh zero-inflated continuous/count-like demand, trong khi vẫn giữ khả năng học phi tuyến mạnh của gradient boosting tree. "
        "Early stopping dùng 28 ngày ngay trước origin làm inner validation, còn test horizon là 28 ngày sau origin; do đó validation không chạm vào actual test horizon.",
    )

    add_h(doc, "5.2. Evaluation criteria", 2)
    add_p(doc, "Các tiêu chí đánh giá được chọn để giảm thiên lệch: WRMSSE phản ánh hierarchy; WAPE và bias hỗ trợ diễn giải sai số; scale-aware stability tránh phóng đại near-zero demand; train-test gap kiểm tra overfitting; DM/Nemenyi/CI kiểm tra ý nghĩa thống kê.")
    add_p(
        doc,
        "Scale-aware stability được thiết kế để đo mức thay đổi của forecast khi origin dịch chuyển, nhưng tránh chia trực tiếp cho mean sales quá nhỏ. "
        "Với các sản phẩm bán gần như bằng 0, thay đổi tuyệt đối rất nhỏ cũng có thể tạo tỷ lệ thay đổi rất lớn nếu mẫu số gần 0. "
        "Stability denominator được chặn bởi một floor để giảm hiện tượng phóng đại ở nhóm intermittent demand.",
    )
    add_p(
        doc,
        "Cụ thể, với hai rolling origins liên tiếp o_prev và o_cur, chỉ các cặp dự báo có cùng id và cùng ngày d trong phần horizon chồng lấn mới được so sánh. "
        "Mức thay đổi tuyệt đối của forecast được chuẩn hóa theo max(mean_sales_i, gamma), trong đó gamma là stability floor. "
        "Cách formalize này khác với coefficient of variation theo nhiều seed: nó đo độ ổn định khi cửa sổ huấn luyện dịch chuyển theo thời gian, phù hợp với quy trình demand planning định kỳ.",
    )
    add_p(
        doc,
        "Ngoài stability loss trung bình, JumpRate(tau) đo tỷ lệ forecast có thay đổi chuẩn hóa vượt ngưỡng tau. "
        "Chỉ số này giúp phát hiện trường hợp trung bình ổn định nhưng tồn tại một phần đáng kể chuỗi bị nhảy dự báo mạnh. "
        "Trong nghiên cứu, tau được đặt ở 0.3 và 0.5 để đánh giá mức thay đổi vừa và lớn theo thang mean-sales đã được chặn floor.",
    )
    add_p(
        doc,
        "Overfitting được đánh giá bằng khoảng cách giữa train, inner validation và test. Nếu train error giảm mạnh nhưng test error tăng bất thường, kết quả có thể phản ánh mô hình học thuộc lịch sử thay vì khái quát tốt. "
        "Ngoài ra, feature importance được lưu theo origin để kiểm tra mô hình có phụ thuộc quá mức vào một feature duy nhất hay không. "
        "Statistical testing gồm bootstrap confidence interval, Friedman/Nemenyi cho so sánh thứ hạng và Diebold-Mariano test cho cặp A0-C.",
    )
    add_eq(doc, "RMSSE_i = sqrt( mean_h((y_i,t - yhat_i,t)^2) / scale_i )")
    add_eq(doc, "WRMSSE = sum_i w_i * RMSSE_i")
    add_eq(doc, "S_m(o_prev,o_cur) = mean_{i,t in overlap} |yhat_{i,t,o_cur}^{(m)} - yhat_{i,t,o_prev}^{(m)}| / max(mean_sales_i, gamma)")
    add_eq(doc, "JumpRate_m(tau) = mean 1[ |Delta yhat_i,t| / max(mean_sales_i, gamma) > tau ]")
    add_eq(doc, "ADI = active_days / nonzero_days")
    add_eq(doc, "CV2 = Var(y | y>0) / Mean(y | y>0)^2")
    add_eq(doc, "WAPE = sum |y - yhat| / sum |y|")
    add_eq(doc, "Bias = sum(yhat - y) / sum(y)")
    add_eq(doc, "DM = mean(d_t) / sqrt(HACVar(d_t)/n), d_t = L_A0,t - L_C,t")

    add_h(doc, "6. Result & Discussion", 1)
    add_h(doc, "6.1. Kết quả tổng hợp", 2)
    result_rows = [
        [MODEL_LABEL[r.model], fmt(r.sampled_wrmsse_12level), fmt(r.validation_origin_wrmsse_mean), fmt(r.scale_aware_stability_loss), fmt(r.wape), fmt(r.bias), fmt(r.mean_test_train_gap)]
        for r in base.itertuples()
    ]
    add_table(doc, ["Mô hình", "Rolling WRMSSE", "Validation WRMSSE", "Stability loss", "WAPE", "Bias", "Train-test gap"], result_rows)
    add_p(doc, f"C đạt WRMSSE thấp nhất ({fmt(c.sampled_wrmsse_12level)}) và stability loss thấp nhất ({fmt(c.scale_aware_stability_loss)}). B1 có WRMSSE {fmt(b1.sampled_wrmsse_12level)}, cao hơn A0 {fmt(a0.sampled_wrmsse_12level)}, cho thấy cluster label đơn lẻ không đủ.")
    add_p(
        doc,
        "Sự khác biệt giữa B1 và C có ý nghĩa phương pháp luận. B1 chỉ cung cấp cluster label cho cùng một global learner, nên mô hình vẫn phải học một hàm dự báo chung cho toàn bộ panel. "
        "Nếu tác động của demand regime chủ yếu nằm ở tương tác phức tạp giữa lag, rolling, calendar, price và hierarchy, một biến cluster label đơn lẻ có thể không đủ để thay đổi cơ chế học. "
        "Trong khi đó, C cho phép từng cụm có cây quyết định, split point và feature interaction riêng, vì vậy có thể phù hợp hơn với dữ liệu dị thể.",
    )
    add_p(
        doc,
        "Lợi thế của C tập trung ở độ chính xác trung bình; phân tích DM và phân đoạn vẫn ghi nhận một phần chuỗi mà A0 có loss thấp hơn. "
        "Mẫu kết quả này phù hợp với bản chất dị thể của retail demand.",
    )
    add_p(
        doc,
        "Các kết quả được báo cáo trong cùng thiết kế rolling-origin local evaluation, cùng feature pipeline và cùng metric cho A0, B1 và C. "
        "Trọng tâm là hiệu ứng của cluster-aware learning trong điều kiện kiểm soát leakage.",
    )
    add_p(
        doc,
        "Audit sau thực nghiệm xác nhận lag/rolling features, clustering features theo origin, recursive forecasting, early stopping và WRMSSE đều tuân thủ thiết kế historical-only.",
    )
    add_fig(doc, FIG / "fig03_accuracy_stability_tradeoff.png", "Hình 3. Trade-off giữa WRMSSE và stability.")

    add_h(doc, "6.2. Kết quả theo rolling origin", 2)
    piv = per_origin.pivot(index="origin", columns="model", values="official_like_wrmsse_12level")
    origin_rows = [[str(idx), fmt(row["A0_global_baseline"]), fmt(row["B1_cluster_label"]), fmt(row["C_cluster_specific"])] for idx, row in piv.iterrows()]
    add_table(doc, ["Origin", "A0", "B1", "C"], origin_rows)
    add_p(
        doc,
        "Kết quả theo origin cho thấy lợi thế của C không đến từ một split đơn lẻ. "
        "Độ dao động giữa origins phản ánh mùa vụ, sự kiện và biến động category/store ở phần cuối chuỗi M5.",
    )
    add_p(
        doc,
        "Năm rolling origins cung cấp kiểm tra temporal robustness ở mức thực nghiệm; mở rộng số origin là hướng tự nhiên cho nghiên cứu tiếp theo.",
    )

    add_h(doc, "6.3. Kết quả theo cluster và demand regime", 2)
    cm = cluster_metrics.groupby(["model", "cluster_label"])[["rmsse_item_store", "mae", "wape", "bias"]].mean().reset_index()
    cluster_rows = [[MODEL_LABEL[r.model], str(int(r.cluster_label)), fmt(r.rmsse_item_store), fmt(r.mae), fmt(r.wape), fmt(r.bias)] for r in cm.itertuples()]
    add_table(doc, ["Mô hình", "Cluster", "RMSSE", "MAE", "WAPE", "Bias"], cluster_rows)
    add_p(
        doc,
        "Kết quả theo cluster cho thấy cách sai số thay đổi giữa các demand regimes. "
        "Phân tích này bổ sung cho WRMSSE tổng thể bằng cách tách high-demand, medium-demand và long-tail/intermittent behavior.",
    )
    dr = residual[residual["group_type"].eq("demand_class")]
    demand_result_rows = [[str(r.group), MODEL_LABEL[r.model], fmt(r.mae), fmt(r.bias), fmt(r.p90_abs_error)] for r in dr.itertuples()]
    add_table(doc, ["Demand regime", "Mô hình", "MAE", "Bias", "P90 AE"], demand_result_rows)
    add_p(
        doc,
        "Demand regime theo ADI-CV2 là lát cắt độc lập với cluster học được từ nhiều feature. "
        "Sự khác biệt MAE, bias và P90 absolute error giữa các regime làm rõ phần đóng góp của intermittent, lumpy, smooth và erratic demand vào sai số chung.",
    )
    add_fig(doc, FIG / "fig06_cluster_profile_k3.png", "Hình 4. Profile các cụm K=3.")

    add_h(doc, "6.4. Kết quả theo category, store và state", 2)
    cat = hierarchy[hierarchy["level"].eq("cat_id")]
    cat_rows = [[str(r.group), MODEL_LABEL[r.model], fmt(r.mae), fmt(r.wape), fmt(r.bias)] for r in cat.itertuples()]
    add_table(doc, ["Category", "Mô hình", "MAE", "WAPE", "Bias"], cat_rows)
    store = hierarchy[hierarchy["level"].eq("store_id")]
    store_rows = [[str(r.group), MODEL_LABEL[r.model], fmt(r.mae), fmt(r.wape), fmt(r.bias)] for r in store.itertuples()]
    add_table(doc, ["Store", "Mô hình", "MAE", "WAPE", "Bias"], store_rows)
    forecasts = pd.read_parquet(BASE / "metrics" / "test_forecasts.parquet")
    forecasts["state_id"] = forecasts["store_id"].str.split("_").str[0]
    state_rows = []
    for (state, model), g in forecasts.groupby(["state_id", "model"], observed=True):
        state_rows.append([state, MODEL_LABEL[model], fmt((g["y"] - g["yhat"]).abs().mean()), fmt((g["y"] - g["yhat"]).abs().sum() / (g["y"].abs().sum() + 1e-9)), fmt((g["yhat"] - g["y"]).sum() / (g["y"].sum() + 1e-9))])
    add_table(doc, ["State", "Mô hình", "MAE", "WAPE", "Bias"], state_rows)
    add_p(
        doc,
        "Category, store và state là các lát cắt vận hành chính của M5. "
        "Category phản ánh loại sản phẩm, store phản ánh bối cảnh bán hàng, còn state liên quan đến SNAP và khác biệt vùng.",
    )

    add_h(doc, "6.5. Ablation, kiểm định thống kê và robustness", 2)
    base_c_wrmsse = float(c.sampled_wrmsse_12level)
    ab_rows = []
    for r in ablation.itertuples():
        label = {"no_price": "No price", "no_calendar": "No calendar/SNAP/event", "no_hierarchy": "No hierarchy/ID"}[r.run]
        ab_rows.append([label, fmt(r.sampled_wrmsse_12level), fmt(r.validation_origin_wrmsse_mean), fmt(r.scale_aware_stability_loss), fmt(r.sampled_wrmsse_12level - base_c_wrmsse)])
    add_table(doc, ["Biến thể C", "Rolling WRMSSE", "Validation WRMSSE", "Stability loss", "Delta WRMSSE"], ab_rows)
    add_p(
        doc,
        "Ablation cho thấy calendar/SNAP/event là nhóm feature có tác động lớn nhất, phù hợp với nhịp bán lẻ theo weekday, sự kiện và SNAP. "
        "Hierarchy/ID cũng đóng góp rõ, trong khi price features tạo ảnh hưởng nhỏ hơn trong cấu hình hiện tại.",
    )
    ci_rows = [[r.metric, MODEL_LABEL[r.model], fmt(r.mean_diff), fmt(r["ci_2.5"]), fmt(r["ci_97.5"])] for _, r in bootstrap.iterrows() if r["model"] in MODEL_ORDER]
    add_table(doc, ["Metric", "Mô hình so với A0", "Mean diff", "CI 2.5%", "CI 97.5%"], ci_rows)
    add_p(
        doc,
        "Confidence interval tóm tắt độ không chắc chắn của chênh lệch metric qua origins. "
        "Các kiểm định được đọc cùng DM-HAC và block/bootstrap diagnostics để phản ánh phụ thuộc thời gian trong forecast errors.",
    )
    dm_rows = [[r.loss, str(int(r.n_series)), fmt(r.mean_loss_diff_a0_minus_c), fmt(r.share_c_better_mean_loss), fmt(r.share_c_better_p05), fmt(r.share_c_better_fdr05)] for r in dm_summary.itertuples()]
    add_table(doc, ["Loss", "Số chuỗi", "Mean diff A0-C", "Share C better", "Share p<0.05", "Share FDR<0.05"], dm_rows)
    add_p(
        doc,
        "DM test dùng HAC/Newey-West correction cho horizon 28 ngày và Benjamini-Hochberg FDR correction cho nhiều chuỗi. "
        "Kết quả ủng hộ H1 ở mức tổng thể, với effect size tập trung ở cải thiện trung bình hơn là cải thiện đồng nhất trên toàn bộ panel.",
    )
    nem_rows = [[str(r["metric"]), MODEL_LABEL[str(r["model_a"])], MODEL_LABEL[str(r["model_b"])], fmt(r["rank_diff"]), fmt(r["critical_difference"]), str(bool(r["significant_alpha_0.05"]))] for _, r in nemenyi.iterrows()]
    add_table(doc, ["Metric", "Model A", "Model B", "Rank diff", "Critical diff", "Significant"], nem_rows)
    add_p(
        doc,
        "Nemenyi post-hoc bổ sung góc nhìn thứ hạng sau Friedman test. "
        "Với số origin nhỏ, phép thử này chủ yếu được dùng như kiểm tra bảo thủ bên cạnh effect size.",
    )
    kmed_rows = [[str(int(r.sample_n)), fmt(r.ari_kmeans_vs_kmedoids), str(int(r.kmedoids_min_cluster_size)), str(int(r.kmedoids_max_cluster_size)), str(r.note)] for r in kmed_sens.itertuples()]
    add_table(doc, ["Sample", "ARI", "Min cluster size", "Max cluster size", "Ghi chú"], kmed_rows)
    add_p(
        doc,
        "K-Medoids robustness full-scale cho ARI ở mức thấp-vừa, phản ánh độ nhạy của phân hoạch cụm khi đổi thuật toán. "
        "K=3 được sử dụng như phân hoạch vận hành theo demand regime.",
    )

    add_h(doc, "6.6. Multi-seed robustness", 2)
    seed_agg_focus = seed_agg[seed_agg["model"].isin(MODEL_ORDER)].copy()
    seed_agg_focus["model"] = pd.Categorical(seed_agg_focus["model"], MODEL_ORDER, ordered=True)
    seed_agg_focus = seed_agg_focus.sort_values("model")
    seed_rows = [
        [
            MODEL_LABEL[r.model],
            fmt(r.rmsse_item_store_mean),
            fmt(r.rmsse_item_store_std),
            fmt(r.sampled_wrmsse_12level_mean),
            fmt(r.sampled_wrmsse_12level_std),
            fmt(r.sampled_wrmsse_12level_ci95_low),
            fmt(r.sampled_wrmsse_12level_ci95_high),
            fmt(r.scale_aware_stability_loss_mean),
            fmt(r.scale_aware_stability_loss_std),
            fmt(r.train_test_gap_mean),
        ]
        for r in seed_agg_focus.itertuples()
    ]
    add_table(doc, ["Mô hình", "RMSSE mean", "RMSSE std", "WRMSSE mean", "WRMSSE std", "WRMSSE CI low", "WRMSSE CI high", "Stability mean", "Stability std", "Gap mean"], seed_rows)
    per_seed_rows = []
    for _, r in seed_per_seed[seed_per_seed["model"].isin(MODEL_ORDER)].iterrows():
        per_seed_rows.append(
            [
                str(int(r["seed"])),
                MODEL_LABEL[str(r["model"])],
                fmt(r["sampled_wrmsse_12level"]),
                fmt(r["scale_aware_stability_loss"]),
                fmt(r["jump_rate_tau_0.3"]),
                fmt(r["jump_rate_tau_0.5"]),
                fmt(r["wape"]),
                fmt(r["bias"]),
                fmt(r["train_test_gap"]),
            ]
        )
    add_table(doc, ["Seed", "Mô hình", "WRMSSE", "Stability", "Jump@0.3", "Jump@0.5", "WAPE", "Bias", "Gap"], per_seed_rows)
    paired_focus = seed_paired[seed_paired["metric"].isin(["sampled_wrmsse_12level", "scale_aware_stability_loss", "wape", "train_test_gap"])].copy()
    paired_rows = [
        [
            str(r.metric),
            str(int(r.n_seeds)),
            fmt(r.mean_diff_A0_minus_C),
            fmt(r.paired_t_p_value),
            fmt(r.wilcoxon_p_value),
        ]
        for r in paired_focus.itertuples()
    ]
    add_table(doc, ["Metric", "Số seed", "Mean diff A0-C", "Paired t p", "Wilcoxon p"], paired_rows)
    nemenyi_focus = seed_nemenyi[
        (seed_nemenyi["metric"].isin(["sampled_wrmsse_12level", "scale_aware_stability_loss", "wape", "train_test_gap"]))
        & (seed_nemenyi["model_a"].eq("A0_global_baseline"))
        & (seed_nemenyi["model_b"].eq("C_cluster_specific"))
    ].copy()
    nemenyi_rows = []
    for _, r in nemenyi_focus.iterrows():
        nemenyi_rows.append(
            [
                str(r["metric"]),
                fmt(r["average_rank_a"]),
                fmt(r["average_rank_b"]),
                fmt(r["rank_diff"]),
                fmt(r["critical_difference"]),
                str(bool(r["significant_alpha_0.05"])),
            ]
        )
    add_table(doc, ["Metric", "Rank A0", "Rank C", "Rank diff", "Critical diff", "Significant"], nemenyi_rows)
    add_p(
        doc,
        "Multi-seed robustness được chạy full-scale với K=3 trên năm seed 42, 52, 62, 72 và 82, cùng năm rolling origins 1885, 1892, 1899, 1906 và 1913. "
        "Mỗi seed dùng historical-only clustering, Mini-batch K-Means và LightGBM theo seed tương ứng, cùng recursive 28-day forecasting. "
        "C có WRMSSE và WAPE trung bình thấp nhất; paired t-test A0-C có ý nghĩa cho WRMSSE và WAPE. Stability improvement nhỏ và không có ý nghĩa thống kê rõ do seed82 làm C có stability loss cao hơn. "
        "Train-test gap của C cao hơn A0, phản ánh trade-off giữa accuracy và complexity.",
    )

    add_h(doc, "6.7. Trả lời câu hỏi nghiên cứu", 2)
    answers = [
        f"RQ1: Có. C có WRMSSE {fmt(c.sampled_wrmsse_12level)}, thấp hơn A0 {fmt(a0.sampled_wrmsse_12level)}.",
        f"RQ2: Có. C có stability loss {fmt(c.scale_aware_stability_loss)}, thấp hơn A0 {fmt(a0.scale_aware_stability_loss)}.",
        f"RQ3: Không. B1 có WRMSSE {fmt(b1.sampled_wrmsse_12level)}, cao hơn A0.",
        "RQ4: Có. No-calendar/SNAP/event làm WRMSSE tăng mạnh nhất; no-hierarchy tăng rõ; no-price tăng nhẹ.",
        f"RQ5: Ổn định ở mức thấp-vừa. K-Medoids full-scale có ARI {fmt(kmed.iloc[0].ari_kmeans_vs_kmedoids)}.",
    ]
    for ans in answers:
        doc.add_paragraph(ans, style="List Bullet")
    add_p(doc, "Knowledge rút ra là cluster-aware forecasting chỉ có giá trị rõ khi cụm được dùng để tách cơ chế học như C; việc gắn cluster label vào global learner như B1 chưa đủ. Đồng thời, stability là tiêu chí cần thiết vì intermittent demand làm các chỉ số phần trăm dễ sai lệch.")

    add_h(doc, "7. Conclusion", 1)
    add_p(
        doc,
        "Nghiên cứu kiểm tra một framework cluster-aware global forecasting trên M5 bằng Mini-batch K-Means K=3, LightGBM Tweedie và rolling-origin evaluation. "
        "C cải thiện WRMSSE so với A0 và B1, trong khi B1 không tạo lợi ích ổn định so với baseline. Full multi-seed củng cố kết luận về accuracy; stability cho thấy lợi thế nhỏ hơn và nhạy với seed. "
        "Kết quả cũng ghi nhận trade-off của C qua train-test gap cao hơn. Hướng phát triển tiếp theo gồm mở rộng rolling origins, thử clustering có khả năng mở rộng khác, đánh giá ngoài M5 và nghiên cứu hierarchical reconciliation.",
    )

    add_h(doc, "8. References", 1)
    references = [
        "Bandara, K., Bergmeir, C., & Smyl, S. (2020). Forecasting across time series databases using recurrent neural networks on groups of similar series: A clustering approach. Expert Systems with Applications, 140, 112896.",
        "Croston, J. D. (1972). Forecasting and stock control for intermittent demands. Operational Research Quarterly, 23(3), 289-303.",
        "E, E., Yu, M., Tian, X., & Tao, Y. (2022). Dynamic model selection based on demand pattern classification in retail sales forecasting. Mathematics, 10(17), 3179.",
        "Fildes, R., Ma, S., & Kolassa, S. (2019). Retail forecasting: Research and practice. International Journal of Forecasting, 35(1), 1-7.",
        "Hoeltgebaum, H. H., Borenstein, D., Fernandes, C., & Veiga, A. (2023). Lumpy and intermittent retail demand forecasts with score-driven models. European Journal of Operational Research.",
        "Hyndman, R. J., & Koehler, A. B. (2006). Another look at measures of forecast accuracy. International Journal of Forecasting, 22(4), 679-688.",
        "Ke, G., Meng, Q., Finley, T., Wang, T., Chen, W., Ma, W., Ye, Q., & Liu, T. Y. (2017). LightGBM: A highly efficient gradient boosting decision tree. Advances in Neural Information Processing Systems.",
        "Klee, S., & Xia, A. (2025). Measuring time series forecast stability for demand planning. Amazon Science.",
        "Ling, J., & Wu, W. B. (2025). Comparative analysis of global and local probabilistic time series forecasting for contiguous spatial demand regions. arXiv:2509.08214.",
        "Ma, S., & Fildes, R. (2022). The performance of the global bottom-up approach in the M5 accuracy competition: A robustness check. International Journal of Forecasting, 38(4), 1492-1499.",
        "Makridakis, S., Spiliotis, E., & Assimakopoulos, V. (2022). M5 accuracy competition: Results, findings, and conclusions. International Journal of Forecasting, 38(4), 1346-1364.",
        "Mitchell, R., Monokroussos, G., Nikzad, A., & Wang, W. (2024). Hierarchical demand forecasting in retail: A view from the trenches. SSRN working paper.",
        "Montero-Manso, P., Athanasopoulos, G., Hyndman, R. J., & Talagala, T. S. (2020). FFORMA: Feature-based forecast model averaging. International Journal of Forecasting, 36(1), 86-92.",
        "Semenoglou, A. A., Spiliotis, E., Makridakis, S., & Assimakopoulos, V. (2021). Investigating the accuracy of cross-learning time series forecasting methods. International Journal of Forecasting.",
        "Syntetos, A. A., Boylan, J. E., & Croston, J. D. (2005). On the categorization of demand patterns. Journal of the Operational Research Society, 56, 495-503.",
        "Talagala, T. S., Hyndman, R. J., & Athanasopoulos, G. (2021). Meta-learning how to forecast time series. Journal of Forecasting.",
        "Tibshirani, R., Walther, G., & Hastie, T. (2001). Estimating the number of clusters in a data set via the gap statistic. Journal of the Royal Statistical Society: Series B, 63(2), 411-423.",
        "Van Ruitenbeek, R. E., Koole, G. M., & Bhulai, S. (2023). A hierarchical agglomerative clustering for product sales forecasting. Decision Analytics Journal, 8, 100318.",
        "Vermorel, J. (2013). Quantile forecasting for retail inventory optimization. International Journal of Forecasting, 29(4), 595-604.",
    ]
    for ref in references:
        doc.add_paragraph(ref, style="List Number")

    doc.save(OUT_DOCX)
    OUT_MD.write_text(
        "# Bài báo nghiên cứu M5 theo cấu trúc chuẩn học thuật\n\n"
        f"File Word: `{OUT_DOCX}`\n\n"
        "Cấu trúc: Abstract, Introduction, Related work, Background, Research methodology, Experiment, Result & Discussion, Conclusion, References.\n",
        encoding="utf-8",
    )
    print(OUT_DOCX)
    print(OUT_MD)
    print("abstract_words", len(abstract.split()))


if __name__ == "__main__":
    main()
