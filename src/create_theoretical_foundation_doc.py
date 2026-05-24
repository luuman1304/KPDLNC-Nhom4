from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX_OUT = ROOT / "Document" / "Co_So_Ly_Thuyet_Nghien_Cuu_M5.docx"
MD_OUT = ROOT / "Document" / "Co_So_Ly_Thuyet_Nghien_Cuu_M5.md"


TITLE = (
    "CƠ SỞ LÝ THUYẾT CHO NGHIÊN CỨU CLUSTER-AWARE GLOBAL FORECASTING "
    "TRÊN DỮ LIỆU M5"
)


SECTIONS = [
    (
        "1. Bối cảnh nghiệp vụ của dự báo nhu cầu bán lẻ",
        [
            (
                "Dự báo nhu cầu bán lẻ là đầu vào quan trọng cho quản trị tồn kho, "
                "bổ sung hàng, phân bổ hàng hóa và lập kế hoạch cung ứng. Trong môi "
                "trường bán lẻ quy mô lớn, sai số dự báo không chỉ làm tăng tồn kho "
                "dư hoặc thiếu hàng, mà còn ảnh hưởng đến quyết định phân bổ nguồn "
                "lực ở cấp cửa hàng, vùng và ngành hàng. Vì vậy, một mô hình dự báo "
                "cần được đánh giá đồng thời theo độ chính xác, độ ổn định và khả "
                "năng khái quát hóa, thay vì chỉ tối ưu một chỉ số sai số trung bình "
                "(Fildes et al., 2019; Vermorel, 2013)."
            ),
            (
                "Trong nghiên cứu này, logic nghiệp vụ được chuyển hóa thành ba yêu "
                "cầu kỹ thuật. Thứ nhất, mô hình phải xử lý được dữ liệu nhiều chuỗi "
                "item-store. Thứ hai, mô hình phải tôn trọng cấu trúc phân cấp của M5. "
                "Thứ ba, dự báo phải đủ ổn định khi rolling origin thay đổi, vì trong "
                "thực tế kế hoạch bán lẻ thường được cập nhật định kỳ."
            ),
        ],
    ),
    (
        "2. Dữ liệu M5 và cấu trúc phân cấp của bài toán",
        [
            (
                "Bộ dữ liệu M5 Forecasting Accuracy được xây dựng từ dữ liệu bán hàng "
                "Walmart, bao gồm chuỗi bán hàng theo ngày, lịch, sự kiện, SNAP, giá "
                "bán và các định danh phân cấp như item, department, category, store "
                "và state. Đặc trưng quan trọng của M5 là cấu trúc phân cấp: dự báo ở "
                "cấp item-store có thể được tổng hợp lên nhiều cấp cao hơn như store, "
                "state, category và department (Makridakis et al., 2022)."
            ),
            (
                "Cấu trúc này làm cho bài toán khác với dự báo một chuỗi đơn lẻ. Một "
                "mô hình có thể dự báo tốt ở cấp item-store nhưng vẫn tạo sai số lớn "
                "ở cấp tổng hợp nếu sai số có cùng hướng trong nhiều chuỗi con. Vì vậy, "
                "nghiên cứu sử dụng WRMSSE để đánh giá mô hình theo nhiều cấp phân cấp, "
                "phù hợp với mục tiêu vận hành của bán lẻ quy mô lớn."
            ),
        ],
    ),
    (
        "3. WRMSSE và đánh giá dự báo phân cấp",
        [
            (
                "WRMSSE, Weighted Root Mean Squared Scaled Error, là chỉ số đánh giá "
                "chính trong M5. Chỉ số này chuẩn hóa sai số bằng scale lịch sử của "
                "chuỗi và gán trọng số theo mức độ quan trọng kinh tế của từng chuỗi. "
                "Về nguyên tắc, RMSSE giúp so sánh sai số giữa các chuỗi có quy mô "
                "bán hàng khác nhau, còn trọng số giúp phản ánh đóng góp doanh thu "
                "của từng chuỗi hoặc cấp phân cấp (Makridakis et al., 2022; Hyndman "
                "& Koehler, 2006)."
            ),
            (
                "Trong nghiên cứu này, WRMSSE được dùng để đánh giá A0, B1 và C trong "
                "cùng một thiết kế rolling-origin. Việc sử dụng metric phân cấp giúp "
                "tránh kết luận lệch nếu mô hình chỉ cải thiện ở cấp bottom-level nhưng "
                "không cải thiện ở các cấp tổng hợp có ý nghĩa vận hành."
            ),
            "Công thức khái quát: RMSSE_i = sqrt(mean((y_i - yhat_i)^2) / scale_i); WRMSSE = sum_i w_i * RMSSE_i.",
        ],
    ),
    (
        "4. Intermittent demand, ADI và CV²",
        [
            (
                "Intermittent demand là dạng nhu cầu có nhiều kỳ bằng 0 xen kẽ với các "
                "kỳ phát sinh bán hàng. Đây là hiện tượng phổ biến trong bán lẻ, đặc "
                "biệt với sản phẩm long-tail hoặc sản phẩm bán không đều. Các phương "
                "pháp đánh giá dựa trên tỷ lệ phần trăm dễ bị méo trong bối cảnh này, "
                "vì mẫu số có thể rất nhỏ hoặc bằng 0 (Croston, 1972; Syntetos et al., "
                "2005)."
            ),
            (
                "ADI, Average Demand Interval, đo khoảng cách trung bình giữa các lần "
                "bán khác 0. CV² đo mức biến động tương đối của lượng bán trong các "
                "ngày có nhu cầu. Hai chỉ số này thường được dùng để phân biệt smooth, "
                "intermittent, erratic và lumpy demand. Trong bài nghiên cứu, ADI và "
                "CV² không được dùng làm nhãn mục tiêu, mà được dùng để mô tả dữ liệu, "
                "hỗ trợ diễn giải cụm và kiểm tra xem clustering có phản ánh khác biệt "
                "về demand behavior hay không."
            ),
            "Công thức sử dụng: ADI = số ngày quan sát / số ngày có bán; CV² = Var(y | y > 0) / Mean(y | y > 0)^2.",
        ],
    ),
    (
        "5. Global forecasting và cross-learning",
        [
            (
                "Global forecasting huấn luyện một mô hình chung trên nhiều chuỗi thời "
                "gian. Cách tiếp cận này cho phép mô hình học các pattern lặp lại giữa "
                "các chuỗi, đặc biệt hữu ích khi mỗi chuỗi đơn lẻ có lịch sử ngắn hoặc "
                "nhiễu lớn. Nhiều nghiên cứu cho thấy global models có thể khai thác "
                "cross-learning hiệu quả trong các tập dữ liệu nhiều chuỗi (Bandara et "
                "al., 2020; Montero-Manso et al., 2020; Semenoglou et al., 2021)."
            ),
            (
                "Hạn chế của global forecasting là mọi chuỗi phải đi qua cùng một hàm "
                "học. Nếu dữ liệu chứa nhiều demand regimes khác nhau, ví dụ sản phẩm "
                "bán đều, bán thưa và bán biến động mạnh, một mô hình chung có thể bị "
                "trung bình hóa hành vi. Đây là cơ sở để nghiên cứu so sánh A0, một "
                "global LightGBM chuẩn, với B1 và C, hai biến thể đưa thông tin cụm vào "
                "quá trình học."
            ),
        ],
    ),
    (
        "6. Cluster-aware forecasting",
        [
            (
                "Cluster-aware forecasting sử dụng phân cụm để nhóm các chuỗi có hành "
                "vi nhu cầu tương tự. Ý tưởng chính là giảm dị thể trong dữ liệu trước "
                "khi huấn luyện hoặc giúp mô hình nhận biết demand regime của từng chuỗi. "
                "Các nghiên cứu về forecasting theo nhóm chuỗi cho thấy clustering có "
                "thể hỗ trợ lựa chọn mô hình, cải thiện diễn giải và khai thác cấu trúc "
                "liên chuỗi (Bandara et al., 2020; Van Ruitenbeek et al., 2023)."
            ),
            (
                "Trong bài nghiên cứu, cluster-aware forecasting được kiểm tra bằng hai "
                "cách. B1 thêm cluster label vào global LightGBM như một categorical "
                "feature. C dùng cluster label để tách dữ liệu và huấn luyện một LightGBM "
                "riêng cho từng cụm. So sánh A0-B1-C giúp xác định liệu cluster label "
                "chỉ như một biến bổ sung đã đủ hay cần thay đổi cơ chế học theo từng "
                "cụm."
            ),
        ],
    ),
    (
        "7. Mini-batch K-Means, K-Medoids và lựa chọn K=3",
        [
            (
                "K-Means là thuật toán phân cụm dựa trên centroid, tối thiểu hóa tổng "
                "bình phương khoảng cách từ điểm dữ liệu đến tâm cụm. Mini-batch K-Means "
                "là biến thể dùng các mini-batch để tăng khả năng mở rộng trên dữ liệu "
                "lớn. Trong nghiên cứu, Mini-batch K-Means phù hợp vì số chuỗi item-store "
                "lớn và clustering phải lặp lại theo từng rolling origin và nhiều seed "
                "(Ke et al., 2017; Tibshirani et al., 2001)."
            ),
            (
                "K-Medoids được dùng như kiểm tra robustness vì medoid là điểm dữ liệu "
                "thật, thường ít nhạy hơn centroid trong một số trường hợp có ngoại lệ. "
                "Nghiên cứu không dùng K-Medoids làm pipeline chính do chi phí full-scale "
                "cao hơn, nhưng dùng kết quả ARI giữa K-Means và K-Medoids để đánh giá "
                "độ ổn định của cấu trúc cụm."
            ),
            (
                "K=3 được chọn theo hai lý do. Về trực quan, bản đồ ADI-CV² cho thấy ba "
                "vùng vận hành chính: long-tail/intermittent, medium demand và high-demand/core. "
                "Về thực nghiệm, K=3 giúp tránh tạo cụm quá nhỏ khi huấn luyện model C, "
                "đồng thời vẫn đủ tách biệt để kiểm tra cluster-specific learning."
            ),
        ],
    ),
    (
        "8. Tiền xử lý feature cho clustering",
        [
            (
                "Các feature dùng cho clustering như mean sales, total sales, ADI, CV², "
                "zero-sales ratio, gap, spike frequency, event lift và price statistics "
                "có thang đo khác nhau và thường lệch phải. Nếu đưa trực tiếp vào K-Means, "
                "các biến có scale lớn có thể chi phối khoảng cách Euclidean. Vì vậy, "
                "pipeline sử dụng log transform, clipping ngoại lệ và RobustScaler trước "
                "khi phân cụm."
            ),
            (
                "RobustScaler chuẩn hóa dữ liệu bằng median và interquartile range, nên "
                "ít nhạy hơn StandardScaler khi dữ liệu có outlier. Trong bài nghiên cứu, "
                "scaler chỉ được fit trên dữ liệu train tại từng rolling origin, tránh "
                "để phân phối của test horizon rò rỉ vào quá trình phân cụm."
            ),
        ],
    ),
    (
        "9. LightGBM và objective Tweedie",
        [
            (
                "LightGBM là thuật toán gradient boosting decision tree được thiết kế "
                "cho dữ liệu lớn, hỗ trợ học phi tuyến, feature interaction và biến "
                "phân loại hiệu quả. Đây là lựa chọn phù hợp cho bài toán M5 vì dữ liệu "
                "có nhiều feature dạng tabular như lag, rolling statistics, calendar, "
                "price và hierarchy (Ke et al., 2017; Makridakis et al., 2022)."
            ),
            (
                "Objective Tweedie được dùng vì nhu cầu bán lẻ là biến không âm, có nhiều "
                "giá trị 0 và phần dương lệch phải. Phân phối Tweedie thường được dùng "
                "cho dữ liệu có khối lượng xác suất tại 0 và giá trị dương liên tục hoặc "
                "count-like. Trong nghiên cứu, Tweedie giúp mô hình phù hợp hơn với đặc "
                "điểm zero-inflated demand so với squared error thuần túy."
            ),
        ],
    ),
    (
        "10. Feature engineering cho forecasting",
        [
            (
                "Forecasting features được chia thành bốn nhóm. Nhóm historical demand "
                "features gồm lag_7, lag_14, lag_28, lag_56 và rolling mean 7/28/56 ngày. "
                "Nhóm calendar features gồm weekday, month, year, event và SNAP. Nhóm "
                "price features gồm sell_price và relative price. Nhóm hierarchy/identity "
                "features gồm item, department, category, store và state."
            ),
            (
                "Nguyên tắc quan trọng là mọi lag và rolling feature phải được tạo từ dữ "
                "liệu quá khứ. Nếu rolling mean được tính trực tiếp trên cửa sổ bao gồm "
                "ngày dự báo, kết quả sẽ bị data leakage và sai số có thể đẹp giả tạo. "
                "Vì vậy, pipeline sử dụng thiết kế historical-only tại từng rolling origin."
            ),
        ],
    ),
    (
        "11. Rolling-origin evaluation và recursive forecasting",
        [
            (
                "Rolling-origin evaluation đánh giá mô hình qua nhiều điểm cắt thời gian. "
                "Tại mỗi origin T, mô hình chỉ dùng dữ liệu đến T để tạo feature, phân "
                "cụm, fit scaler và huấn luyện. Horizon sau T được giữ làm test. Cách "
                "thiết kế này phản ánh tốt hơn tình huống triển khai thực tế so với một "
                "split cố định duy nhất (Hyndman & Koehler, 2006)."
            ),
            (
                "Recursive forecasting nghĩa là dự báo của bước h+1 được đưa vào history "
                "để tạo lag và rolling feature cho bước h+2. Thiết kế này tránh teacher "
                "forcing, tức không dùng actual value trong test horizon để tạo feature "
                "cho các bước dự báo sau. Đây là điều kiện bắt buộc để đánh giá mô hình "
                "forecasting nhiều bước một cách hợp lệ."
            ),
        ],
    ),
    (
        "12. Scale-aware stability và JumpRate",
        [
            (
                "Forecast stability đo mức thay đổi của dự báo khi rolling origin dịch "
                "chuyển. Trong demand planning, dự báo quá dao động có thể gây bất ổn "
                "cho kế hoạch bổ sung hàng dù sai số trung bình thấp. Các nghiên cứu gần "
                "đây về forecast stability nhấn mạnh rằng độ ổn định là một tiêu chí bổ "
                "sung quan trọng bên cạnh accuracy (Klee & Xia, 2025)."
            ),
            (
                "Trong bài nghiên cứu, stability được chuẩn hóa theo scale của từng chuỗi "
                "nhưng có floor ở mẫu số để tránh phóng đại các chuỗi gần như không bán. "
                "JumpRate@0.3 và JumpRate@0.5 đo tỷ lệ dự báo có thay đổi vượt ngưỡng, "
                "giúp phát hiện trường hợp stability trung bình thấp nhưng vẫn có một "
                "phần chuỗi dao động mạnh."
            ),
        ],
    ),
    (
        "13. Overfitting diagnostics và feature importance",
        [
            (
                "Overfitting xảy ra khi mô hình học quá sát dữ liệu huấn luyện nhưng "
                "khái quát kém trên test horizon. Trong nghiên cứu, overfitting được "
                "kiểm tra bằng train-test gap, inner validation và so sánh kết quả qua "
                "nhiều rolling origins. Nếu model C cải thiện WRMSSE nhưng train-test "
                "gap tăng quá mạnh, kết luận cần được diễn giải cùng trade-off này."
            ),
            (
                "Feature importance được lưu theo origin và seed để kiểm tra mô hình "
                "có phụ thuộc quá mức vào một nhóm feature duy nhất hay không. Diễn giải "
                "feature importance chỉ được dùng như bằng chứng hỗ trợ, không thay thế "
                "cho đánh giá bằng metric và kiểm định thống kê."
            ),
        ],
    ),
    (
        "14. Ablation study",
        [
            (
                "Ablation study kiểm tra vai trò của từng nhóm feature bằng cách loại bỏ "
                "một nhóm và đo mức thay đổi metric. Trong bài nghiên cứu, ablation được "
                "thực hiện cho price, calendar/SNAP/event và hierarchy/ID. Mục tiêu không "
                "phải tìm feature tốt nhất riêng lẻ, mà là xác định nhóm tín hiệu nào làm "
                "mô hình suy giảm nhiều nhất khi bị loại bỏ."
            ),
            (
                "Kết quả ablation giúp trả lời câu hỏi liệu cải thiện của model C đến từ "
                "cluster-specific learning hay phụ thuộc quá mức vào một nhóm feature. "
                "Nó cũng hỗ trợ diễn giải business: nếu calendar/SNAP/event gây suy giảm "
                "lớn khi bỏ đi, điều này phù hợp với hành vi bán lẻ chịu tác động của "
                "ngày trong tuần, sự kiện và chương trình SNAP."
            ),
        ],
    ),
    (
        "15. Kiểm định thống kê: Diebold-Mariano, Friedman và Nemenyi",
        [
            (
                "Diebold-Mariano test được dùng để so sánh độ chính xác dự báo giữa hai "
                "mô hình trên cùng chuỗi sai số. Trong bài nghiên cứu, DM test tập trung "
                "vào cặp A0-C để kiểm tra liệu cải thiện của C có ý nghĩa thống kê hay "
                "không. Vì forecast errors có thể phụ thuộc theo thời gian, kiểm định cần "
                "được đọc cùng HAC/Newey-West correction và horizon correction (Diebold "
                "& Mariano, 1995)."
            ),
            (
                "Friedman test và Nemenyi post-hoc được dùng khi so sánh thứ hạng của "
                "nhiều mô hình trên nhiều folds hoặc seeds. Friedman test kiểm tra liệu "
                "có khác biệt tổng thể về thứ hạng mô hình hay không; Nemenyi test kiểm "
                "tra cặp mô hình nào khác biệt đáng kể sau đó. Cách tiếp cận này phù hợp "
                "khi nghiên cứu so sánh A0, B1 và C qua nhiều origin/seed (Demšar, 2006)."
            ),
            (
                "Vì bài nghiên cứu có nhiều metric và nhiều phép so sánh, kết quả thống "
                "kê cần được đọc cùng hiệu ứng thực tế, không chỉ p-value. Một khác biệt "
                "có ý nghĩa thống kê nhưng biên độ nhỏ vẫn cần được cân nhắc theo giá trị "
                "vận hành."
            ),
        ],
    ),
    (
        "16. Multi-seed robustness",
        [
            (
                "Multi-seed robustness kiểm tra xem kết quả có phụ thuộc vào một random "
                "seed duy nhất hay không. Trong pipeline này, seed ảnh hưởng đến Mini-batch "
                "K-Means, LightGBM và một số bước sampling/nội bộ của thuật toán. Nếu C "
                "chỉ thắng ở seed=42 nhưng không ổn định ở seed khác, kết luận về cải "
                "thiện accuracy sẽ yếu hơn."
            ),
            (
                "Nghiên cứu vì vậy chạy các seed 42, 52, 62, 72 và 82 trên cùng năm rolling "
                "origins. Kết quả multi-seed được tổng hợp bằng mean, standard deviation, "
                "min/max, confidence interval, paired tests và ranking tests. Đây là bước "
                "quan trọng để tăng độ tin cậy của kết luận thực nghiệm."
            ),
        ],
    ),
]


REFERENCES = [
    "Bandara, K., Bergmeir, C., & Smyl, S. (2020). Forecasting across time series databases using recurrent neural networks on groups of similar series: A clustering approach. Expert Systems with Applications, 140, 112896.",
    "Croston, J. D. (1972). Forecasting and stock control for intermittent demands. Operational Research Quarterly, 23(3), 289-303.",
    "Demšar, J. (2006). Statistical comparisons of classifiers over multiple data sets. Journal of Machine Learning Research, 7, 1-30.",
    "Diebold, F. X., & Mariano, R. S. (1995). Comparing predictive accuracy. Journal of Business & Economic Statistics, 13(3), 253-263.",
    "Fildes, R., Ma, S., & Kolassa, S. (2019). Retail forecasting: Research and practice. International Journal of Forecasting, 35(1), 1-7.",
    "Hyndman, R. J., & Koehler, A. B. (2006). Another look at measures of forecast accuracy. International Journal of Forecasting, 22(4), 679-688.",
    "Ke, G., Meng, Q., Finley, T., Wang, T., Chen, W., Ma, W., Ye, Q., & Liu, T. Y. (2017). LightGBM: A highly efficient gradient boosting decision tree. Advances in Neural Information Processing Systems.",
    "Klee, S., & Xia, A. (2025). Measuring time series forecast stability for demand planning. Amazon Science.",
    "Ma, S., & Fildes, R. (2022). The performance of the global bottom-up approach in the M5 accuracy competition: A robustness check. International Journal of Forecasting, 38(4), 1492-1499.",
    "Makridakis, S., Spiliotis, E., & Assimakopoulos, V. (2022). M5 accuracy competition: Results, findings, and conclusions. International Journal of Forecasting, 38(4), 1346-1364.",
    "Montero-Manso, P., Athanasopoulos, G., Hyndman, R. J., & Talagala, T. S. (2020). FFORMA: Feature-based forecast model averaging. International Journal of Forecasting, 36(1), 86-92.",
    "Semenoglou, A. A., Spiliotis, E., Makridakis, S., & Assimakopoulos, V. (2021). Investigating the accuracy of cross-learning time series forecasting methods. International Journal of Forecasting.",
    "Syntetos, A. A., Boylan, J. E., & Croston, J. D. (2005). On the categorization of demand patterns. Journal of the Operational Research Society, 56, 495-503.",
    "Tibshirani, R., Walther, G., & Hastie, T. (2001). Estimating the number of clusters in a data set via the gap statistic. Journal of the Royal Statistical Society: Series B, 63(2), 411-423.",
    "Van Ruitenbeek, R. E., Koole, G. M., & Bhulai, S. (2023). A hierarchical agglomerative clustering for product sales forecasting. Decision Analytics Journal, 8, 100318.",
    "Vermorel, J. (2013). Quantile forecasting for retail inventory optimization. International Journal of Forecasting, 29(4), 595-604.",
]


def add_paragraphs(doc: Document, paragraphs: list[str]) -> None:
    for text in paragraphs:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.3


def create_docx() -> None:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(13)
    styles["Heading 1"].font.name = "Times New Roman"
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.name = "Times New Roman"
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.bold = True

    title = doc.add_heading(TITLE, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro = doc.add_paragraph(
        "File này trình bày riêng phần cơ sở lý thuyết cho bài nghiên cứu. Nội dung "
        "được tổ chức theo từng nhóm kiến thức, thuật toán và ứng dụng thực tế đã "
        "được sử dụng trong pipeline thực nghiệm."
    )
    intro.paragraph_format.line_spacing = 1.3

    for heading, paragraphs in SECTIONS:
        doc.add_heading(heading, level=2)
        add_paragraphs(doc, paragraphs)

    doc.add_heading("Tài liệu tham khảo", level=2)
    for ref in REFERENCES:
        doc.add_paragraph(ref, style="List Number")

    doc.save(DOCX_OUT)


def create_markdown() -> None:
    lines = [f"# {TITLE}", ""]
    lines.append(
        "File này trình bày riêng phần cơ sở lý thuyết cho bài nghiên cứu. Nội dung "
        "được tổ chức theo từng nhóm kiến thức, thuật toán và ứng dụng thực tế đã "
        "được sử dụng trong pipeline thực nghiệm."
    )
    lines.append("")

    for heading, paragraphs in SECTIONS:
        lines.append(f"## {heading}")
        for paragraph in paragraphs:
            lines.append(paragraph)
            lines.append("")

    lines.append("## Tài liệu tham khảo")
    for idx, ref in enumerate(REFERENCES, start=1):
        lines.append(f"{idx}. {ref}")
    lines.append("")
    MD_OUT.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    create_docx()
    create_markdown()
    print(f"Created: {DOCX_OUT}")
    print(f"Created: {MD_OUT}")


if __name__ == "__main__":
    main()
