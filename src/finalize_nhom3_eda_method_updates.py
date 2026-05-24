from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOC_PATH = ROOT / "Document" / "Bai_Bao_Nghien_Cuu_M5_Cau_Truc_Chuan_Hoc_Thuat.docx"
OUTPUT_COPY = ROOT / "Document" / "Bai_Bao_Nghien_Cuu_M5_Nhom3_Format.docx"


THEORY_INTRO = (
    "Phần này hệ thống hóa các nền tảng business và technical được sử dụng trong "
    "nghiên cứu. Các mục lý thuyết được trình bày theo đúng vai trò trong pipeline: "
    "từ logic vận hành của dự báo bán lẻ, cấu trúc dữ liệu phân cấp, đặc điểm "
    "intermittent demand, đến mô hình LightGBM, clustering, rolling-origin và hệ "
    "tiêu chí đánh giá."
)


METHOD_DESCRIPTION = [
    (
        "Hình 2 mô tả pipeline đã dùng trong thực nghiệm sau các cập nhật kiểm soát "
        "leakage. Luồng xử lý bắt đầu từ dữ liệu sales, calendar, price và hierarchy, "
        "sau đó được chia thành hai nhánh feature song song tại từng rolling origin T."
    ),
    (
        "Nhánh forecasting tạo lag/rolling features, calendar/SNAP/event, price, "
        "availability và hierarchy features cho từng ngày dự báo. Nhánh clustering "
        "chỉ dùng dữ liệu lịch sử đến T để tính ADI, CV2, mean sales, zero-sales "
        "ratio, gap, event lift và price statistics; các biến này được log-transform, "
        "clipping và fit RobustScaler trên train fold trước khi chạy Mini-batch "
        "K-Means K=3."
    ),
    (
        "Từ cùng một feature pipeline, ba nhánh mô hình được huấn luyện để tách riêng "
        "hiệu ứng của cluster information. A0 là global LightGBM không dùng cụm. B1 "
        "là global LightGBM nhưng thêm cluster label như categorical feature. C tách "
        "dữ liệu theo cluster và huấn luyện một LightGBM riêng cho từng cụm. Forecast "
        "của cả ba nhánh được sinh đệ quy trong 28 ngày."
    ),
    (
        "Khối đánh giá cuối cùng gồm WRMSSE close-to-official theo hierarchy, WAPE, "
        "bias, scale-aware stability, JumpRate, train-test gap, ablation, K-Medoids "
        "robustness, DM test và multi-seed robustness. Thiết kế này phản ánh đúng thực "
        "nghiệm đã chạy: trọng tâm là A0, B1 và C, không so sánh theo leaderboard."
    ),
]


def find_heading(doc: Document, text: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading") and paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Heading not found: {text}")


def insert_after(anchor: Paragraph, text: str, style: str = "Normal") -> Paragraph:
    p = anchor._parent.add_paragraph(text, style=style)
    anchor._p.addnext(p._p)
    return p


def delete_element(element) -> None:
    element.getparent().remove(element)


def main() -> None:
    doc = Document(DOC_PATH)

    theory = find_heading(doc, "Cơ sở lý thuyết")
    next_para = Paragraph(theory._p.getnext(), doc)
    if next_para.text.strip() != THEORY_INTRO:
        insert_after(theory, THEORY_INTRO)

    method_heading = find_heading(doc, "Mô hình triển khai")
    cursor = method_heading._p.getnext()

    # Skip image and caption, then replace prose until the pseudocode marker.
    after_caption = False
    to_delete = []
    insertion_anchor = method_heading
    while cursor is not None:
        next_cursor = cursor.getnext()
        tag = cursor.tag.split("}")[-1]
        if tag == "p":
            p = Paragraph(cursor, doc)
            if p.style.name.startswith("Heading"):
                break
            if p.text.strip().startswith("Hình 2."):
                after_caption = True
                insertion_anchor = p
            elif after_caption and p.text.strip() == "Mã giả quy trình nghiên cứu:":
                break
            elif after_caption:
                to_delete.append(cursor)
        cursor = next_cursor

    # Avoid duplicating if already replaced.
    existing = "\n".join(p.text for p in doc.paragraphs)
    if "Nhánh forecasting tạo lag/rolling features" not in existing:
        for element in to_delete:
            delete_element(element)
        anchor = insertion_anchor
        for text in METHOD_DESCRIPTION:
            anchor = insert_after(anchor, text)

    doc.save(DOC_PATH)
    OUTPUT_COPY.write_bytes(DOC_PATH.read_bytes())
    print(f"Updated: {DOC_PATH}")
    print(f"Copy: {OUTPUT_COPY}")


if __name__ == "__main__":
    main()
