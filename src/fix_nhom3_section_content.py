from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
TARGET = ROOT / "Document" / "Bai_Bao_Nghien_Cuu_M5_Cau_Truc_Chuan_Hoc_Thuat.docx"


INTRO_RQ_TEXTS = {
    "RQ1: C có làm giảm WRMSSE so với A0 không?",
    "RQ2: C có làm giảm scale-aware stability loss so với A0 không?",
    "RQ3: B1 có cải thiện WRMSSE so với A0 không?",
    "RQ4: Ablation feature có làm thay đổi WRMSSE của C không?",
    "RQ5: Phân cụm K=3 có ổn định khi thay đổi thuật toán không?",
    (
        "Các giả thuyết tương ứng là H1: C có WRMSSE thấp hơn A0; "
        "H2: C có stability loss thấp hơn A0; H3: B1 không nhất thiết vượt A0; "
        "H4: calendar/SNAP/event và hierarchy/ID là các nhóm feature quan trọng; "
        "H5: cấu trúc cụm có độ nhạy khi đổi thuật toán."
    ),
}


RQ_BLOCK = [
    (
        "Nghiên cứu được triển khai nhằm kiểm định vai trò của cluster-aware "
        "learning trong dự báo nhu cầu bán lẻ quy mô lớn. Các câu hỏi nghiên "
        "cứu được đặt ra theo hướng tách riêng độ chính xác, độ ổn định, vai "
        "trò của nhãn cụm và độ bền của kết quả thực nghiệm."
    ),
    "RQ1: Mô hình C có làm giảm WRMSSE so với baseline A0 không?",
    "RQ2: Mô hình C có làm giảm scale-aware stability loss so với baseline A0 không?",
    "RQ3: Việc thêm cluster label trong mô hình B1 có cải thiện WRMSSE so với A0 không?",
    "RQ4: Các nhóm feature calendar/SNAP/event, price và hierarchy/ID ảnh hưởng thế nào đến WRMSSE của C?",
    "RQ5: Phân cụm K=3 có ổn định khi thay đổi thuật toán clustering không?",
    (
        "Từ các câu hỏi trên, nghiên cứu kiểm định các giả thuyết: H1, C có "
        "WRMSSE thấp hơn A0; H2, C có stability loss thấp hơn A0; H3, B1 "
        "không nhất thiết vượt A0 nếu cluster label chỉ được dùng như một "
        "feature rời rạc; H4, calendar/SNAP/event và hierarchy/ID là các "
        "nhóm feature quan trọng; H5, cấu trúc cụm có độ nhạy khi thay đổi "
        "thuật toán clustering."
    ),
]


RELATED_PROBLEM_BLOCK = [
    (
        "Các nghiên cứu gần đây về forecasting bán lẻ cho thấy global models "
        "và cross-learning có khả năng khai thác thông tin liên chuỗi tốt hơn "
        "nhiều mô hình cục bộ, đặc biệt trong dữ liệu lớn và có cấu trúc phân "
        "cấp. Các tổng kết về M5 cũng ghi nhận vai trò nổi bật của gradient "
        "boosting, ensemble và feature engineering trong bài toán demand "
        "forecasting quy mô lớn."
    ),
    (
        "Một nhánh nghiên cứu khác dùng phân loại demand pattern hoặc "
        "clustering để xử lý intermittent và lumpy demand. Các phương pháp "
        "này thường cải thiện diễn giải hoặc chọn mô hình theo nhóm chuỗi, "
        "nhưng ít nghiên cứu đánh giá đồng thời độ chính xác phân cấp, độ ổn "
        "định forecast, kiểm soát leakage theo rolling origin và robustness "
        "khi thay đổi seed hoặc thuật toán clustering."
    ),
    (
        "Khoảng trống nghiên cứu của bài này nằm ở việc đóng khung cluster-aware "
        "learning như một framework thực nghiệm có kiểm soát: clustering được "
        "tính historical-only tại từng origin, mô hình C huấn luyện riêng theo "
        "cụm, và kết quả được đánh giá bằng WRMSSE, scale-aware stability, "
        "ablation, overfitting diagnostics, Diebold-Mariano test và multi-seed "
        "robustness."
    ),
]


THEORY_BLOCK = [
    (
        "Cơ sở lý thuyết của nghiên cứu dựa trên ba nhóm khái niệm: global "
        "forecasting, phân cụm demand regimes và đánh giá dự báo theo cấu "
        "trúc phân cấp. Global forecasting học một hàm chung trên nhiều chuỗi, "
        "giúp chia sẻ thông tin giữa các item-store nhưng có thể bị suy giảm "
        "khi dữ liệu gồm nhiều demand regimes dị biệt."
    ),
    (
        "Cluster-aware forecasting xử lý vấn đề dị thể bằng cách nhóm các "
        "chuỗi có hành vi nhu cầu tương tự trước khi học mô hình. Trong nghiên "
        "cứu này, B1 dùng cluster label như một biến giải thích trong global "
        "LightGBM, còn C dùng cluster label để tách tập huấn luyện và huấn "
        "luyện LightGBM riêng cho từng cụm. So sánh A0-B1-C cho phép tách "
        "riêng hiệu ứng của thông tin cụm và hiệu ứng của cơ chế học theo cụm."
    ),
    (
        "Intermittent demand được mô tả bằng ADI và CV2. ADI đo khoảng cách "
        "trung bình giữa các lần bán khác 0, còn CV2 đo biến động tương đối "
        "của lượng bán trong các ngày có phát sinh nhu cầu. Nền tảng này giúp "
        "diễn giải bản đồ ADI-CV2 và hỗ trợ lựa chọn K=3 theo ba vùng vận hành "
        "chính: long-tail/intermittent, medium demand và high-demand/core."
    ),
    (
        "WRMSSE được dùng vì dữ liệu M5 có cấu trúc phân cấp. Chỉ số này kết "
        "hợp sai số đã chuẩn hóa theo scale với trọng số doanh thu, nhờ đó "
        "đánh giá đồng thời bottom-level item-store và các cấp tổng hợp như "
        "store, state, category và department. Stability loss bổ sung góc nhìn "
        "về mức dao động forecast giữa các rolling origins, đặc biệt quan "
        "trọng với các chuỗi có doanh số gần 0."
    ),
]


RELATED_FOUNDATION_BLOCK = [
    (
        "Bảng tổng hợp nghiên cứu liên quan ở phần trước cho thấy hai hướng "
        "tiếp cận chính: tối ưu hóa độ chính xác bằng global learning và xử lý "
        "dị thể bằng demand classification hoặc clustering. Bài nghiên cứu này "
        "kết hợp hai hướng đó nhưng đặt trọng tâm vào thiết kế leakage-aware "
        "rolling-origin và kiểm tra độ bền của kết quả."
    ),
    (
        "So với các nghiên cứu chỉ báo cáo accuracy trung bình, phần thực nghiệm "
        "còn đánh giá stability, train-test gap, ablation, K-Medoids robustness "
        "và multi-seed robustness. Cách tổ chức này giúp kết quả phù hợp hơn "
        "với yêu cầu của demand planning, nơi dự báo cần vừa chính xác vừa đủ "
        "ổn định để dùng trong quyết định vận hành."
    ),
]


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def find_heading(doc: Document, text: str, occurrence: int = 1) -> Paragraph:
    count = 0
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading") and paragraph.text.strip() == text:
            count += 1
            if count == occurrence:
                return paragraph
    raise ValueError(f"Heading not found: {text!r}, occurrence={occurrence}")


def insert_after(paragraph: Paragraph, entries: list[str], styles: list[str] | None = None) -> None:
    """Insert entries after paragraph in the same order."""
    styles = styles or ["Normal"] * len(entries)
    anchor = paragraph
    for text, style in zip(entries, styles):
        new_paragraph = paragraph._parent.add_paragraph(text, style=style)
        anchor._p.addnext(new_paragraph._p)
        anchor = new_paragraph


def style_existing_document(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.style.name == "Heading 1":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif paragraph.style.name == "Heading 2":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def section_has_content_until_next_heading(heading: Paragraph) -> bool:
    cursor = heading._p.getnext()
    while cursor is not None:
        tag = cursor.tag.split("}")[-1]
        if tag == "tbl":
            return True
        if tag == "p":
            p = Paragraph(cursor, heading._parent)
            if p.style.name.startswith("Heading"):
                return False
            if p.text.strip():
                return True
        cursor = cursor.getnext()
    return False


def main() -> None:
    doc = Document(TARGET)

    # Move research questions out of Introduction into the proper section.
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip() in INTRO_RQ_TEXTS:
            delete_paragraph(paragraph)

    insert_after(
        find_heading(doc, "Câu hỏi nghiên cứu"),
        RQ_BLOCK,
        ["Normal", "List Bullet", "List Bullet", "List Bullet", "List Bullet", "List Bullet", "Normal"],
    )

    insert_after(find_heading(doc, "Nghiên cứu liên quan", occurrence=1), RELATED_PROBLEM_BLOCK)
    insert_after(find_heading(doc, "Cơ sở lý thuyết"), THEORY_BLOCK)
    insert_after(find_heading(doc, "Nghiên cứu liên quan", occurrence=2), RELATED_FOUNDATION_BLOCK)

    style_existing_document(doc)
    doc.save(TARGET)

    # Validation output for review.
    doc = Document(TARGET)
    print(f"Updated: {TARGET}")
    print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} inline_shapes={len(doc.inline_shapes)}")
    print("Headings without immediate content:")
    missing = []
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading") and not section_has_content_until_next_heading(paragraph):
            missing.append(paragraph.text.strip())
    if missing:
        for text in missing:
            print(f"- {text}")
    else:
        print("- None")


if __name__ == "__main__":
    main()
