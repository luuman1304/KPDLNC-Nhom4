from __future__ import annotations

import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


TARGET = Path("Document/Bai_Bao_Nghien_Cuu_M5_Cau_Truc_Chuan_Hoc_Thuat.docx")
BACKUP = Path("Document/Bai_Bao_Nghien_Cuu_M5_Cau_Truc_Chuan_Hoc_Thuat.before_format.docx")
OUT = TARGET


def set_run_font(run, size: float = 13, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic


def set_paragraph_format(paragraph, align=None, first_line=None, before=0, after=0, line=1.5) -> None:
    pf = paragraph.paragraph_format
    if align is not None:
        paragraph.alignment = align
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if first_line is not None:
        pf.first_line_indent = Cm(first_line)


def clear_paragraph(paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        p.remove(child)


def add_field(paragraph, instr: str, placeholder: str = "") -> None:
    run_begin = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run_begin._r.append(fld_begin)

    run_instr = paragraph.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    run_instr._r.append(instr_text)

    run_sep = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run_sep._r.append(fld_sep)

    if placeholder:
        result = paragraph.add_run(placeholder)
        set_run_font(result, 13)

    run_end = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_end._r.append(fld_end)


def add_page_break(paragraph) -> None:
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def insert_block_before(anchor, paragraphs) -> None:
    parent = anchor._p.getparent()
    idx = parent.index(anchor._p)
    for para in paragraphs:
        parent.insert(idx, para._p)
        idx += 1


def make_para(doc: Document, text: str = "", style: str | None = None, align=None, size=13, bold=False, italic=False, line=1.5):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, italic=italic)
    set_paragraph_format(p, align=align, before=0, after=0, line=line)
    return p


def make_toc_para(doc: Document, instr: str):
    p = doc.add_paragraph()
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, line=1.5)
    add_field(p, instr)
    return p


def apply_styles(doc: Document) -> None:
    # Page setup follows the sample document.
    for section in doc.sections:
        section.top_margin = Cm(1.91)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        section.page_width = Cm(21.59)
        section.page_height = Cm(27.94)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
    normal.font.size = Pt(13)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)

    for name, size, align in [
        ("Title", 16, WD_ALIGN_PARAGRAPH.CENTER),
        ("Heading 1", 13, WD_ALIGN_PARAGRAPH.CENTER),
        ("Heading 2", 13, WD_ALIGN_PARAGRAPH.LEFT),
        ("Heading 3", 13, WD_ALIGN_PARAGRAPH.LEFT),
    ]:
        if name not in styles:
            continue
        st = styles[name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        st._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
        st.font.size = Pt(size)
        st.font.bold = True
        st.paragraph_format.line_spacing = 1.5
        st.paragraph_format.space_before = Pt(6 if name != "Title" else 12)
        st.paragraph_format.space_after = Pt(0)
        st.paragraph_format.alignment = align
        if name.startswith("Heading"):
            st.paragraph_format.first_line_indent = Cm(-0.635)
            st.paragraph_format.left_indent = Cm(0.635 if name == "Heading 1" else 1.27)


def format_body(doc: Document) -> None:
    for p in doc.paragraphs:
        text = p.text.strip()
        style = p.style.name
        if style == "Title":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=12, after=6, line=1.5)
            for r in p.runs:
                set_run_font(r, 16, bold=True)
        elif style == "Heading 1":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, after=0, line=1.5)
            for r in p.runs:
                set_run_font(r, 13, bold=True)
        elif style in {"Heading 2", "Heading 3", "Heading 4"}:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, before=6, after=0, line=1.5)
            for r in p.runs:
                set_run_font(r, 13, bold=True)
        elif text.startswith("Hình ") or text.startswith("Chú thích:"):
            try:
                p.style = doc.styles["Caption"]
            except Exception:
                pass
            set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=0, after=6, line=1.15)
            for r in p.runs:
                set_run_font(r, 11, italic=True)
        elif style.startswith("List"):
            set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=0, line=1.5)
            for r in p.runs:
                set_run_font(r, 13)
        else:
            set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=0, line=1.5)
            for r in p.runs:
                set_run_font(r, 13)

    for table in doc.tables:
        table.style = "Table Grid"
        for row_i, row in enumerate(table.rows):
            for cell in row.cells:
                for p in cell.paragraphs:
                    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=0, line=1.15)
                    for r in p.runs:
                        set_run_font(r, 10, bold=(row_i == 0))


def set_header_page_number(doc: Document) -> None:
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        clear_paragraph(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("Trang ")
        set_run_font(r, 13)
        add_field(p, "PAGE")


def add_update_fields_flag(doc: Document) -> None:
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def add_front_matter(doc: Document) -> None:
    title_text = doc.paragraphs[0].text.strip()
    temp = Document()
    for section in temp.sections:
        section.top_margin = Cm(1.91)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    cover = [
        make_para(temp, "BỘ GIÁO DỤC VÀ ĐÀO TẠO", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True),
        make_para(temp, "ĐẠI HỌC KINH TẾ THÀNH PHỐ HỒ CHÍ MINH", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True),
        make_para(temp, ""),
        make_para(temp, ""),
        make_para(temp, "BÁO CÁO NGHIÊN CỨU", align=WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True),
        make_para(temp, "Môn học: Khai phá dữ liệu nâng cao", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True),
        make_para(temp, "Nhóm 3", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True),
        make_para(temp, ""),
        make_para(temp, title_text.upper(), align=WD_ALIGN_PARAGRAPH.CENTER, size=15, bold=True),
        make_para(temp, ""),
        make_para(temp, "TP. Hồ Chí Minh, 2026", align=WD_ALIGN_PARAGRAPH.CENTER),
    ]
    add_page_break(cover[-1])

    toc_title = make_para(temp, "MỤC LỤC", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    toc = make_toc_para(temp, 'TOC \\o "1-3" \\h \\z \\u')
    add_page_break(toc)

    fig_title = make_para(temp, "MỤC LỤC HÌNH", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    fig_toc = make_toc_para(temp, 'TOC \\h \\z \\t "Caption,1"')
    make_para(temp, "")
    table_title = make_para(temp, "MỤC LỤC BẢNG", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    table_toc = make_toc_para(temp, 'TOC \\h \\z \\c "Table"')
    add_page_break(table_toc)

    front = [deepcopy(p) for p in temp.paragraphs]
    insert_block_before(doc.paragraphs[0], front)


def main() -> None:
    if not BACKUP.exists():
        shutil.copy2(TARGET, BACKUP)
    source = BACKUP if BACKUP.exists() else TARGET
    doc = Document(source)
    apply_styles(doc)
    format_body(doc)
    add_front_matter(doc)
    set_header_page_number(doc)
    add_update_fields_flag(doc)
    doc.save(OUT)
    print(OUT)
    print(BACKUP)


if __name__ == "__main__":
    main()
