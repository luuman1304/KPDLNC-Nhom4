from __future__ import annotations

from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import pandas as pd
import seaborn as sns
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


FIG_DIR = Path("outputs_research_figures")
DOCX_OUT = Path("Document/Phu_Luc_Bieu_Do_Ket_Qua_Nghien_Cuu_M5.docx")
MD_OUT = Path("Document/Phu_Luc_Bieu_Do_Ket_Qua_Nghien_Cuu_M5.md")

BASE_METRICS = Path("outputs_full_scope_a0_b1_c/metrics/full_scope_all_metrics.csv")
EXT_METRICS = Path("outputs_extended_fullscale_checks/metrics/extended_fullscale_all_metrics.csv")
CLUSTER_PROFILE = Path("outputs_extended_fullscale_checks/metrics/extended_fullscale_cluster_regime_profile.csv")
TOP_FEATURES = Path("outputs_full_feature_importance_drift_tweedie/feature_importance_top_features.csv")


MODEL_LABELS = {
    "A0_global_baseline": "A0 Global",
    "B1_cluster_label": "B1 Cluster label",
    "C_cluster_specific": "C Cluster-specific",
}

RUN_LABELS = {
    "base_tweedie": "Base",
    "no_price": "No price",
    "no_calendar": "No calendar/SNAP/event",
    "no_hierarchy": "No hierarchy/ID",
}


def setup() -> None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    sns.set_theme(style="whitegrid", context="paper", font_scale=1.1)
    plt.rcParams["figure.dpi"] = 150
    plt.rcParams["savefig.dpi"] = 220
    plt.rcParams["axes.unicode_minus"] = False


def savefig(name: str) -> Path:
    path = FIG_DIR / name
    plt.tight_layout()
    plt.savefig(path, bbox_inches="tight")
    plt.close()
    return path


def add_value_labels(ax, fmt="{:.3f}", rotation=0) -> None:
    for container in ax.containers:
        ax.bar_label(container, fmt=fmt, padding=3, fontsize=8, rotation=rotation)


def figure_base_wrmsse(base: pd.DataFrame) -> Path:
    df = base.loc[base["run"] == "base_tweedie"].copy()
    df["model_label"] = df["model"].map(MODEL_LABELS)
    order = ["A0 Global", "B1 Cluster label", "C Cluster-specific"]
    fig, ax = plt.subplots(figsize=(7.2, 4.2))
    sns.barplot(data=df, x="model_label", y="rolling_wrmsse_mean", order=order, ax=ax, color="#4C78A8")
    add_value_labels(ax, "{:.3f}")
    ax.set_title("Rolling close-to-official WRMSSE của A0/B1/C")
    ax.set_xlabel("Mô hình")
    ax.set_ylabel("Rolling WRMSSE thấp hơn là tốt hơn")
    ax.set_ylim(max(0.55, df["rolling_wrmsse_mean"].min() - 0.02), df["rolling_wrmsse_mean"].max() + 0.02)
    return savefig("fig01_base_wrmsse_a0_b1_c.png")


def figure_base_stability(base: pd.DataFrame) -> Path:
    df = base.loc[base["run"] == "base_tweedie"].copy()
    df["model_label"] = df["model"].map(MODEL_LABELS)
    order = ["A0 Global", "B1 Cluster label", "C Cluster-specific"]
    fig, ax = plt.subplots(figsize=(7.2, 4.2))
    sns.barplot(data=df, x="model_label", y="scale_aware_stability_loss", order=order, ax=ax, color="#59A14F")
    add_value_labels(ax, "{:.3f}")
    ax.set_title("Scale-aware stability loss của A0/B1/C")
    ax.set_xlabel("Mô hình")
    ax.set_ylabel("Stability loss thấp hơn là ổn định hơn")
    ax.set_ylim(max(0.11, df["scale_aware_stability_loss"].min() - 0.01), df["scale_aware_stability_loss"].max() + 0.015)
    return savefig("fig02_base_stability_a0_b1_c.png")


def figure_accuracy_stability(base: pd.DataFrame) -> Path:
    df = base.loc[base["run"] == "base_tweedie"].copy()
    df["model_label"] = df["model"].map(MODEL_LABELS)
    fig, ax = plt.subplots(figsize=(7.2, 4.6))
    sns.scatterplot(
        data=df,
        x="rolling_wrmsse_mean",
        y="scale_aware_stability_loss",
        hue="model_label",
        s=180,
        ax=ax,
        palette=["#F28E2B", "#4E79A7", "#59A14F"],
    )
    for _, row in df.iterrows():
        ax.annotate(row["model_label"], (row["rolling_wrmsse_mean"], row["scale_aware_stability_loss"]), xytext=(7, 5), textcoords="offset points")
    ax.set_title("Trade-off giữa accuracy và stability")
    ax.set_xlabel("Rolling WRMSSE thấp hơn là tốt hơn")
    ax.set_ylabel("Stability loss thấp hơn là tốt hơn")
    ax.legend().remove()
    return savefig("fig03_accuracy_stability_tradeoff.png")


def figure_ablation_wrmsse(ext: pd.DataFrame) -> Path:
    df = ext.loc[ext["model"] == "C_cluster_specific"].copy()
    df["run_label"] = df["run"].map(RUN_LABELS)
    order = ["Base", "No price", "No calendar/SNAP/event", "No hierarchy/ID"]
    fig, ax = plt.subplots(figsize=(8.8, 4.6))
    sns.barplot(data=df, x="run_label", y="rolling_wrmsse_mean", order=order, ax=ax, palette=["#59A14F", "#F28E2B", "#E15759", "#B07AA1"])
    add_value_labels(ax, "{:.3f}")
    ax.set_title("Tác động của ablation lên rolling WRMSSE của model C")
    ax.set_xlabel("Cấu hình")
    ax.set_ylabel("Rolling WRMSSE thấp hơn là tốt hơn")
    ax.tick_params(axis="x", rotation=15)
    return savefig("fig04_ablation_wrmsse_c.png")


def figure_ablation_delta(ext: pd.DataFrame) -> Path:
    df = ext.loc[ext["model"] == "C_cluster_specific"].copy()
    base = float(df.loc[df["run"] == "base_tweedie", "rolling_wrmsse_mean"].iloc[0])
    df = df.loc[df["run"] != "base_tweedie"].copy()
    df["delta"] = df["rolling_wrmsse_mean"] - base
    df["run_label"] = df["run"].map(RUN_LABELS)
    order = ["No price", "No calendar/SNAP/event", "No hierarchy/ID"]
    fig, ax = plt.subplots(figsize=(8.8, 4.4))
    sns.barplot(data=df, x="run_label", y="delta", order=order, ax=ax, palette=["#F28E2B", "#E15759", "#B07AA1"])
    add_value_labels(ax, "{:+.3f}")
    ax.axhline(0, color="black", linewidth=0.8)
    ax.set_title("Mức xấu đi WRMSSE của model C khi loại từng nhóm feature")
    ax.set_xlabel("Ablation")
    ax.set_ylabel("Δ Rolling WRMSSE so với base")
    ax.tick_params(axis="x", rotation=15)
    return savefig("fig05_ablation_delta_wrmsse_c.png")


def figure_cluster_profile(profile: pd.DataFrame) -> Path:
    df = profile.copy()
    df["regime"] = df["regime_interpretation"]
    long = df.melt(
        id_vars=["regime"],
        value_vars=["mean_sales", "zero_sales_ratio", "adi", "cv2"],
        var_name="metric",
        value_name="value",
    )
    fig, axes = plt.subplots(2, 2, figsize=(10, 7))
    axes = axes.flatten()
    metrics = [
        ("mean_sales", "Mean sales"),
        ("zero_sales_ratio", "Zero-sales ratio"),
        ("adi", "ADI"),
        ("cv2", "CV2"),
    ]
    for ax, (metric, title) in zip(axes, metrics):
        sub = long.loc[long["metric"] == metric]
        sns.barplot(data=sub, x="regime", y="value", ax=ax, color="#4C78A8")
        ax.set_title(title)
        ax.set_xlabel("")
        ax.set_ylabel("")
        ax.tick_params(axis="x", rotation=20)
        add_value_labels(ax, "{:.2f}", rotation=0)
    fig.suptitle("Đặc trưng hành vi của ba cụm nhu cầu K=3", y=1.02, fontsize=13)
    return savefig("fig06_cluster_profile_k3.png")


def figure_overfitting_gap(base: pd.DataFrame) -> Path:
    df = base.loc[base["run"] == "base_tweedie"].copy()
    df["model_label"] = df["model"].map(MODEL_LABELS)
    order = ["A0 Global", "B1 Cluster label", "C Cluster-specific"]
    fig, ax = plt.subplots(figsize=(7.2, 4.2))
    sns.barplot(data=df, x="model_label", y="mean_test_train_gap", order=order, ax=ax, color="#B07AA1")
    add_value_labels(ax, "{:.3f}")
    ax.set_title("Mean test-train gap của A0/B1/C")
    ax.set_xlabel("Mô hình")
    ax.set_ylabel("Test-train gap")
    ax.set_ylim(max(0.09, df["mean_test_train_gap"].min() - 0.005), df["mean_test_train_gap"].max() + 0.008)
    return savefig("fig07_overfitting_gap_a0_b1_c.png")


def figure_feature_importance(top_features: pd.DataFrame) -> Path:
    df = top_features.loc[top_features["model"] == "C_cluster_specific"].head(12).copy()
    df["importance_gain_m"] = df["importance_gain"] / 1_000_000
    fig, ax = plt.subplots(figsize=(8, 5.2))
    sns.barplot(data=df, y="feature", x="importance_gain_m", ax=ax, color="#4E79A7")
    ax.set_title("Top feature importance của C_cluster_specific")
    ax.set_xlabel("Importance gain (triệu)")
    ax.set_ylabel("Feature")
    return savefig("fig08_feature_importance_c.png")


def build_markdown(figures: list[tuple[Path, str, str]]) -> None:
    lines = [
        "# Phụ lục biểu đồ kết quả nghiên cứu M5",
        "",
        "Phụ lục này bổ sung các biểu đồ trực quan cho phần kết quả nghiên cứu. Các biểu đồ được sinh trực tiếp từ output full-scale đã chạy, không huấn luyện lại mô hình.",
        "",
    ]
    for idx, (path, title, caption) in enumerate(figures, start=1):
        lines.extend(
            [
                f"## Hình {idx}. {title}",
                "",
                f"![{title}](../{path.as_posix()})",
                "",
                caption,
                "",
            ]
        )
    MD_OUT.write_text("\n".join(lines), encoding="utf-8")


def build_docx(figures: list[tuple[Path, str, str]]) -> None:
    doc = Document()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Phụ lục biểu đồ kết quả nghiên cứu M5")
    run.bold = True
    run.font.size = Pt(16)

    intro = doc.add_paragraph(
        "Các biểu đồ dưới đây được tạo từ kết quả full-scale đã chạy. Mục đích là trực quan hóa các kết luận chính về accuracy, stability, ablation, phân cụm, overfitting và feature importance."
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for idx, (path, title_text, caption) in enumerate(figures, start=1):
        doc.add_heading(f"Hình {idx}. {title_text}", level=2)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(6.3))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in cap.runs:
            run.font.size = Pt(10)

    doc.save(DOCX_OUT)


def main() -> None:
    setup()
    base = pd.read_csv(BASE_METRICS)
    ext = pd.read_csv(EXT_METRICS)
    profile = pd.read_csv(CLUSTER_PROFILE)
    top_features = pd.read_csv(TOP_FEATURES)

    figures = [
        (
            figure_base_wrmsse(base),
            "So sánh rolling WRMSSE của A0, B1 và C",
            "Biểu đồ cho thấy C_cluster_specific có rolling WRMSSE thấp nhất, tức độ chính xác tốt nhất trong ba mô hình chính.",
        ),
        (
            figure_base_stability(base),
            "So sánh scale-aware stability loss của A0, B1 và C",
            "C_cluster_specific có stability loss thấp nhất, cho thấy mô hình này không chỉ chính xác hơn mà còn ổn định hơn qua rolling origins.",
        ),
        (
            figure_accuracy_stability(base),
            "Trade-off giữa accuracy và stability",
            "Điểm nằm gần góc dưới bên trái là tốt hơn. C_cluster_specific chiếm ưu thế đồng thời về WRMSSE và stability loss.",
        ),
        (
            figure_ablation_wrmsse(ext),
            "Tác động của ablation lên WRMSSE của model C",
            "Biểu đồ cho thấy no-calendar/SNAP/event làm WRMSSE tăng mạnh nhất, xác nhận nhóm feature này là thành phần thiết yếu.",
        ),
        (
            figure_ablation_delta(ext),
            "Mức tăng WRMSSE khi loại từng nhóm feature",
            "Calendar/SNAP/event có tác động lớn nhất, tiếp theo là hierarchy/ID features; price features có tác động nhỏ hơn nhưng vẫn làm kết quả xấu đi.",
        ),
        (
            figure_cluster_profile(profile),
            "Đặc trưng ba cụm nhu cầu K=3",
            "Ba cụm tương ứng với long-tail/intermittent demand, medium demand và high-demand/core products, phù hợp với giả thuyết demand regimes của proposal.",
        ),
        (
            figure_overfitting_gap(base),
            "Mean test-train gap của A0, B1 và C",
            "C có gap cao hơn A0/B1 nhẹ nhưng không bất thường; kết quả không cho thấy overfitting nghiêm trọng.",
        ),
        (
            figure_feature_importance(top_features),
            "Top feature importance của C_cluster_specific",
            "Các rolling mean và định danh item là những feature quan trọng nhất, cho thấy mô hình dựa nhiều vào tín hiệu nhu cầu gần đây và đặc trưng sản phẩm.",
        ),
    ]
    build_markdown(figures)
    build_docx(figures)

    manifest = pd.DataFrame(
        [{"figure": idx, "title": title, "path": path.as_posix(), "caption": caption} for idx, (path, title, caption) in enumerate(figures, start=1)]
    )
    manifest.to_csv(FIG_DIR / "figure_manifest.csv", index=False)
    print(DOCX_OUT)
    print(MD_OUT)
    print(FIG_DIR)


if __name__ == "__main__":
    main()
