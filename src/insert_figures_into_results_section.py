from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


DOCX = Path("Document/Bai_Bao_Nghien_Cuu_M5_Cau_Truc_Chuan_Hoc_Thuat.docx")
FIGS = Path("outputs_research_figures")
EDA = Path("outputs_eda_figures")


def paragraph_after(ref: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new = OxmlElement("w:p")
    ref._element.addnext(new)
    p = Paragraph(new, ref._parent)
    if style:
        p.style = style
    if text:
        run = p.add_run(text)
        run.font.size = Pt(10 if text.startswith("Chú thích:") else 11)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_fig_after(ref: Paragraph, path: Path, caption: str, width: float = 6.3) -> Paragraph:
    cap = paragraph_after(ref, f"Chú thích: {caption}")
    pic = paragraph_after(ref)
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic.add_run().add_picture(str(path), width=Inches(width))
    return cap


def find_heading(doc: Document, text: str) -> Paragraph:
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    raise RuntimeError(f"Không tìm thấy heading: {text}")


def main() -> None:
    doc = Document(DOCX)

    # Insert in reverse order inside each section so final order follows the call list comments.
    h61 = find_heading(doc, "6.1. Kết quả tổng hợp")
    anchor = h61
    anchor = add_fig_after(anchor, FIGS / "fig03_accuracy_stability_tradeoff.png", "Trade-off giữa accuracy và stability. C nằm ở vùng tốt hơn vì có WRMSSE và stability loss thấp hơn A0/B1.")
    anchor = add_fig_after(anchor, FIGS / "fig02_base_stability_a0_b1_c.png", "So sánh scale-aware stability loss của A0, B1 và C. C có stability loss thấp nhất.")
    anchor = add_fig_after(anchor, FIGS / "fig01_base_wrmsse_a0_b1_c.png", "So sánh rolling WRMSSE của A0, B1 và C. C là mô hình có WRMSSE thấp nhất trong ba cấu hình trọng tâm.")

    h63 = find_heading(doc, "6.3. Kết quả theo cluster và demand regime")
    anchor = h63
    anchor = add_fig_after(anchor, EDA / "eda_fig02_zero_sales_ratio_distribution.png", "Phân phối zero-sales ratio cho thấy mức độ sparse demand cao; đây là cơ sở để dùng scale-aware stability.")
    anchor = add_fig_after(anchor, EDA / "eda_fig01_demand_class_counts.png", "Phân bố chuỗi theo demand regime. Intermittent và lumpy chiếm tỷ trọng lớn trong dữ liệu.")
    anchor = add_fig_after(anchor, FIGS / "fig06_cluster_profile_k3.png", "Profile các cụm K=3, thể hiện khác biệt về mean sales, zero-sales ratio, ADI và CV2.")

    h64 = find_heading(doc, "6.4. Kết quả theo category, store và state")
    anchor = h64
    anchor = add_fig_after(anchor, EDA / "eda_fig05_store_sales_zero_ratio.png", "Khác biệt nhu cầu theo cửa hàng, cho thấy store/state là tín hiệu quan trọng trong hierarchy features.")
    anchor = add_fig_after(anchor, EDA / "eda_fig04_category_sales_zero_ratio.png", "Khác biệt nhu cầu theo category, cho thấy FOODS, HOBBIES và HOUSEHOLD có demand profile khác nhau.")

    h65 = find_heading(doc, "6.5. Ablation, kiểm định thống kê và robustness")
    anchor = h65
    anchor = add_fig_after(anchor, FIGS / "fig08_feature_importance_c.png", "Top feature importance của C_cluster_specific, cho thấy rolling demand và ID/hierarchy features đóng vai trò lớn.")
    anchor = add_fig_after(anchor, FIGS / "fig07_overfitting_gap_a0_b1_c.png", "Mean test-train gap của A0, B1 và C. Kết quả không cho thấy C bị overfitting nghiêm trọng.")
    anchor = add_fig_after(anchor, FIGS / "fig05_ablation_delta_wrmsse_c.png", "Mức tăng WRMSSE khi loại từng nhóm feature khỏi C. Calendar/SNAP/event gây suy giảm lớn nhất.")
    anchor = add_fig_after(anchor, FIGS / "fig04_ablation_wrmsse_c.png", "Tác động của ablation lên WRMSSE của model C.")

    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
