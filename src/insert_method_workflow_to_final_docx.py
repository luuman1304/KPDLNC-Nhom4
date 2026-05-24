from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


DOCX = Path("Document/Bai_Viet_Nghien_Cuu_Khoa_Hoc_Chi_Tiet_Hoc_Thuat_M5_Co_Bieu_Do.docx")
FIGURE = Path("outputs_method_figures/method_workflow.png")
HEADING = "5.0 Sơ đồ quy trình phương pháp thực hiện"


def paragraph_after(ref: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new = OxmlElement("w:p")
    ref._element.addnext(new)
    paragraph = Paragraph(new, ref._parent)
    if style:
        paragraph.style = style
    if text:
        run = paragraph.add_run(text)
        run.font.size = Pt(11)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return paragraph


def main() -> None:
    doc = Document(DOCX)

    # Remove previous inserted workflow block when rerunning the script.
    body = doc._body._element
    start = None
    end = None
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() == HEADING:
            start = idx
        if start is not None and paragraph.text.strip() == "5.1 Thiết kế rolling-origin evaluation":
            end = idx
            break
    if start is not None and end is not None:
        for paragraph in doc.paragraphs[start:end]:
            body.remove(paragraph._element)

    target = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "5. Phương pháp nghiên cứu":
            target = paragraph
            break
    if target is None:
        raise RuntimeError("Không tìm thấy mục 5. Phương pháp nghiên cứu")

    caption_text = (
        "Hình quy trình phương pháp thực hiện. Sơ đồ tóm tắt toàn bộ luồng nghiên cứu từ dữ liệu M5, "
        "kiểm tra dữ liệu, EDA, rolling-origin split, feature engineering, phân cụm theo từng origin, "
        "huấn luyện A0/B1/C, recursive forecasting, đánh giá WRMSSE và scale-aware stability, đến các kiểm tra bổ sung như overfitting, ablation và robustness. "
        "Điểm quan trọng của quy trình là tất cả feature và clustering tại origin T chỉ sử dụng dữ liệu có sẵn đến T để tránh data leakage."
    )

    # Insert in reverse order after section 5 heading so the final order is heading, explanation, figure, caption.
    cap = paragraph_after(target, caption_text)
    cap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in cap.runs:
        run.font.size = Pt(10)

    pic_p = paragraph_after(target)
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.add_run().add_picture(str(FIGURE), width=Inches(6.6))

    explanation = paragraph_after(
        target,
        "Sơ đồ dưới đây được bổ sung để trực quan hóa phương pháp nghiên cứu. Quy trình được thiết kế theo hướng anti-leakage: dữ liệu được kiểm tra và phân tích trước, sau đó mỗi rolling origin tạo feature, cluster assignment và mô hình dự báo riêng trên phần dữ liệu lịch sử. Kết quả cuối cùng được đánh giá đồng thời theo độ chính xác, độ ổn định và các kiểm tra độ tin cậy.",
    )
    explanation.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    paragraph_after(target, HEADING, "Heading 3")

    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
