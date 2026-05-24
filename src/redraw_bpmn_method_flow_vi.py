from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOC_PATH = ROOT / "Document" / "Bai_Bao_Nghien_Cuu_M5_Cau_Truc_Chuan_Hoc_Thuat.docx"
OUTPUT_COPY = ROOT / "Document" / "Bai_Bao_Nghien_Cuu_M5_Nhom3_Format.docx"
FIG_DIR = ROOT / "Document" / "figures"
FIG_PATH = FIG_DIR / "method_flow_bpmn_vi.png"


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    paths = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Times New Roman Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Times New Roman.ttf",
    ]
    for path in paths:
        try:
            return ImageFont.truetype(path, size)
        except OSError:
            continue
    return ImageFont.load_default()


TITLE_FONT = load_font(42, True)
TASK_FONT = load_font(24)
TASK_BOLD = load_font(27, True)
LANE_FONT = load_font(26, True)
NOTE_FONT = load_font(21)


def centered_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], lines: list[str], title: str | None = None) -> None:
    x1, y1, x2, y2 = box
    all_lines = ([title] if title else []) + lines
    heights = [34 if i == 0 and title else 30 for i, _ in enumerate(all_lines)]
    y = y1 + ((y2 - y1) - sum(heights)) / 2
    for i, line in enumerate(all_lines):
        f = TASK_BOLD if i == 0 and title else TASK_FONT
        bbox = draw.textbbox((0, 0), line, font=f)
        draw.text((x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2, y), line, fill="#17202A", font=f)
        y += heights[i]


def task(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], title: str, lines: list[str], fill: str = "#F8F9F9") -> None:
    draw.rounded_rectangle(box, radius=16, fill=fill, outline="#2C3E50", width=4)
    centered_text(draw, box, lines, title)


def event(draw: ImageDraw.ImageDraw, center: tuple[int, int], radius: int, text: str, end: bool = False) -> None:
    x, y = center
    width = 6 if end else 4
    draw.ellipse((x - radius, y - radius, x + radius, y + radius), fill="#FFFFFF", outline="#2C3E50", width=width)
    bbox = draw.textbbox((0, 0), text, font=TASK_BOLD)
    draw.text((x - (bbox[2] - bbox[0]) / 2, y - (bbox[3] - bbox[1]) / 2 - 2), text, fill="#17202A", font=TASK_BOLD)


def gateway(draw: ImageDraw.ImageDraw, center: tuple[int, int], size: int, title: str, lines: list[str]) -> None:
    x, y = center
    pts = [(x, y - size), (x + size, y), (x, y + size), (x - size, y)]
    draw.polygon(pts, fill="#FCF3CF", outline="#2C3E50")
    draw.line((x - 18, y - 18, x + 18, y + 18), fill="#2C3E50", width=4)
    draw.line((x + 18, y - 18, x - 18, y + 18), fill="#2C3E50", width=4)
    label_box = (x - 190, y + size + 8, x + 190, y + size + 85)
    centered_text(draw, label_box, lines, title)


def arrow_head(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int]) -> None:
    sx, sy = start
    ex, ey = end
    if abs(ex - sx) >= abs(ey - sy):
        pts = [(ex, ey), (ex - 20, ey - 11), (ex - 20, ey + 11)] if ex >= sx else [(ex, ey), (ex + 20, ey - 11), (ex + 20, ey + 11)]
    else:
        pts = [(ex, ey), (ex - 11, ey - 20), (ex + 11, ey - 20)] if ey >= sy else [(ex, ey), (ex - 11, ey + 20), (ex + 11, ey + 20)]
    draw.polygon(pts, fill="#2C3E50")


def flow(draw: ImageDraw.ImageDraw, points: list[tuple[int, int]]) -> None:
    for a, b in zip(points, points[1:]):
        draw.line((a, b), fill="#2C3E50", width=4)
    arrow_head(draw, points[-2], points[-1])


def create_figure() -> None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    width, height = 3300, 2300
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)

    title = "Quy trình BPMN triển khai mô hình dự báo cluster-aware trên dữ liệu M5"
    bbox = draw.textbbox((0, 0), title, font=TITLE_FONT)
    draw.text(((width - (bbox[2] - bbox[0])) / 2, 45), title, fill="#17202A", font=TITLE_FONT)

    lanes = [
        (120, 140, 3180, 500, "Chuẩn bị dữ liệu"),
        (120, 535, 3180, 980, "Tạo đặc trưng"),
        (120, 1015, 3180, 1530, "Huấn luyện mô hình"),
        (120, 1565, 3180, 2160, "Dự báo và đánh giá"),
    ]
    for x1, y1, x2, y2, name in lanes:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=20, fill="#FBFCFC", outline="#B2BABB", width=2)
        draw.rectangle((x1, y1, x1 + 70, y2), fill="#EAECEE", outline="#B2BABB", width=2)
        # vertical-ish lane label split into short lines
        parts = name.split()
        yy = y1 + 40
        for part in parts:
            tb = draw.textbbox((0, 0), part, font=LANE_FONT)
            draw.text((x1 + 35 - (tb[2] - tb[0]) / 2, yy), part, fill="#17202A", font=LANE_FONT)
            yy += 34

    # Row 1
    event(draw, (240, 320), 55, "Bắt đầu")
    task(draw, (380, 230, 760, 410), "Thu thập dữ liệu", ["sales, calendar", "price, hierarchy"], "#EAF2F8")
    task(draw, (900, 230, 1280, 410), "Kiểm tra dữ liệu", ["schema, missing", "range, duplicate"], "#EAF2F8")
    task(draw, (1420, 230, 1840, 410), "Phân tích EDA", ["zero-sales, ADI/CV²", "category, store, state"], "#EAF2F8")
    task(draw, (1980, 230, 2420, 410), "Chọn mốc rolling", ["1885, 1892, 1899", "1906, 1913"], "#EAF2F8")
    task(draw, (2580, 230, 3080, 410), "Khóa dữ liệu lịch sử", ["chỉ dùng dữ liệu <= T", "tách train/validation/test"], "#FDEDEC")

    flow(draw, [(295, 320), (380, 320)])
    flow(draw, [(760, 320), (900, 320)])
    flow(draw, [(1280, 320), (1420, 320)])
    flow(draw, [(1840, 320), (1980, 320)])
    flow(draw, [(2420, 320), (2580, 320)])

    # Gateway 1 and features
    gateway(draw, (500, 735), 62, "Tách nhánh", ["đặc trưng"])
    task(draw, (760, 610, 1250, 820), "Tạo đặc trưng dự báo", ["shift lag/rolling", "calendar/SNAP/event", "price, availability", "hierarchy"], "#F4F6F7")
    task(draw, (1470, 610, 1960, 820), "Tạo đặc trưng phân cụm", ["mean sales, ADI/CV²", "zero ratio, gap", "spike, event lift", "price statistics"], "#F4F6F7")
    task(draw, (2180, 610, 2630, 820), "Tiền xử lý cụm", ["log transform", "clipping", "fit RobustScaler"], "#F4F6F7")
    task(draw, (2800, 610, 3100, 820), "Gán cụm", ["Mini-batch K-Means", "K = 3"], "#FCF3CF")

    flow(draw, [(2830, 410), (2830, 510), (500, 510), (500, 673)])
    flow(draw, [(562, 735), (760, 735)])
    flow(draw, [(500, 797), (500, 900), (1715, 900), (1715, 820)])
    flow(draw, [(1960, 735), (2180, 735)])
    flow(draw, [(2630, 735), (2800, 735)])

    # Gateway 2 and models
    gateway(draw, (500, 1250), 62, "Tách nhánh", ["training"])
    task(draw, (760, 1085, 1180, 1275), "Huấn luyện A0", ["Global LightGBM", "không dùng cluster", "baseline chính"], "#E8F6F3")
    task(draw, (1440, 1190, 1860, 1380), "Huấn luyện B1", ["Global LightGBM", "thêm cluster label", "kiểm tra vai trò cụm"], "#E8F6F3")
    task(draw, (2120, 1295, 2540, 1485), "Huấn luyện C", ["tách dữ liệu theo cụm", "một LightGBM mỗi cụm", "cluster-specific"], "#E8F6F3")
    gateway(draw, (2665, 1285), 45, "Gom nhánh", ["mô hình"])
    task(draw, (2820, 1190, 3140, 1380), "Lưu đầu ra", ["mô hình", "forecast", "metric"], "#E8F6F3")

    # Forecasting features and cluster labels enter the training split from above.
    flow(draw, [(1005, 820), (1005, 1010), (500, 1010), (500, 1188)])
    flow(draw, [(2950, 820), (2950, 1010), (500, 1010)])

    # Three clean parallel branches from the training gateway.
    flow(draw, [(562, 1250), (660, 1250), (660, 1180), (760, 1180)])
    flow(draw, [(500, 1188), (500, 1065), (1650, 1065), (1650, 1190)])
    flow(draw, [(500, 1188), (500, 1035), (2330, 1035), (2330, 1295)])

    # Merge branches below/around the tasks without crossing task text.
    flow(draw, [(1180, 1180), (1240, 1180), (1240, 1530), (2665, 1530), (2665, 1330)])
    flow(draw, [(1860, 1285), (2665, 1285)])
    flow(draw, [(2540, 1390), (2665, 1390), (2665, 1330)])
    flow(draw, [(2710, 1285), (2820, 1285)])

    # Forecast and evaluation
    task(draw, (360, 1690, 880, 1905), "Sinh dự báo", ["recursive 28 ngày", "không teacher forcing", "không dùng actual tương lai"], "#EBF5FB")
    task(draw, (1070, 1690, 1590, 1905), "Tính metric", ["WRMSSE, WAPE, Bias", "Stability, JumpRate", "train-test gap"], "#FEF5E7")
    task(draw, (1780, 1690, 2320, 1905), "Kiểm tra độ tin cậy", ["ablation", "K-Medoids robustness", "DM, Friedman/Nemenyi"], "#FEF5E7")
    task(draw, (2510, 1690, 2980, 1905), "Tổng hợp kết quả", ["bảng, biểu đồ", "multi-seed summary", "kết luận A0/B1/C"], "#FEF5E7")
    event(draw, (3100, 1798), 58, "Kết thúc", end=True)

    flow(draw, [(2980, 1380), (2980, 1600), (620, 1600), (620, 1690)])
    flow(draw, [(880, 1798), (1070, 1798)])
    flow(draw, [(1590, 1798), (1780, 1798)])
    flow(draw, [(2320, 1798), (2510, 1798)])
    flow(draw, [(2980, 1798), (3042, 1798)])

    note = "Ghi chú: Calendar/SNAP/event và price là biến ngoại sinh biết trước; mọi feature từ actual sales chỉ dùng dữ liệu lịch sử tại rolling origin T."
    nb = draw.textbbox((0, 0), note, font=NOTE_FONT)
    draw.text(((width - (nb[2] - nb[0])) / 2, 2205), note, fill="#566573", font=NOTE_FONT)

    img.save(FIG_PATH)


def find_heading(doc: Document, text: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading") and paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Heading not found: {text}")


def replace_figure() -> None:
    doc = Document(DOC_PATH)
    heading = find_heading(doc, "Mô hình triển khai")
    cursor = heading._p.getnext()
    while cursor is not None:
        next_cursor = cursor.getnext()
        if cursor.tag.split("}")[-1] == "p":
            p = Paragraph(cursor, doc)
            if p.style.name.startswith("Heading"):
                break
            if cursor.xpath(".//*[local-name()='drawing']") or p.text.strip().startswith("Hình 2."):
                cursor.getparent().remove(cursor)
        cursor = next_cursor

    image_p = heading._parent.add_paragraph()
    image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_p.add_run().add_picture(str(FIG_PATH), width=Inches(6.8))
    heading._p.addnext(image_p._p)
    caption = heading._parent.add_paragraph(
        "Hình 2. Quy trình BPMN triển khai mô hình dự báo cluster-aware trên dữ liệu M5.",
        style="Caption",
    )
    image_p._p.addnext(caption._p)
    doc.save(DOC_PATH)
    OUTPUT_COPY.write_bytes(DOC_PATH.read_bytes())


def main() -> None:
    create_figure()
    replace_figure()
    doc = Document(DOC_PATH)
    print(f"Updated figure: {FIG_PATH}")
    print(f"Updated document: {DOC_PATH}")
    print(f"Copy: {OUTPUT_COPY}")
    print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} inline_shapes={len(doc.inline_shapes)}")


if __name__ == "__main__":
    main()
