from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOC_PATH = ROOT / "Document" / "Bai_Bao_Nghien_Cuu_M5_Cau_Truc_Chuan_Hoc_Thuat.docx"
OUTPUT_COPY = ROOT / "Document" / "Bai_Bao_Nghien_Cuu_M5_Nhom3_Format.docx"
FIG_DIR = ROOT / "Document" / "figures"
FIG_PATH = FIG_DIR / "method_flow_detailed_vi.png"


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Times New Roman Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Times New Roman.ttf",
        "/Library/Fonts/Arial Unicode.ttf",
    ]
    for path in candidates:
        try:
            return ImageFont.truetype(path, size)
        except OSError:
            continue
    return ImageFont.load_default()


def draw_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    lines: list[str],
    fill: str,
    outline: str = "#2C3E50",
) -> None:
    x1, y1, x2, y2 = box
    draw.rounded_rectangle((x1, y1, x2, y2), radius=18, fill=fill, outline=outline, width=4)
    title_font = font(28, bold=True)
    body_font = font(24)

    title_bbox = draw.textbbox((0, 0), title, font=title_font)
    draw.text((x1 + (x2 - x1 - title_bbox[2] + title_bbox[0]) / 2, y1 + 18), title, fill="#17202A", font=title_font)

    y = y1 + 60
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=body_font)
        draw.text((x1 + (x2 - x1 - bbox[2] + bbox[0]) / 2, y), line, fill="#1B2631", font=body_font)
        y += 33


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int]) -> None:
    draw.line((start, end), fill="#2C3E50", width=4)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) > abs(ey - sy):
        if ex >= sx:
            pts = [(ex, ey), (ex - 22, ey - 12), (ex - 22, ey + 12)]
        else:
            pts = [(ex, ey), (ex + 22, ey - 12), (ex + 22, ey + 12)]
    else:
        if ey >= sy:
            pts = [(ex, ey), (ex - 12, ey - 22), (ex + 12, ey - 22)]
        else:
            pts = [(ex, ey), (ex - 12, ey + 22), (ex + 12, ey + 22)]
    draw.polygon(pts, fill="#2C3E50")


def path_arrow(draw: ImageDraw.ImageDraw, points: list[tuple[int, int]]) -> None:
    for start, end in zip(points, points[1:]):
        draw.line((start, end), fill="#2C3E50", width=4)
    start, end = points[-2], points[-1]
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) > abs(ey - sy):
        pts = [(ex, ey), (ex - 22, ey - 12), (ex - 22, ey + 12)] if ex >= sx else [(ex, ey), (ex + 22, ey - 12), (ex + 22, ey + 12)]
    else:
        pts = [(ex, ey), (ex - 12, ey - 22), (ex + 12, ey - 22)] if ey >= sy else [(ex, ey), (ex - 12, ey + 22), (ex + 12, ey + 22)]
    draw.polygon(pts, fill="#2C3E50")


def create_figure() -> None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (3200, 2250), "white")
    draw = ImageDraw.Draw(img)

    title = "Quy trình triển khai mô hình dự báo cluster-aware trên dữ liệu M5"
    title_font = font(42, bold=True)
    title_bbox = draw.textbbox((0, 0), title, font=title_font)
    draw.text(((3200 - title_bbox[2] + title_bbox[0]) / 2, 45), title, fill="#17202A", font=title_font)

    boxes = {
        "data": ((90, 150, 600, 350), "1. Dữ liệu đầu vào", ["Sales theo ngày", "Calendar, SNAP, event", "Sell price, hierarchy"], "#EAF2F8"),
        "eda": ((760, 150, 1330, 350), "2. EDA dữ liệu", ["Zero-sales ratio", "ADI, CV²", "Category, store, state", "Demand regime"], "#EAF2F8"),
        "origin": ((1490, 150, 2060, 350), "3. Rolling origin", ["T = 1885, 1892, 1899", "1906, 1913", "Chỉ dùng lịch sử đến T"], "#EAF2F8"),
        "leakage": ((2220, 150, 3110, 350), "4. Kiểm soát leakage", ["Lag/rolling shift trước", "Scaler fit trên train fold", "Không dùng actual test horizon"], "#FDEDEC"),
        "forecast_feat": ((180, 560, 980, 830), "5A. Feature dự báo", ["Lag 7/14/28/56", "Rolling mean 7/28/56", "Calendar, SNAP, event", "Price, availability, hierarchy"], "#F4F6F7"),
        "cluster_feat": ((1320, 560, 2120, 830), "5B. Feature phân cụm", ["Mean sales, total sales", "ADI, CV², zero ratio", "Gap, spike, event lift", "Price statistics"], "#F4F6F7"),
        "prep": ((2320, 560, 3060, 830), "6. Tiền xử lý cụm", ["Log transform", "Clipping ngoại lệ", "RobustScaler", "Fit theo từng origin"], "#F4F6F7"),
        "kmeans": ((2320, 940, 3060, 1135), "7. Mini-batch K-Means", ["K = 3", "Seed: 42, 52, 62, 72, 82", "Sinh cluster label"], "#FCF3CF"),
        "a0": ((190, 1280, 890, 1500), "8A. Model A0", ["Global LightGBM", "Không dùng cluster", "Baseline chính"], "#E8F6F3"),
        "b1": ((1250, 1280, 1950, 1500), "8B. Model B1", ["Global LightGBM", "Thêm cluster label", "Kiểm tra vai trò nhãn cụm"], "#E8F6F3"),
        "c": ((2310, 1280, 3010, 1500), "8C. Model C", ["Tách dữ liệu theo cụm", "Một LightGBM cho mỗi cụm", "Cluster-specific learning"], "#E8F6F3"),
        "recursive": ((620, 1660, 2580, 1825), "9. Dự báo đệ quy 28 ngày", ["Dự báo h+1 đưa vào history để tạo feature cho h+2", "Không teacher forcing, không dùng actual tương lai"], "#EBF5FB"),
        "eval": ((330, 1990, 2870, 2165), "10. Đánh giá và kiểm định", ["WRMSSE, WAPE, Bias, Stability loss, JumpRate", "Train-test gap, ablation, K-Medoids robustness", "DM test, Friedman/Nemenyi, multi-seed robustness"], "#FEF5E7"),
    }

    for key, (coords, title_text, lines, fill) in boxes.items():
        draw_box(draw, coords, title_text, lines, fill)

    def mid_bottom(key: str) -> tuple[int, int]:
        x1, y1, x2, y2 = boxes[key][0]
        return ((x1 + x2) // 2, y2)

    def mid_top(key: str) -> tuple[int, int]:
        x1, y1, x2, y2 = boxes[key][0]
        return ((x1 + x2) // 2, y1)

    def mid_left(key: str) -> tuple[int, int]:
        x1, y1, x2, y2 = boxes[key][0]
        return (x1, (y1 + y2) // 2)

    def mid_right(key: str) -> tuple[int, int]:
        x1, y1, x2, y2 = boxes[key][0]
        return (x2, (y1 + y2) // 2)

    arrow(draw, mid_right("data"), mid_left("eda"))
    arrow(draw, mid_right("eda"), mid_left("origin"))
    arrow(draw, mid_right("origin"), mid_left("leakage"))
    path_arrow(draw, [mid_bottom("origin"), (1775, 455), mid_top("forecast_feat")])
    path_arrow(draw, [mid_bottom("origin"), (1775, 455), mid_top("cluster_feat")])
    arrow(draw, mid_right("cluster_feat"), mid_left("prep"))
    arrow(draw, mid_bottom("prep"), mid_top("kmeans"))
    path_arrow(draw, [mid_bottom("forecast_feat"), (580, 1180), mid_top("a0")])
    path_arrow(draw, [mid_bottom("forecast_feat"), (580, 1180), (1600, 1180), mid_top("b1")])
    path_arrow(draw, [mid_bottom("forecast_feat"), (580, 1180), (2660, 1180), mid_top("c")])
    path_arrow(draw, [mid_bottom("kmeans"), (2690, 1230), mid_top("c")])
    path_arrow(draw, [mid_bottom("kmeans"), (2690, 1230), (1600, 1230), mid_top("b1")])
    arrow(draw, mid_bottom("a0"), (1080, 1660))
    arrow(draw, mid_bottom("b1"), (1600, 1660))
    arrow(draw, mid_bottom("c"), (2120, 1660))
    arrow(draw, mid_bottom("recursive"), mid_top("eval"))

    note_font = font(22)
    draw.text(
        (90, 2200),
        "Ghi chú: Calendar/SNAP/event và price được xem là biến ngoại sinh biết trước; mọi feature từ actual sales chỉ dùng dữ liệu lịch sử tại origin T.",
        fill="#566573",
        font=note_font,
    )

    img.save(FIG_PATH)


def find_heading(doc: Document, text: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading") and paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Heading not found: {text}")


def replace_figure_in_doc() -> None:
    doc = Document(DOC_PATH)
    heading = find_heading(doc, "Mô hình triển khai")
    cursor = heading._p.getnext()
    insertion_anchor = heading
    while cursor is not None:
        next_cursor = cursor.getnext()
        if cursor.tag.split("}")[-1] == "p":
            p = Paragraph(cursor, doc)
            if p.style.name.startswith("Heading"):
                break
            if cursor.xpath(".//*[local-name()='drawing']") or p.text.strip().startswith("Hình 2."):
                cursor.getparent().remove(cursor)
            else:
                insertion_anchor = p
        cursor = next_cursor

    img_p = heading._parent.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p.add_run().add_picture(str(FIG_PATH), width=Inches(6.8))
    heading._p.addnext(img_p._p)

    caption = heading._parent.add_paragraph(
        "Hình 2. Quy trình triển khai mô hình dự báo cluster-aware trên dữ liệu M5.",
        style="Caption",
    )
    img_p._p.addnext(caption._p)

    doc.save(DOC_PATH)
    OUTPUT_COPY.write_bytes(DOC_PATH.read_bytes())


def main() -> None:
    create_figure()
    replace_figure_in_doc()
    doc = Document(DOC_PATH)
    print(f"Updated figure: {FIG_PATH}")
    print(f"Updated document: {DOC_PATH}")
    print(f"Copy: {OUTPUT_COPY}")
    print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} inline_shapes={len(doc.inline_shapes)}")


if __name__ == "__main__":
    main()
