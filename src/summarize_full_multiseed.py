from __future__ import annotations

import json
import os
from itertools import combinations
from pathlib import Path

_cache_root = Path("outputs_full_multiseed_summary") / ".cache"
_cache_root.mkdir(parents=True, exist_ok=True)
os.environ.setdefault("MPLCONFIGDIR", str((_cache_root / "matplotlib").resolve()))
os.environ.setdefault("XDG_CACHE_HOME", str((_cache_root / "xdg").resolve()))

import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches
from scipy import stats


SEEDS = [42, 52, 62, 72, 82]
MODELS = ["A0_global_baseline", "B1_cluster_label", "C_cluster_specific"]
MODEL_LABEL = {
    "A0_global_baseline": "A0",
    "B1_cluster_label": "B1",
    "C_cluster_specific": "C",
}
CONFIG_TEMPLATE = "configs/research_config_full_multiseed_k3_seed{seed}_tweedie_a0_b1_c.json"
OUT = Path("outputs_full_multiseed_summary")
METRICS = OUT / "metrics"
FIGS = OUT / "figures"
REPORTS = OUT / "reports"


def fmt(x: object, digits: int = 6) -> str:
    if pd.isna(x):
        return "NA"
    if isinstance(x, (float, np.floating)):
        return f"{float(x):.{digits}f}"
    return str(x)


def md_table(df: pd.DataFrame) -> str:
    if df.empty:
        return "(empty)"
    cols = list(df.columns)
    lines = ["| " + " | ".join(map(str, cols)) + " |", "| " + " | ".join(["---"] * len(cols)) + " |"]
    for _, row in df.iterrows():
        lines.append("| " + " | ".join(fmt(row[c]) for c in cols) + " |")
    return "\n".join(lines)


def add_table(doc: Document, df: pd.DataFrame, title: str) -> None:
    doc.add_paragraph(title)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            cells[i].text = fmt(row[col])


def load_seed(seed: int) -> dict[str, pd.DataFrame | dict]:
    config_path = Path(CONFIG_TEMPLATE.format(seed=seed))
    config = json.loads(config_path.read_text(encoding="utf-8"))
    metrics_dir = Path(config["outputs_dir"]) / "metrics"
    return {
        "config": config,
        "model": pd.read_csv(metrics_dir / "model_test_metrics.csv"),
        "wrmsse": pd.read_csv(metrics_dir / "sampled_wrmsse_overall.csv"),
        "stability": pd.read_csv(metrics_dir / "forecast_stability_metrics.csv"),
        "overfit": pd.read_csv(metrics_dir / "overfitting_gap_summary.csv"),
        "feature_importance": pd.read_csv(metrics_dir / "feature_importance_by_origin.csv"),
        "clustering": pd.read_csv(metrics_dir / "clustering_quality_metrics.csv"),
        "runtime": pd.read_csv(metrics_dir / "runtime_costs.csv"),
    }


def per_origin(seed: int, run: dict[str, pd.DataFrame | dict]) -> pd.DataFrame:
    model = run["model"]  # type: ignore[assignment]
    wr = run["wrmsse"]  # type: ignore[assignment]
    gap = run["overfit"]  # type: ignore[assignment]
    out = model.loc[model["model"].isin(MODELS)].merge(
        wr.loc[wr["model"].isin(MODELS)], on=["origin", "model"], how="left"
    )
    out = out.merge(
        gap.loc[gap["model"].isin(MODELS), ["origin", "model", "test_train_gap"]],
        on=["origin", "model"],
        how="left",
    )
    out.insert(0, "seed", seed)
    return out


def per_seed(seed: int, run: dict[str, pd.DataFrame | dict]) -> pd.DataFrame:
    po = per_origin(seed, run)
    st = run["stability"]  # type: ignore[assignment]
    st_summary = (
        st.loc[st["model"].isin(MODELS)]
        .groupby("model")[["scale_aware_stability_loss", "jump_rate_tau_0.3", "jump_rate_tau_0.5"]]
        .mean()
        .reset_index()
    )
    out = (
        po.groupby("model")
        .agg(
            rmsse_item_store=("rmsse_item_store", "mean"),
            sampled_wrmsse_12level=("sampled_wrmsse_12level", "mean"),
            wape=("wape", "mean"),
            bias=("bias", "mean"),
            train_test_gap=("test_train_gap", "mean"),
        )
        .reset_index()
        .merge(st_summary, on="model", how="left")
    )
    out.insert(0, "seed", seed)
    return out


def ci95(values: pd.Series) -> tuple[float, float, float]:
    x = values.dropna().astype(float).to_numpy()
    mean = float(np.mean(x))
    if len(x) <= 1:
        return mean, np.nan, np.nan
    se = stats.sem(x)
    lo, hi = stats.t.interval(0.95, len(x) - 1, loc=mean, scale=se)
    return mean, float(lo), float(hi)


def aggregate(per_seed_df: pd.DataFrame) -> pd.DataFrame:
    metrics = [
        "rmsse_item_store",
        "sampled_wrmsse_12level",
        "scale_aware_stability_loss",
        "jump_rate_tau_0.3",
        "jump_rate_tau_0.5",
        "wape",
        "bias",
        "train_test_gap",
    ]
    rows = []
    for model, grp in per_seed_df.groupby("model"):
        row = {"model": model, "n_seeds": grp["seed"].nunique()}
        for metric in metrics:
            mean, lo, hi = ci95(grp[metric])
            row[f"{metric}_mean"] = mean
            row[f"{metric}_std"] = float(grp[metric].std())
            row[f"{metric}_min"] = float(grp[metric].min())
            row[f"{metric}_max"] = float(grp[metric].max())
            row[f"{metric}_ci95_low"] = lo
            row[f"{metric}_ci95_high"] = hi
        rows.append(row)
    return pd.DataFrame(rows).sort_values("sampled_wrmsse_12level_mean")


def paired_tests(per_seed_df: pd.DataFrame) -> pd.DataFrame:
    metrics = [
        "rmsse_item_store",
        "sampled_wrmsse_12level",
        "scale_aware_stability_loss",
        "jump_rate_tau_0.3",
        "jump_rate_tau_0.5",
        "wape",
        "bias",
        "train_test_gap",
    ]
    rows = []
    for metric in metrics:
        piv = per_seed_df.pivot(index="seed", columns="model", values=metric).dropna()
        if {"A0_global_baseline", "C_cluster_specific"}.issubset(piv.columns):
            diff = piv["A0_global_baseline"] - piv["C_cluster_specific"]
            t_stat, t_p = stats.ttest_rel(piv["A0_global_baseline"], piv["C_cluster_specific"])
            try:
                w_stat, w_p = stats.wilcoxon(piv["A0_global_baseline"], piv["C_cluster_specific"])
            except ValueError:
                w_stat, w_p = np.nan, np.nan
            rows.append(
                {
                    "metric": metric,
                    "n_seeds": len(piv),
                    "mean_diff_A0_minus_C": float(diff.mean()),
                    "median_diff_A0_minus_C": float(diff.median()),
                    "paired_t_stat": float(t_stat),
                    "paired_t_p_value": float(t_p),
                    "wilcoxon_stat": float(w_stat) if np.isfinite(w_stat) else np.nan,
                    "wilcoxon_p_value": float(w_p) if np.isfinite(w_p) else np.nan,
                }
            )
    return pd.DataFrame(rows)


def friedman_nemenyi(per_seed_df: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame]:
    metrics = [
        "rmsse_item_store",
        "sampled_wrmsse_12level",
        "scale_aware_stability_loss",
        "jump_rate_tau_0.3",
        "jump_rate_tau_0.5",
        "wape",
        "train_test_gap",
    ]
    friedman_rows = []
    nemenyi_rows = []
    q_alpha = 2.343
    for metric in metrics:
        piv = per_seed_df.pivot(index="seed", columns="model", values=metric).dropna(axis=1)
        piv = piv[[m for m in MODELS if m in piv.columns]].dropna()
        stat, pval = stats.friedmanchisquare(*[piv[m].to_numpy() for m in piv.columns])
        ranks = piv.rank(axis=1, ascending=True, method="average")
        avg_ranks = ranks.mean()
        n_blocks = ranks.shape[0]
        n_models = ranks.shape[1]
        cd = q_alpha * np.sqrt(n_models * (n_models + 1) / (6.0 * n_blocks))
        friedman_rows.append({"metric": metric, "n_seeds": n_blocks, "statistic": float(stat), "p_value": float(pval)})
        for a, b in combinations(piv.columns, 2):
            diff = abs(float(avg_ranks[a] - avg_ranks[b]))
            nemenyi_rows.append(
                {
                    "metric": metric,
                    "model_a": a,
                    "model_b": b,
                    "average_rank_a": float(avg_ranks[a]),
                    "average_rank_b": float(avg_ranks[b]),
                    "rank_diff": diff,
                    "critical_difference": cd,
                    "significant_alpha_0.05": diff > cd,
                }
            )
    return pd.DataFrame(friedman_rows), pd.DataFrame(nemenyi_rows)


def feature_summary(runs: dict[int, dict[str, pd.DataFrame | dict]]) -> pd.DataFrame:
    frames = []
    for seed, run in runs.items():
        fi = run["feature_importance"]  # type: ignore[assignment]
        fi = fi.loc[fi["model"].isin(MODELS)].copy()
        fi["seed"] = seed
        frames.append(fi)
    all_fi = pd.concat(frames, ignore_index=True)
    return (
        all_fi.groupby(["model", "feature"])
        .agg(mean_gain=("importance_gain", "mean"), std_gain=("importance_gain", "std"), mean_split=("importance_split", "mean"))
        .reset_index()
        .sort_values(["model", "mean_gain"], ascending=[True, False])
    )


def _svg_text(x: float, y: float, text: str, size: int = 12, anchor: str = "middle") -> str:
    return f'<text x="{x:.1f}" y="{y:.1f}" font-size="{size}" text-anchor="{anchor}" font-family="Arial">{text}</text>'


def _metric_scale(values: list[float]) -> tuple[float, float]:
    lo = min(values)
    hi = max(values)
    pad = (hi - lo) * 0.12 if hi > lo else max(abs(hi) * 0.1, 0.01)
    return lo - pad, hi + pad


def _save_svg(path: Path, body: str, width: int = 760, height: int = 430) -> None:
    path.write_text(
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}">'
        '<rect width="100%" height="100%" fill="white"/>'
        f"{body}</svg>",
        encoding="utf-8",
    )


def plot_outputs(per_seed_df: pd.DataFrame, aggregate_df: pd.DataFrame) -> None:
    metric_labels = {
        "sampled_wrmsse_12level": "WRMSSE 12-level",
        "scale_aware_stability_loss": "Stability loss",
        "jump_rate_tau_0.3": "JumpRate@0.3",
        "jump_rate_tau_0.5": "JumpRate@0.5",
        "wape": "WAPE",
        "train_test_gap": "Train-test gap",
    }
    for metric, label in metric_labels.items():
        data = {model: per_seed_df.loc[per_seed_df["model"].eq(model), metric].astype(float).to_numpy() for model in MODELS}
        all_values = [float(v) for arr in data.values() for v in arr]
        ymin, ymax = _metric_scale(all_values)
        width, height = 760, 430
        left, right, top, bottom = 70, 30, 50, 70
        plot_h = height - top - bottom

        def ymap(v: float) -> float:
            return top + (ymax - v) / (ymax - ymin) * plot_h

        body = [
            _svg_text(width / 2, 25, f"Full multi-seed {label}", 16),
            f'<line x1="{left}" y1="{top}" x2="{left}" y2="{height-bottom}" stroke="#222"/>',
            f'<line x1="{left}" y1="{height-bottom}" x2="{width-right}" y2="{height-bottom}" stroke="#222"/>',
            _svg_text(10, top + 5, fmt(ymax, 4), 10, "start"),
            _svg_text(10, height - bottom, fmt(ymin, 4), 10, "start"),
        ]
        x_positions = np.linspace(left + 110, width - right - 110, len(MODELS))
        for x, model in zip(x_positions, MODELS):
            arr = np.sort(data[model])
            q1, med, q3 = np.percentile(arr, [25, 50, 75])
            vmin, vmax = float(arr.min()), float(arr.max())
            box_w = 54
            body.extend(
                [
                    f'<line x1="{x:.1f}" y1="{ymap(vmin):.1f}" x2="{x:.1f}" y2="{ymap(vmax):.1f}" stroke="#1565c0" stroke-width="2"/>',
                    f'<line x1="{x-box_w/3:.1f}" y1="{ymap(vmin):.1f}" x2="{x+box_w/3:.1f}" y2="{ymap(vmin):.1f}" stroke="#1565c0" stroke-width="2"/>',
                    f'<line x1="{x-box_w/3:.1f}" y1="{ymap(vmax):.1f}" x2="{x+box_w/3:.1f}" y2="{ymap(vmax):.1f}" stroke="#1565c0" stroke-width="2"/>',
                    f'<rect x="{x-box_w/2:.1f}" y="{ymap(q3):.1f}" width="{box_w}" height="{max(1, ymap(q1)-ymap(q3)):.1f}" fill="#bbdefb" stroke="#1565c0" stroke-width="2"/>',
                    f'<line x1="{x-box_w/2:.1f}" y1="{ymap(med):.1f}" x2="{x+box_w/2:.1f}" y2="{ymap(med):.1f}" stroke="#0d47a1" stroke-width="3"/>',
                    _svg_text(x, height - 35, MODEL_LABEL[model], 12),
                ]
            )
        _save_svg(FIGS / f"boxplot_{metric}.svg", "\n".join(body), width, height)

        agg = aggregate_df.set_index("model")
        means = [float(agg.at[m, f"{metric}_mean"]) for m in MODELS]
        ci_lows = [float(agg.at[m, f"{metric}_ci95_low"]) for m in MODELS]
        ci_highs = [float(agg.at[m, f"{metric}_ci95_high"]) for m in MODELS]
        ymin, ymax = _metric_scale(ci_lows + ci_highs + means)
        body = [
            _svg_text(width / 2, 25, f"Mean and 95% CI: {label}", 16),
            f'<line x1="{left}" y1="{top}" x2="{left}" y2="{height-bottom}" stroke="#222"/>',
            f'<line x1="{left}" y1="{height-bottom}" x2="{width-right}" y2="{height-bottom}" stroke="#222"/>',
            _svg_text(10, top + 5, fmt(ymax, 4), 10, "start"),
            _svg_text(10, height - bottom, fmt(ymin, 4), 10, "start"),
        ]
        for x, model, mean, lo, hi in zip(x_positions, MODELS, means, ci_lows, ci_highs):
            body.extend(
                [
                    f'<line x1="{x:.1f}" y1="{ymap(lo):.1f}" x2="{x:.1f}" y2="{ymap(hi):.1f}" stroke="#2e7d32" stroke-width="2"/>',
                    f'<line x1="{x-18:.1f}" y1="{ymap(lo):.1f}" x2="{x+18:.1f}" y2="{ymap(lo):.1f}" stroke="#2e7d32" stroke-width="2"/>',
                    f'<line x1="{x-18:.1f}" y1="{ymap(hi):.1f}" x2="{x+18:.1f}" y2="{ymap(hi):.1f}" stroke="#2e7d32" stroke-width="2"/>',
                    f'<circle cx="{x:.1f}" cy="{ymap(mean):.1f}" r="6" fill="#1b5e20"/>',
                    _svg_text(x, height - 35, MODEL_LABEL[model], 12),
                ]
            )
        _save_svg(FIGS / f"errorbar_{metric}.svg", "\n".join(body), width, height)


def main() -> None:
    METRICS.mkdir(parents=True, exist_ok=True)
    FIGS.mkdir(parents=True, exist_ok=True)
    REPORTS.mkdir(parents=True, exist_ok=True)
    runs = {seed: load_seed(seed) for seed in SEEDS}
    per_origin_df = pd.concat([per_origin(seed, run) for seed, run in runs.items()], ignore_index=True)
    per_seed_df = pd.concat([per_seed(seed, run) for seed, run in runs.items()], ignore_index=True)
    aggregate_df = aggregate(per_seed_df)
    paired_df = paired_tests(per_seed_df)
    friedman_df, nemenyi_df = friedman_nemenyi(per_seed_df)
    feature_df = feature_summary(runs)

    per_origin_df.to_csv(METRICS / "full_multiseed_per_origin_metrics.csv", index=False)
    per_seed_df.to_csv(METRICS / "full_multiseed_per_seed_metrics.csv", index=False)
    aggregate_df.to_csv(METRICS / "full_multiseed_aggregate_metrics.csv", index=False)
    paired_df.to_csv(METRICS / "full_multiseed_paired_tests_a0_vs_c.csv", index=False)
    friedman_df.to_csv(METRICS / "full_multiseed_friedman_tests.csv", index=False)
    nemenyi_df.to_csv(METRICS / "full_multiseed_nemenyi_posthoc.csv", index=False)
    feature_df.to_csv(METRICS / "full_multiseed_feature_importance_summary.csv", index=False)
    plot_outputs(per_seed_df, aggregate_df)

    top_feature = feature_df.groupby("model").head(10).reset_index(drop=True)
    lines = [
        "# Full Multi-Seed Assessment",
        "",
        "Seeds: 42, 52, 62, 72, 82. Rolling origins: 1885, 1892, 1899, 1906, 1913.",
        "All runs use the leakage-aware recursive forecasting pipeline with A0, B1, and C.",
        "",
        "## Per-Seed Metrics",
        "",
        md_table(per_seed_df),
        "",
        "## Aggregate Metrics",
        "",
        md_table(aggregate_df),
        "",
        "## Paired Tests A0 vs C",
        "",
        md_table(paired_df),
        "",
        "## Friedman Test",
        "",
        md_table(friedman_df),
        "",
        "## Nemenyi Post-Hoc",
        "",
        md_table(nemenyi_df),
        "",
        "## Top Feature Importance",
        "",
        md_table(top_feature),
    ]
    (REPORTS / "full_multiseed_assessment.md").write_text("\n".join(lines), encoding="utf-8")

    doc = Document()
    doc.add_heading("Full Multi-Seed Assessment", 0)
    doc.add_paragraph("Seeds: 42, 52, 62, 72, 82. Rolling origins: 1885, 1892, 1899, 1906, 1913.")
    doc.add_paragraph("Pipeline preserves historical-only clustering/features and recursive 28-day forecasting.")
    add_table(doc, per_seed_df, "Per-seed metrics")
    add_table(doc, aggregate_df, "Aggregate metrics with 95% CI")
    add_table(doc, paired_df, "Paired tests A0 vs C")
    add_table(doc, friedman_df, "Friedman tests")
    add_table(doc, nemenyi_df, "Nemenyi post-hoc")
    add_table(doc, top_feature, "Top feature importance")
    for section in doc.sections:
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
    doc.save(REPORTS / "full_multiseed_assessment.docx")
    print(METRICS / "full_multiseed_aggregate_metrics.csv")
    print(REPORTS / "full_multiseed_assessment.md")
    print(REPORTS / "full_multiseed_assessment.docx")


if __name__ == "__main__":
    main()
