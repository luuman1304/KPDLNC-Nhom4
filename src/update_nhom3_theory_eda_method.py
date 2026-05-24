from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOC_PATH = ROOT / "Document" / "Bai_Bao_Nghien_Cuu_M5_Cau_Truc_Chuan_Hoc_Thuat.docx"
OUTPUT_COPY = ROOT / "Document" / "Bai_Bao_Nghien_Cuu_M5_Nhom3_Format.docx"
FIG_DIR = ROOT / "Document" / "figures"
METHOD_FIG = FIG_DIR / "method_flow_updated.png"


INTRO_APPEND = [
    (
        "Lý do chọn đề tài xuất phát từ yêu cầu vận hành của dự báo bán lẻ hiện đại: "
        "mô hình không chỉ cần giảm sai số trung bình mà còn phải duy trì độ ổn định "
        "khi kế hoạch bổ sung hàng được cập nhật định kỳ. Tổng quan về retail "
        "forecasting cho thấy sai số dự báo tác động trực tiếp đến quyết định tồn kho, "
        "phân bổ hàng hóa và service level (Fildes et al., 2019). Dữ liệu M5 cung cấp "
        "một bối cảnh phù hợp để nghiên cứu vấn đề này vì gồm 30,490 chuỗi item-store, "
        "10 cửa hàng, 3 bang, 3 ngành hàng và hệ thống phân cấp đánh giá bằng WRMSSE "
        "(Makridakis et al., 2022)."
    ),
    (
        "Trong các nghiên cứu về M5, global learning và gradient boosting thường đạt "
        "hiệu quả tốt nhờ khai thác thông tin liên chuỗi, nhưng dữ liệu bán lẻ vẫn có "
        "độ dị thể lớn giữa sản phẩm bán đều, sản phẩm bán thưa và sản phẩm có nhu cầu "
        "lumpy. Ma và Fildes (2022) chỉ ra rằng các tiếp cận global bottom-up cần được "
        "kiểm tra độ bền trong bối cảnh M5, còn Syntetos et al. (2005) cho thấy demand "
        "pattern classification là nền tảng quan trọng khi xử lý intermittent demand. "
        "Vì vậy, đề tài chọn hướng cluster-aware global forecasting để kiểm tra liệu "
        "phân cụm theo demand behavior có giúp cải thiện accuracy và stability so với "
        "một global LightGBM duy nhất hay không."
    ),
]


THEORY_BLOCK: list[tuple[str, str]] = [
    (
        "3.1.1. Logic nghiệp vụ của dự báo nhu cầu bán lẻ",
        "Dự báo nhu cầu bán lẻ là đầu vào cho tồn kho, bổ sung hàng, phân bổ hàng hóa "
        "và kế hoạch cung ứng. Một mô hình có WRMSSE thấp nhưng dự báo dao động mạnh "
        "giữa các kỳ cập nhật vẫn có thể gây khó khăn vận hành vì planner phải liên tục "
        "điều chỉnh kế hoạch. Do đó nghiên cứu đánh giá đồng thời độ chính xác, độ ổn "
        "định và mức overfitting thay vì chỉ tối ưu một metric sai số."
    ),
    (
        "3.1.2. Cấu trúc phân cấp của dữ liệu bán lẻ",
        "M5 là dữ liệu phân cấp: chuỗi item-store ở cấp thấp được tổng hợp lên store, "
        "state, category và department. Business logic của bài toán đòi hỏi dự báo tốt "
        "ở nhiều cấp, vì sai số ở cấp sản phẩm ảnh hưởng tồn kho chi tiết, còn sai số ở "
        "cấp tổng hợp ảnh hưởng phân bổ nguồn lực theo cửa hàng, bang và ngành hàng."
    ),
    (
        "3.1.3. Intermittent demand, ADI và CV2",
        "Intermittent demand xuất hiện khi chuỗi có nhiều ngày không bán, làm các chỉ số "
        "phần trăm dễ bị phóng đại nếu mẫu số nhỏ. ADI đo khoảng cách trung bình giữa "
        "các lần bán khác 0; CV2 đo mức biến động tương đối của lượng bán trong các ngày "
        "có nhu cầu. Hai chỉ số này giúp diễn giải các demand regimes như smooth, "
        "intermittent, erratic và lumpy."
    ),
    (
        "3.1.4. Global forecasting và cluster-aware learning",
        "Global forecasting huấn luyện một mô hình chung trên nhiều chuỗi để tận dụng "
        "cross-learning. Hạn chế của cách tiếp cận này là các demand regimes dị biệt "
        "phải đi qua cùng một hàm học. Cluster-aware learning giảm dị thể bằng cách "
        "nhóm chuỗi có hành vi gần nhau. Trong nghiên cứu, A0 là global baseline, B1 "
        "thêm cluster label vào global model, còn C huấn luyện LightGBM riêng cho từng "
        "cụm."
    ),
    (
        "3.1.5. LightGBM Tweedie cho nhu cầu bán lẻ",
        "LightGBM phù hợp với dữ liệu tabular lớn và có nhiều feature phân loại, lag, "
        "rolling, calendar, price và hierarchy. Objective Tweedie được chọn vì nhu cầu "
        "bán lẻ là biến không âm, có nhiều giá trị 0 và phần dương lệch phải. Thiết kế "
        "này phù hợp hơn squared error thuần túy trong bối cảnh zero-inflated demand."
    ),
    (
        "3.1.6. Rolling-origin và kiểm soát data leakage",
        "Rolling-origin evaluation mô phỏng việc huấn luyện lại mô hình khi thời gian "
        "dịch chuyển. Tại origin T, lag, rolling statistics, ADI, CV2, mean sales, "
        "zero-sales ratio và scaler chỉ được tính từ dữ liệu lịch sử đến T. Forecast "
        "28 ngày được sinh đệ quy, nghĩa là dự báo của bước trước được dùng để tạo "
        "feature cho bước sau, không dùng actual value trong test horizon."
    ),
    (
        "3.1.7. Tiêu chí đánh giá và kiểm định",
        "WRMSSE đo sai số theo cấu trúc phân cấp và trọng số doanh thu; WAPE và bias hỗ "
        "trợ diễn giải sai số tổng thể; scale-aware stability đo mức dao động forecast "
        "giữa các origins nhưng chặn mẫu số để tránh phóng đại near-zero demand. Train-test "
        "gap, ablation, Diebold-Mariano test, Friedman/Nemenyi và multi-seed robustness "
        "được dùng để kiểm tra độ tin cậy của kết luận."
    ),
]


EDA_PARAGRAPHS = [
    (
        "Mục này trình bày EDA dữ liệu M5 trước khi xây dựng mô hình. Dữ liệu được lấy "
        "từ bộ M5 Forecasting Accuracy, gồm bảng sales theo ngày, lịch, sự kiện, SNAP "
        "và giá bán theo item-store-week. Bài nghiên cứu sử dụng dữ liệu evaluation có "
        "actual sales đến d_1941 để đánh giá rolling-origin sau d_1913."
    ),
    (
        "Ở cấp chuỗi, dữ liệu gồm 30,490 item-store series, 3,049 sản phẩm, 7 department, "
        "3 category, 10 store và 3 state. Đây là cấu trúc phù hợp với mục tiêu nghiên "
        "cứu vì vừa có cross-sectional scale đủ lớn cho global learning, vừa có nhiều "
        "nguồn dị thể để kiểm tra cluster-aware learning."
    ),
    (
        "EDA cho thấy nhu cầu bán lẻ trong M5 có độ thưa cao. Median zero-sales ratio "
        "là 0.6337, median ADI là 2.7300 và median CV2 là 0.3485. Phân loại demand "
        "regime theo ADI-CV2 cho thấy intermittent demand chiếm phần lớn, tiếp theo là "
        "lumpy demand; đây là lý do bài nghiên cứu dùng scale-aware stability thay vì "
        "chỉ dựa trên phần trăm thay đổi forecast."
    ),
    (
        "Các biến calendar, SNAP, event và price được xem là biến ngoại sinh có thể "
        "biết trước trong bối cảnh M5. Ngược lại, mọi đặc trưng tạo từ actual sales như "
        "lag, rolling mean, ADI, CV2 và zero-sales ratio chỉ được tính từ dữ liệu lịch "
        "sử tại từng rolling origin để tránh data leakage."
    ),
]


EDA_TABLE = [
    ("Số chuỗi item-store", "30,490", "Quy mô panel dùng cho global learning"),
    ("Sản phẩm", "3,049", "Định danh item trong hierarchy"),
    ("Cửa hàng", "10", "Bối cảnh bán hàng theo store"),
    ("Bang", "3", "CA, TX, WI; liên quan SNAP và vùng"),
    ("Category", "3", "FOODS, HOBBIES, HOUSEHOLD"),
    ("Department", "7", "Cấp phân cấp giữa item và category"),
    ("Số ngày lịch", "1,969", "Calendar, event, SNAP, wm_yr_wk"),
    ("Median zero-sales ratio", "0.6337", "Mức độ thưa của nhu cầu"),
    ("Median ADI", "2.7300", "Khoảng cách trung bình giữa các lần bán"),
    ("Median CV2", "0.3485", "Biến động lượng bán khi có nhu cầu"),
]


METHOD_DESCRIPTION = [
    (
        "Hình 2 mô tả pipeline đã dùng trong thực nghiệm sau các cập nhật kiểm soát "
        "leakage. Luồng xử lý bắt đầu từ dữ liệu sales, calendar, price và hierarchy, "
        "sau đó được chia thành hai nhánh feature song song tại từng rolling origin T."
    ),
    (
        "Nhánh forecasting tạo lag/rolling features, calendar/SNAP/event, price, "
        "availability và hierarchy features cho từng ngày dự báo. Nhánh clustering chỉ "
        "dùng dữ liệu lịch sử đến T để tính ADI, CV2, mean sales, zero-sales ratio, gap, "
        "event lift và price statistics; các biến này được log-transform, clipping và "
        "fit RobustScaler trên train fold trước khi chạy Mini-batch K-Means K=3."
    ),
    (
        "Từ cùng một feature pipeline, ba nhánh mô hình được huấn luyện để tách riêng "
        "hiệu ứng của cluster information. A0 là global LightGBM không dùng cụm. B1 là "
        "global LightGBM nhưng thêm cluster label như categorical feature. C tách dữ "
        "liệu theo cluster và huấn luyện một LightGBM riêng cho từng cụm. Forecast của "
        "cả ba nhánh được sinh đệ quy trong 28 ngày."
    ),
    (
        "Khối đánh giá cuối cùng gồm WRMSSE close-to-official theo hierarchy, WAPE, "
        "bias, scale-aware stability, JumpRate, train-test gap, ablation, K-Medoids "
        "robustness, DM test và multi-seed robustness. Thiết kế này phản ánh đúng thực "
        "nghiệm đã chạy: trọng tâm là A0, B1 và C, không so sánh theo leaderboard."
    ),
]


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def find_heading(doc: Document, text: str, occurrence: int = 1) -> Paragraph:
    count = 0
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading") and paragraph.text.strip() == text:
            count += 1
            if count == occurrence:
                return paragraph
    raise ValueError(f"Heading not found: {text!r}, occurrence={occurrence}")


def remove_until_next_heading(heading: Paragraph, same_or_higher: bool = False) -> None:
    level = int(heading.style.name.split()[-1]) if heading.style.name.startswith("Heading") else 9
    cursor = heading._p.getnext()
    while cursor is not None:
        next_cursor = cursor.getnext()
        tag = cursor.tag.split("}")[-1]
        stop = False
        if tag == "p":
            p = Paragraph(cursor, heading._parent)
            if p.style.name.startswith("Heading"):
                next_level = int(p.style.name.split()[-1])
                stop = next_level <= level if same_or_higher else True
        if stop:
            break
        cursor.getparent().remove(cursor)
        cursor = next_cursor


def insert_after(anchor: Paragraph, entries: list[tuple[str, str]]) -> Paragraph:
    current = anchor
    for text, style in entries:
        p = anchor._parent.add_paragraph(text, style=style)
        current._p.addnext(p._p)
        current = p
    return current


def insert_table_after(paragraph: Paragraph, rows: list[tuple[str, str, str]]) -> None:
    table = paragraph._parent.add_table(rows=1, cols=3, width=Inches(6.5))
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "Nội dung EDA"
    header[1].text = "Giá trị"
    header[2].text = "Ý nghĩa trong nghiên cứu"
    for item, value, meaning in rows:
        cells = table.add_row().cells
        cells[0].text = item
        cells[1].text = value
        cells[2].text = meaning
    paragraph._p.addnext(table._tbl)


def create_method_figure() -> None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    width, height = 2400, 1600
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    try:
        title_font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 36)
        font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 26)
        small_font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 24)
    except OSError:
        title_font = ImageFont.load_default()
        font = ImageFont.load_default()
        small_font = ImageFont.load_default()

    boxes = {
        "data": (120, 180, 610, 360, "M5 data\nsales + calendar/SNAP/event\nprice + hierarchy", "#EAF2F8"),
        "origin": (895, 190, 1485, 340, "Rolling origin T\nhistorical-only split", "#EAF2F8"),
        "forecast_feat": (120, 590, 760, 860, "Forecasting features\nlag/rolling shifted\ncalendar, price\navailability, hierarchy", "#F4F6F7"),
        "cluster_feat": (850, 590, 1490, 860, "Clustering branch\nADI, CV2, mean sales\nzero ratio, gap, event lift\nlog + clip + RobustScaler", "#F4F6F7"),
        "cluster": (1640, 620, 2180, 830, "Mini-batch K-Means\nK=3, seed-specific\ncluster labels", "#F4F6F7"),
        "a0": (130, 1010, 690, 1230, "A0\nGlobal LightGBM\nno cluster", "#E8F6F3"),
        "b1": (920, 1010, 1480, 1230, "B1\nGlobal LightGBM\n+ cluster label", "#E8F6F3"),
        "c": (1700, 1010, 2260, 1230, "C\nCluster-specific LightGBM\none model per cluster", "#E8F6F3"),
        "recursive": (520, 1320, 1880, 1440, "Recursive 28-day forecasting\nprediction h+1 updates history for h+2; no teacher forcing", "#F4F6F7"),
        "eval": (400, 1490, 2000, 1585, "Evaluation: WRMSSE, WAPE, bias, stability, JumpRate, gap,\nablation, K-Medoids robustness, DM test, multi-seed", "#FEF5E7"),
    }

    def center_text(text: str, box: tuple[int, int, int, int, str, str]) -> None:
        x1, y1, x2, y2, _, _ = box
        lines = text.split("\n")
        line_h = 34
        start_y = y1 + ((y2 - y1) - line_h * len(lines)) / 2
        for idx, line in enumerate(lines):
            bbox = draw.textbbox((0, 0), line, font=font)
            draw.text((x1 + (x2 - x1 - (bbox[2] - bbox[0])) / 2, start_y + idx * line_h), line, fill="#1B2631", font=font)

    title = "Leakage-aware cluster-aware forecasting pipeline"
    title_bbox = draw.textbbox((0, 0), title, font=title_font)
    draw.text(((width - (title_bbox[2] - title_bbox[0])) / 2, 70), title, fill="#1B2631", font=title_font)

    for box in boxes.values():
        x1, y1, x2, y2, label, fill = box
        draw.rounded_rectangle((x1, y1, x2, y2), radius=18, fill=fill, outline="#2C3E50", width=4)
        center_text(label, box)

    def arrow(src: str, dst: str) -> None:
        sx1, sy1, sx2, sy2, *_ = boxes[src]
        dx1, dy1, dx2, dy2, *_ = boxes[dst]
        start = ((sx1 + sx2) // 2, sy2)
        end = ((dx1 + dx2) // 2, dy1)
        draw.line((start, end), fill="#2C3E50", width=4)
        ex, ey = end
        draw.polygon([(ex, ey), (ex - 12, ey - 22), (ex + 12, ey - 22)], fill="#2C3E50")

    arrow("data", "origin")
    arrow("origin", "forecast_feat")
    arrow("origin", "cluster_feat")
    arrow("cluster_feat", "cluster")
    arrow("forecast_feat", "a0")
    arrow("forecast_feat", "b1")
    arrow("forecast_feat", "c")
    arrow("cluster", "b1")
    arrow("cluster", "c")
    arrow("a0", "recursive")
    arrow("b1", "recursive")
    arrow("c", "recursive")
    arrow("recursive", "eval")

    img.save(METHOD_FIG)


def replace_method_figure(doc: Document) -> None:
    heading = find_heading(doc, "Mô hình triển khai")
    cursor = heading._p.getnext()
    removed_caption = False
    while cursor is not None:
        next_cursor = cursor.getnext()
        tag = cursor.tag.split("}")[-1]
        if tag == "p":
            p = Paragraph(cursor, doc)
            if p.style.name.startswith("Heading"):
                break
            has_drawing = bool(cursor.xpath(".//*[local-name()='drawing']"))
            if has_drawing or p.text.strip() == "Hình 2. Mô hình phương pháp nghiên cứu.":
                cursor.getparent().remove(cursor)
                if p.text.strip().startswith("Hình 2."):
                    removed_caption = True
        cursor = next_cursor

    image_p = heading._parent.add_paragraph()
    image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_p.add_run().add_picture(str(METHOD_FIG), width=Inches(6.5))
    heading._p.addnext(image_p._p)

    caption = heading._parent.add_paragraph(
        "Hình 2. Mô hình triển khai cluster-aware forecasting theo thiết kế leakage-aware rolling-origin.",
        style="Caption",
    )
    image_p._p.addnext(caption._p)


def main() -> None:
    create_method_figure()
    doc = Document(DOC_PATH)

    # Introduction: add topic rationale after the first problem-formulation paragraph.
    intro = find_heading(doc, "Giới thiệu")
    existing_intro_text = "\n".join(p.text for p in doc.paragraphs[:20])
    if "Lý do chọn đề tài xuất phát" not in existing_intro_text:
        first_intro_para = Paragraph(intro._p.getnext(), doc)
        insert_after(first_intro_para, [(text, "Normal") for text in INTRO_APPEND])

    # Theory section: replace with explicit business and technical theory items.
    theory = find_heading(doc, "Cơ sở lý thuyết")
    remove_until_next_heading(theory)
    theory_entries: list[tuple[str, str]] = []
    for heading, body in THEORY_BLOCK:
        theory_entries.append((heading, "Heading 3"))
        theory_entries.append((body, "Normal"))
    insert_after(theory, theory_entries)

    # Add chapter-3 data description and EDA before the existing "Nền tảng nghiên cứu".
    existing_text = "\n".join(p.text for p in doc.paragraphs)
    if "Mô tả dữ liệu và EDA" not in existing_text:
        foundation_heading = find_heading(doc, "Nền tảng nghiên cứu")
        eda_heading = foundation_heading._parent.add_paragraph("Mô tả dữ liệu và EDA", style="Heading 2")
        foundation_heading._p.addprevious(eda_heading._p)
        last = insert_after(eda_heading, [(text, "Normal") for text in EDA_PARAGRAPHS])
        insert_table_after(last, EDA_TABLE)

    replace_method_figure(doc)

    for paragraph in doc.paragraphs:
        if paragraph.style.name == "Heading 1":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif paragraph.style.name in {"Heading 2", "Heading 3"}:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.save(DOC_PATH)
    OUTPUT_COPY.write_bytes(DOC_PATH.read_bytes())

    doc = Document(DOC_PATH)
    print(f"Updated: {DOC_PATH}")
    print(f"Copy: {OUTPUT_COPY}")
    print(f"Figure: {METHOD_FIG}")
    print(f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} inline_shapes={len(doc.inline_shapes)}")


if __name__ == "__main__":
    main()
