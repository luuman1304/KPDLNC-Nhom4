from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.text.paragraph import Paragraph


DOCX = Path("Document/Bai_Viet_Nghien_Cuu_Khoa_Hoc_Chi_Tiet_Hoc_Thuat_M5.docx")
MARKER = "Để làm rõ hơn đặc điểm dữ liệu, nghiên cứu bổ sung bảy biểu đồ EDA"


EDA_PARAGRAPHS = [
    (
        "Để làm rõ hơn đặc điểm dữ liệu, nghiên cứu bổ sung bảy biểu đồ EDA trong bản Word có biểu đồ. "
        "Các biểu đồ này không dùng để lựa chọn mô hình theo kết quả test, mà dùng để kiểm tra giả định dữ liệu trước khi huấn luyện: "
        "mức độ thưa của nhu cầu, khác biệt theo hierarchy, biến động giá và cấu trúc demand regimes. "
        "Đây là bước cần thiết vì dữ liệu M5 có quy mô lớn, nhiều chuỗi item-store có doanh số rất thấp, và nếu chỉ nhìn metric tổng hợp thì có thể bỏ qua nguy cơ sai lệch khi đánh giá stability."
    ),
    (
        "Hình EDA 1 mô tả phân bố các chuỗi theo bốn nhóm nhu cầu: intermittent, lumpy, smooth và erratic. "
        "Trong 30,490 chuỗi, có 22,175 chuỗi intermittent, tương đương khoảng 72.7% tổng số chuỗi. "
        "Nhóm lumpy có 5,580 chuỗi, nhóm smooth có 1,883 chuỗi, và nhóm erratic có 852 chuỗi. "
        "Kết quả này xác nhận bài toán không phải là một bài toán demand forecasting cân bằng giữa các chuỗi bán đều và bán thưa; phần lớn chuỗi có nhu cầu không liên tục. "
        "Vì vậy, việc dùng LightGBM với objective Tweedie là hợp lý hơn so với chỉ giả định nhiễu Gaussian, do Tweedie phù hợp với dữ liệu không âm, lệch phải và có nhiều giá trị 0."
    ),
    (
        "Hình EDA 2 biểu diễn phân phối zero-sales ratio. Median bằng 0.6337, phân vị 75% xấp xỉ 0.7960, cho thấy nhiều chuỗi có phần lớn ngày không bán được. "
        "Đây là bằng chứng trực tiếp cho quyết định không dùng các chỉ số stability dạng phần trăm thay đổi đơn giản. "
        "Với chuỗi có doanh số gần 0, một thay đổi nhỏ về forecast tuyệt đối có thể tạo ra tỷ lệ phần trăm rất lớn, làm stability bị phóng đại. "
        "Vì vậy, scale-aware stability loss trong nghiên cứu dùng mẫu số có scale floor để giảm sai lệch do near-zero demand."
    ),
    (
        "Hình EDA 3 là bản đồ ADI-CV2. ADI phản ánh khoảng cách trung bình giữa các lần phát sinh nhu cầu, còn CV2 phản ánh mức độ biến động của demand size khi có bán. "
        "Hai đường ngưỡng ADI 1.32 và CV2 0.49 được dùng để phân biệt các kiểu nhu cầu phổ biến trong intermittent demand literature. "
        "Biểu đồ cho thấy các chuỗi không nằm trong một vùng hành vi đồng nhất; có nhóm bán đều hơn, nhóm intermittent, nhóm biến động mạnh và nhóm vừa thưa vừa biến động. "
        "Điều này củng cố giả thuyết của proposal rằng một global model duy nhất có thể chưa đủ linh hoạt, và cluster-aware forecasting có cơ sở thực nghiệm."
    ),
    (
        "Hình EDA 4 so sánh tổng doanh số và median zero-sales ratio theo ngành hàng. "
        "FOODS chiếm phần lớn tổng doanh số và có zero-sales ratio thấp hơn so với HOBBIES và HOUSEHOLD. "
        "Điều này cho thấy hierarchy không chỉ là thông tin định danh, mà còn phản ánh khác biệt có ý nghĩa về demand regime. "
        "Vì vậy, các biến category, department, store, state và các thống kê theo hierarchy được giữ trong mô hình chính. "
        "Kết quả ablation sau đó cũng xác nhận nếu bỏ nhóm hierarchy/ID features thì WRMSSE tăng."
    ),
    (
        "Hình EDA 5 mô tả khác biệt ở cấp cửa hàng. CA_3 có tổng doanh số cao nhất và median zero-sales ratio thấp hơn nhiều cửa hàng khác, trong khi các cửa hàng như CA_4, TX_1 hoặc WI_3 có mức thưa lớn hơn. "
        "Khác biệt này giải thích vì sao mô hình phải học cả tín hiệu store/state thay vì chỉ học ở cấp item. "
        "Nó cũng cho thấy cần đánh giá WRMSSE theo weighted hierarchy, vì một sai số ở chuỗi/store có doanh số lớn có ảnh hưởng vận hành khác với sai số ở chuỗi long-tail."
    ),
    (
        "Hình EDA 6 trình bày đặc điểm giá bán theo cửa hàng. Median price giữa các cửa hàng khá gần nhau, nhưng max price có một số giá trị rất cao, đặc biệt ở WI_3. "
        "Điều này cho thấy price features có thể chứa tín hiệu hữu ích nhưng cũng có outlier. "
        "Do đó, nghiên cứu không kết luận vai trò của price chỉ từ EDA, mà kiểm tra bằng ablation no-price. "
        "Kết quả full-scale cho thấy bỏ price làm WRMSSE tăng nhẹ, tức price có đóng góp nhưng không phải nguồn cải thiện lớn nhất."
    ),
    (
        "Hình EDA 7 mô tả phân phối mean sales ở cấp chuỗi item-store. Phân phối lệch phải rõ rệt: đa số chuỗi có mean sales thấp, trong khi một số ít chuỗi có doanh số cao hơn đáng kể. "
        "Đặc điểm này giải thích vì sao mô hình global có lợi thế về chia sẻ thông tin giữa các chuỗi, nhưng đồng thời cũng tạo động lực cho cluster-specific modeling: "
        "nếu tất cả chuỗi long-tail và core products bị ép vào cùng một cơ chế học, mô hình có thể thiên về nhóm chiếm trọng số lớn hoặc nhóm có tín hiệu mạnh hơn."
    ),
    (
        "Tổng hợp các kết quả EDA, dữ liệu M5 trong nghiên cứu có bốn đặc điểm chính: nhu cầu thưa, phân phối doanh số lệch phải, khác biệt rõ theo hierarchy và tồn tại nhiều demand regimes. "
        "Các đặc điểm này phù hợp với định hướng của proposal: dùng global forecasting để tận dụng thông tin toàn bộ panel, dùng clustering để phân tách hành vi nhu cầu, "
        "dùng calendar/SNAP/event và hierarchy features để mô tả ngữ cảnh bán lẻ, và dùng scale-aware stability để tránh kết luận sai trên các chuỗi có doanh số gần 0."
    ),
]


def add_paragraph_before(ref_paragraph, text: str) -> None:
    element = ref_paragraph._element
    new_element = OxmlElement("w:p")
    element.addprevious(new_element)
    paragraph = Paragraph(new_element, ref_paragraph._parent)
    paragraph.style = "Normal"
    run = paragraph.add_run(text)
    run.font.size = Pt(11)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def main() -> None:
    doc = Document(DOCX)
    prefixes_to_remove = (MARKER, "Hình EDA ", "Tổng hợp các kết quả EDA")
    for paragraph in list(doc.paragraphs):
        if paragraph.text.startswith(prefixes_to_remove):
            paragraph._element.getparent().remove(paragraph._element)

    target = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "5. Phương pháp nghiên cứu":
            target = paragraph
            break
    if target is None:
        raise RuntimeError("Could not find section 5 heading in research docx")

    for text in EDA_PARAGRAPHS:
        add_paragraph_before(target, text)

    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
